import io
import zipfile
import json
import qrcode
import re
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

# ============================================================================
# KHU VỰC 0: CẤU HÌNH & HÀM DÙNG CHUNG (KẾ THỪA 100% TỪ KHTN)
# ============================================================================
FONT_NAME = "Times New Roman"
FONT_SIZE = 12
COEFF_TEXT = 0.55
PIXELS_TO_CM = 360000
PAGE_WIDTH_CM = 18.96 
WIDTH_COL_MADE = 2.0
THRESH_MODE_4 = 4.2  
THRESH_MODE_2 = 8.9  
SAFETY_FACTOR = 1.0  
COLOR_BLUE = RGBColor(0, 0, 255)
COLOR_PURPLE = RGBColor(112, 48, 160)
COLOR_PURPLE_TIT = RGBColor(93, 99, 211)
HEX_BG_PART1 = "7030A0"

NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main', 
         'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math', 
         'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing', 
         'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
         'v': 'urn:schemas-microsoft-com:vml'}

def get_tag(element): return element.tag.split('}')[-1]
def measure_text_cm(text): return len(text) * COEFF_TEXT * FONT_SIZE / 28.35 if text else 0

def count_math_chars_recursive(element, in_fraction=False):
    count = 0.0
    tag = get_tag(element)
    current_in_fraction = in_fraction or (tag == 'f')
    if tag == 't':
        text = element.text if element.text else ""
        length = len(text)
        if current_in_fraction: count += length / 2.0
        else: count += length
    for child in element: count += count_math_chars_recursive(child, current_in_fraction)
    return count

def measure_math_cm(math_element): return count_math_chars_recursive(math_element) * COEFF_TEXT * FONT_SIZE / 28.35

def measure_ole_object_cm(object_element):
    shapes = object_element.findall('.//v:shape', NSMAP)
    total_w = 0.0
    for shape in shapes:
        style = shape.get('style', '')
        if 'width:' in style:
            try:
                width_part = style.split('width:')[1].split(';')[0].strip()
                if 'pt' in width_part: val = float(width_part.replace('pt', '')); total_w += val / 28.35
                elif 'in' in width_part: val = float(width_part.replace('in', '')); total_w += val * 2.54
                elif 'cm' in width_part: val = float(width_part.replace('cm', '')); total_w += val
            except: pass
    return total_w

def measure_image_cm(drawing_element):
    extents = drawing_element.findall('.//wp:extent', NSMAP)
    max_w = 0
    for ext in extents:
        try:
            cx = int(ext.get('cx', 0)); w = cx / PIXELS_TO_CM
            if w > max_w: max_w = w
        except: pass
    return max_w

def calculate_cell_width(cell):
    total = 0.0
    for p in cell.paragraphs:
        for child in p._element:
            tag = get_tag(child)
            if tag == 'r':
                for t in child.findall('.//w:t', NSMAP): total += measure_text_cm(t.text)
                for d in child.findall('.//w:drawing', NSMAP): total += measure_image_cm(d)
                for obj in child.findall('.//w:object', NSMAP): total += measure_ole_object_cm(obj)
            elif tag in ['oMath', 'oMathPara']: total += measure_math_cm(child)
    return total

def calculate_option_width(cell): return calculate_cell_width(cell)

def format_run(run, bold=False, italic=False, color=None, size=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size if size else FONT_SIZE)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color

# [BỌC THÉP XML]
def set_cell_shading(cell, hex_str):
    if not hex_str: return
    tcPr = cell._element.get_or_add_tcPr()
    old_shd = tcPr.find(qn('w:shd'))
    if old_shd is not None: tcPr.remove(old_shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), str(hex_str)) 
    insert_idx = len(tcPr)
    for i, child in enumerate(tcPr):
        if child.tag in [qn('w:noWrap'), qn('w:tcMar'), qn('w:textDirection'), qn('w:tcFitText'), qn('w:vAlign'), qn('w:hideMark')]:
            insert_idx = i; break
    tcPr.insert(insert_idx, shd)

def add_table_borders(table):
    tbl = table._tbl; tblPr = tbl.tblPr
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None: tblPr.remove(old_borders)
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4'); border.set(qn('w:space'), '0'); border.set(qn('w:color'), 'auto')
        borders.append(border)
    insert_idx = len(tblPr)
    for i, child in enumerate(tblPr):
        if child.tag in [qn('w:shd'), qn('w:tblLayout'), qn('w:tblCellMar'), qn('w:tblLook'), qn('w:tblCaption')]:
            insert_idx = i; break
    tblPr.insert(insert_idx, borders)

def add_field(p, code, italic=False):
    r = p.add_run(); f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); r._r.append(f1)
    if italic: r.font.italic = True
    r = p.add_run(); i = OxmlElement('w:instrText'); i.text = code; r._r.append(i)
    if italic: r.font.italic = True
    r = p.add_run(); f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end'); r._r.append(f2)
    if italic: r.font.italic = True

def set_narrow_layout(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(1.02)
        s.page_width, s.page_height = Cm(21), Cm(29.7)

def safe_insert_element(doc, element): doc.element.body.insert(-1, deepcopy(element))

def copy_paragraph_content(source_p, target_p):
    if not source_p: return
    for child in source_p._element:
        if get_tag(child) in ['r', 'oMath', 'oMathPara', 'hyperlink', 'drawing', 'object']: 
            target_p._element.append(deepcopy(child))
    for run in target_p.runs:
        run.font.name = FONT_NAME; run.font.size = Pt(FONT_SIZE)

def is_p_empty(p):
    if p.text.strip(): return False
    for child in p._element.iter():
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ['oMath', 'oMathPara', 'drawing', 'object', 'pict', 'AlternateContent', 'shape']: return False 
    return True

def force_compact_p(p_obj):
    p_obj.paragraph_format.space_before = Pt(0)
    p_obj.paragraph_format.space_after = Pt(0)
    p_obj.paragraph_format.line_spacing = 1.0

def render_cell_content(doc, cell, first_line_prefix=None, first_line_color=None):
    if not cell: return
    elements = []
    for child in cell._element.iterchildren():
        tag = get_tag(child)
        if tag in ['p', 'tbl']: elements.append(child)
        
    last_valid_idx = -1
    for i in range(len(elements) - 1, -1, -1):
        if get_tag(elements[i]) == 'tbl':
            last_valid_idx = i; break
        elif get_tag(elements[i]) == 'p':
            p_obj = Paragraph(elements[i], cell)
            if not is_p_empty(p_obj):
                last_valid_idx = i; break
                
    valid_elements = elements[:last_valid_idx + 1]

    if not valid_elements:
        if first_line_prefix:
            new_p = doc.add_paragraph(); force_compact_p(new_p)
            doc.element.body.remove(new_p._element); doc.element.body.insert(-1, new_p._element)
            new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
            r = new_p.add_run(first_line_prefix); format_run(r, bold=True, color=first_line_color)
        return

    for i, elem in enumerate(valid_elements):
        tag = get_tag(elem)
        if tag == 'p':
            p = Paragraph(elem, cell)
            new_p = doc.add_paragraph(); force_compact_p(new_p)
            doc.element.body.remove(new_p._element); doc.element.body.insert(-1, new_p._element)
            if i == 0 and first_line_prefix:
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
                r = new_p.add_run(first_line_prefix); format_run(r, bold=True, color=first_line_color)
            elif p.alignment: new_p.alignment = p.alignment
            copy_paragraph_content(p, new_p)
        elif tag == 'tbl':
            new_tbl = deepcopy(elem)
            doc.element.body.insert(-1, new_tbl)
            if i == 0 and first_line_prefix:
                new_p = doc.add_paragraph(); force_compact_p(new_p)
                new_p._element.addprevious(new_p._element)
                r = new_p.add_run(first_line_prefix); format_run(r, bold=True, color=first_line_color)

def clear_body_preserve_section(doc):
    body = doc.element.body
    for child in list(body):
        if not child.tag.endswith('sectPr'): body.remove(child)

# ============================================================================
# KHU VỰC 1: HEADER, FOOTER, IMAGE (KẾ THỪA 100% TỪ KHTN)
# ============================================================================
def insert_header_and_student_info(doc, config, exam_id):
    if config.get('co_header', True):
        hd = config.get('header_data', {})
        tbl = doc.add_table(rows=2, cols=2); tbl.autofit = False
        tbl.columns[0].width, tbl.columns[1].width = Cm(8.0), Cm(10.96)
        
        c = tbl.cell(0,0); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; force_compact_p(p)
        format_run(p.add_run(f"{hd.get('so','').upper()}\n"), size=12, color=COLOR_BLUE)
        format_run(p.add_run(f"{hd.get('truong','').upper()}\n"), bold=True, size=12, color=COLOR_BLUE)
        format_run(p.add_run(hd.get('to_chuyen_mon','').upper()), bold=True, size=12, color=COLOR_BLUE)
        
        c = tbl.cell(0,1); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; force_compact_p(p)
        format_run(p.add_run(f"{hd.get('ky_thi','').upper()}\n"), bold=True, size=12, color=COLOR_BLUE)
        format_run(p.add_run(f"{hd.get('nam_hoc','')}\n"), size=12, color=COLOR_BLUE)
        format_run(p.add_run(f"Môn: {hd.get('mon','')}"), bold=True, color=COLOR_PURPLE_TIT)
        
        c = tbl.cell(1,0); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; force_compact_p(p)
        format_run(p.add_run("ĐỀ CHÍNH THỨC\n"), bold=True, color=COLOR_PURPLE_TIT)
        format_run(p.add_run("(Đề thi có "), italic=True); add_field(p, "NUMPAGES", italic=True); format_run(p.add_run(" trang)"), italic=True)
        
        c = tbl.cell(1,1); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; force_compact_p(p)
        t_val = "".join(re.findall(r'\d+', str(hd.get('thoi_gian','')))) or "..."
        p.add_run(f"Thời gian làm bài: {t_val} phút\n"); format_run(p.add_run("(Không kể thời gian phát đề)"), italic=True, size=12)
        
        for r in tbl.rows:
            for cl in r.cells: 
                tcPr = cl._element.get_or_add_tcPr(); tcB = OxmlElement('w:tcBorders')
                for b in ['top', 'left', 'bottom', 'right']: e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'nil'); tcB.append(e)
                tcPr.append(tcB)
                
        p_blank = doc.add_paragraph(); force_compact_p(p_blank)
    
    tbl_i = doc.add_table(rows=1, cols=2); tbl_i.autofit = False; tbl_i.allow_autofit = False
    tbl_i.columns[0].width = Cm(15.46); tbl_i.columns[1].width = Cm(3.5)
    tbl_i.cell(0,0).width = Cm(15.46); tbl_i.cell(0,1).width = Cm(3.5)
    
    p = tbl_i.cell(0,0).paragraphs[0]; force_compact_p(p)
    format_run(p.add_run("Họ và tên: ...................................................................... SBD: .........................................."), bold=True, color=COLOR_BLUE)
    
    c = tbl_i.cell(0,1); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; force_compact_p(p)
    format_run(p.add_run(f"Mã đề thi: {exam_id}"), bold=True, color=COLOR_PURPLE)
    
    tcPr0 = tbl_i.cell(0,0)._element.get_or_add_tcPr(); tcB0 = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']: e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'nil'); tcB0.append(e)
    tcPr0.append(tcB0)
    tcPr1 = c._element.get_or_add_tcPr(); tcB1 = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']: e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'double'); e.set(qn('w:sz'), '6'); tcB1.append(e)
    tcPr1.append(tcB1)
    
def insert_common_footer(doc, exam_id):
    sect = doc.sections[0]; sect.footer.is_linked_to_previous = False
    for p in sect.footer.paragraphs: p._element.getparent().remove(p._element)
    p = sect.footer.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    force_compact_p(p) 
    r = p.add_run("Trang "); format_run(r, bold=True, color=COLOR_PURPLE)
    add_field(p, "PAGE"); r = p.add_run("/"); format_run(r, bold=True, color=COLOR_PURPLE)
    add_field(p, "NUMPAGES"); r = p.add_run(f" - Mã đề {exam_id}"); format_run(r, bold=True, color=COLOR_PURPLE)

def insert_signature_footer(doc, config):
    if not config.get('co_footer', True): return
    def add_sig_line(title, is_first=False):
        p = doc.add_paragraph(); doc.element.body.remove(p._element); doc.element.body.insert(-1, p._element)
        force_compact_p(p) 
        if is_first: p.paragraph_format.space_before = Pt(6) 
        else: p.paragraph_format.space_before = Pt(6)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Cm(9.48), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS)
        tabs.add_tab_stop(Cm(18.96), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(f"{title}\t Ký tên:\t"); format_run(r, size=12)
    add_sig_line("Giám thị 1:", is_first=True); add_sig_line("Giám thị 2:", is_first=False)

def insert_images(doc, config):
    img_tn = config.get('img_phieu_to'); img_tl = config.get('img_tu_luan')
    co_header = config.get('co_header', True)
    def add_pic(img, is_page_1=True):
        if not img: return
        try:
            from PIL import Image as PILImage
            img.seek(0)
            with PILImage.open(img) as pi: w_orig, h_orig = pi.size
            max_w = 18.96
            if is_page_1: max_h = 22.4 if co_header else 25.8
            else: max_h = 27.4 
            ratio_w = max_w / w_orig; ratio_h = max_h / h_orig; ratio = min(ratio_w, ratio_h) 
            final_w = w_orig * ratio; final_h = h_orig * ratio
            
            img.seek(0); p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            from docx.enum.text import WD_BREAK
            r = p.add_run()
            if not is_page_1:
                r.add_break(WD_BREAK.PAGE); p.paragraph_format.space_before = Pt(0)
            else: p.paragraph_format.space_before = Pt(12) 
            p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
            r.add_picture(img, width=Cm(final_w), height=Cm(final_h))
        except Exception as e: 
            pass
    if img_tn: add_pic(img_tn, is_page_1=True)
    if img_tl: add_pic(img_tl, is_page_1=False)

def insert_quy_uoc(doc, config):
    file_qu = config.get('file_quy_uoc')
    if not file_qu: return
    try:
        from docx import Document
        from docx.table import Table
        import io
        file_qu.seek(0)
        doc_qu = Document(io.BytesIO(file_qu.read()))
        for element in doc_qu.element.body.iterchildren():
            tag = get_tag(element)
            if tag in ['p', 'tbl']:
                new_element = deepcopy(element)
                bad_tags = ['drawing', 'object', 'hyperlink', 'pict']
                for bad_tag in bad_tags:
                    for bad_elem in new_element.findall(f'.//w:{bad_tag}', NSMAP):
                        parent = bad_elem.getparent()
                        if parent is not None: parent.remove(bad_elem)
                doc.element.body.insert(-1, new_element)
                try:
                    if tag == 'p':
                        p = Paragraph(new_element, doc)
                        force_compact_p(p)
                        for run in p.runs: run.font.name = FONT_NAME; run.font.size = Pt(FONT_SIZE)
                    elif tag == 'tbl':
                        tbl = Table(new_element, doc)
                        for row in tbl.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    force_compact_p(p)
                                    for run in p.runs: run.font.name = FONT_NAME; run.font.size = Pt(FONT_SIZE)
                except: pass 
    except Exception as e:
        p_err = doc.add_paragraph()
        r_err = p_err.add_run(f"[HỆ THỐNG BÁO LỖI CHÈN QUY ƯỚC: {str(e)}]")
        r_err.font.color.rgb = RGBColor(255, 0, 0)

# ============================================================================
# KHU VỰC 2: RENDER CÂU HỎI TIẾNG ANH (RENDER OPTIONS)
# ============================================================================
def determine_mode(options_data):
    for opt in options_data:
        cell = opt['cell']
        if not cell: continue
        valid_count = 0
        for child in cell._element.iterchildren():
            tag = get_tag(child)
            if tag == 'tbl': return 1
            elif tag == 'p':
                p = Paragraph(child, cell)
                if not is_p_empty(p): valid_count += 1
        if valid_count > 1: return 1 

    max_w = 0
    for opt in options_data:
        w = calculate_option_width(opt['cell']) + 0.1 
        if w > max_w: max_w = w
    final_w = max_w * SAFETY_FACTOR
    if final_w <= THRESH_MODE_4: return 4
    if final_w <= THRESH_MODE_2: return 2
    return 1

def render_options_with_tabs(doc, options_data):
    if not options_data: return
    mode = determine_mode(options_data)
    
    def create_p():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(0.5) 
        return p
        
    def copy_cell_to_paragraph_inline(source_cell, target_p):
        if not source_cell: return
        for child in source_cell._element.iterchildren():
            if get_tag(child) == 'p':
                p = Paragraph(child, source_cell)
                if not is_p_empty(p):
                    copy_paragraph_content(p, target_p); break
                
    if mode == 4:
        p = create_p(); ts = p.paragraph_format.tab_stops
        ts.add_tab_stop(Cm(4.74)); ts.add_tab_stop(Cm(9.48)); ts.add_tab_stop(Cm(14.22))
        for i in range(4):
            if i > 0: p.add_run("\t")
            # [ÉP MÀU XANH TẠI ĐÂY]
            r = p.add_run(options_data[i]['lbl'] + " "); format_run(r, bold=True, color=COLOR_BLUE)
            copy_cell_to_paragraph_inline(options_data[i]['cell'], p)
    elif mode == 2:
        p1 = create_p(); p1.paragraph_format.tab_stops.add_tab_stop(Cm(9.48)) 
        # [ÉP MÀU XANH TẠI ĐÂY]
        r = p1.add_run(options_data[0]['lbl'] + " "); format_run(r, bold=True, color=COLOR_BLUE)
        copy_cell_to_paragraph_inline(options_data[0]['cell'], p1); p1.add_run("\t")
        r = p1.add_run(options_data[1]['lbl'] + " "); format_run(r, bold=True, color=COLOR_BLUE)
        copy_cell_to_paragraph_inline(options_data[1]['cell'], p1)
        
        p2 = create_p(); p2.paragraph_format.tab_stops.add_tab_stop(Cm(9.48)) 
        # [ÉP MÀU XANH TẠI ĐÂY]
        r = p2.add_run(options_data[2]['lbl'] + " "); format_run(r, bold=True, color=COLOR_BLUE)
        copy_cell_to_paragraph_inline(options_data[2]['cell'], p2); p2.add_run("\t")
        r = p2.add_run(options_data[3]['lbl'] + " "); format_run(r, bold=True, color=COLOR_BLUE)
        copy_cell_to_paragraph_inline(options_data[3]['cell'], p2)
    else:
        for opt in options_data:
            start_p_count = len(doc.paragraphs)
            # [ÉP MÀU XANH TẠI ĐÂY]
            render_cell_content(doc, opt['cell'], first_line_prefix=opt['lbl'] + " ", first_line_color=COLOR_BLUE)
            for i in range(start_p_count, len(doc.paragraphs)):
                doc.paragraphs[i].paragraph_format.left_indent = Cm(0.5)

# ============================================================================
# KHU VỰC 3: MÁY QUÉT ĐỌC DỮ LIỆU ĐÃ CHUẨN HÓA (DÀNH RIÊNG CHO AV/DGNL)
# ============================================================================
def extract_data_from_raw_av(doc_raw):
    """
    Parser độc quyền cho môn Tiếng Anh: Nhận diện bảng từ xuly_degoc_av.py
    """
    items = []
    current_q = None
    q_counter = 1

    def flush_q():
        nonlocal current_q, q_counter
        if current_q:
            # Gán nhãn cho nó là TN hay TL
            if not current_q['options']: current_q['q_type'] = 'TL'
            else: current_q['q_type'] = 'TN'
            
            # Lưu lại ID in ra
            current_q['print_id'] = q_counter
            q_counter += 1
            
            items.append(current_q)
            current_q = None

    for block in doc_raw.element.body.iterchildren():
        tag = get_tag(block)
        if tag == 'p':
            txt = "".join([t.text for t in block.findall('.//w:t', NSMAP) if t.text]).strip()
            if txt:
                clean_txt = re.sub(r'\s+', '', txt)
                # Phát hiện thẻ nhóm -> Nhắm mắt làm ngơ (Tàng hình khi in)
                if ("@" in clean_txt and "#" in clean_txt) or ("*" in clean_txt and "#" in clean_txt):
                    continue
                flush_q()
                items.append({'type': 'TEXT', 'obj': deepcopy(block)})
                
        elif tag == 'tbl':
            tbl = Table(block, doc_raw)
            c0_text = tbl.rows[0].cells[0].text.strip().upper()

            if c0_text == 'P': # Bảng chứa đoạn văn Reading (Context)
                flush_q()
                items.append({'type': 'CONTEXT', 'cell': tbl.rows[0].cells[1]})
                
            elif "CÂU" in c0_text or "QUESTION" in c0_text:
                flush_q()
                current_q = {
                    'type': 'QUESTION', 'label': c0_text, 
                    'stem': tbl.rows[0].cells[1], 'options': [], 
                    'key': '', 'solution': None
                }
                # Quét các dòng còn lại để lấy A, B, C, D và Lời giải
                for row in tbl.rows[1:]:
                    r0_txt = row.cells[0].text.strip().upper()
                    if r0_txt in ['A', 'B', 'C', 'D', 'A.', 'B.', 'C.', 'D.']:
                        lbl = r0_txt[:1] + '.'
                        current_q['options'].append({'lbl': lbl, 'cell': row.cells[1]})
                    elif "KEY" in r0_txt:
                        # Rút trích nội dung KEY từ ổ cell[1] (với câu TN) hoặc từ chính cell[0] (với TL)
                        if "KEY:" in r0_txt:
                            current_q['key'] = r0_txt.replace("KEY:", "").strip()
                            if len(row.cells) > 1: current_q['solution'] = row.cells[1]
                        else:
                            current_q['key'] = row.cells[1].text.strip()
                            current_q['solution'] = row.cells[1]
                    elif "LỜI GIẢI" in r0_txt or "SOL" in r0_txt:
                        if len(row.cells) > 1: current_q['solution'] = row.cells[1]
                        else: current_q['solution'] = row.cells[0]
                        
            elif "KEY" in c0_text:
                # Trường hợp bảng Tự luận gộp (Row 1: Câu, Row 2: KEY + Giải)
                if current_q:
                    current_q['key'] = c0_text.replace("KEY:", "").strip()
                    if len(tbl.rows[0].cells) > 1:
                        current_q['solution'] = tbl.rows[0].cells[1]
                        
    flush_q()
    return items


# ============================================================================
# KHU VỰC 4: SINH FILE WORD
# ============================================================================
def generate_made_file_av(doc, items, config, exam_id):
    insert_header_and_student_info(doc, config, exam_id); insert_images(doc, config)
    insert_quy_uoc(doc, config)
    
    # --- CÔNG TẮC SONG NGỮ ---
    is_eng = config.get('loai_mon') in ['ENG', 'AV']
    lbl_end = "---------------THE END-----------------" if is_eng else "---------------HẾT-----------------"
    
    for item in items:
        if item['type'] == 'TEXT':
            new_element = deepcopy(item['obj'])
            safe_insert_element(doc, new_element)
            if get_tag(new_element) == 'p':
                p = Paragraph(new_element, doc)
                force_compact_p(p)
                for run in p.runs:
                    run.font.name = FONT_NAME; run.font.size = Pt(FONT_SIZE)
                    
        elif item['type'] == 'CONTEXT':
            # Render đoạn văn Reading
            render_cell_content(doc, item['cell'])
            
        elif item['type'] == 'QUESTION':
            q_label = f"Question {item['print_id']}. " if is_eng else f"Câu {item['print_id']}. "
            render_cell_content(doc, item['stem'], first_line_prefix=q_label, first_line_color=COLOR_BLUE)
            render_options_with_tabs(doc, item['options'])
            
    p_end = doc.add_paragraph(lbl_end)
    doc.element.body.remove(p_end._element); doc.element.body.insert(-1, p_end._element)
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER; format_run(p_end.runs[0], bold=True)
    p_end.paragraph_format.space_after = Pt(6)
    insert_signature_footer(doc, config); insert_common_footer(doc, exam_id)

def generate_dapan_tonghop_av(doc_base, all_results, config):
    if not all_results: return
    
    clear_body_preserve_section(doc_base)
    style = doc_base.styles['Normal']; font = style.font
    font.name = FONT_NAME; font.size = Pt(FONT_SIZE)

    # --- CÔNG TẮC SONG NGỮ ---
    is_eng = config.get('loai_mon') in ['ENG', 'AV']
    lbl_title = "SUMMARY ANSWER SHEET" if is_eng else "BẢNG ĐÁP ÁN TỔNG HỢP"
    lbl_exam = "CODE" if is_eng else "MÃ ĐỀ"
    lbl_tn = "MULTIPLE CHOICE" if is_eng else "PHẦN TRẮC NGHIỆM"
    lbl_tl = "WRITTEN / SHORT ANSWER" if is_eng else "PHẦN TỰ LUẬN / TRẢ LỜI NGẮN"
    lbl_q = "Question" if is_eng else "Câu"
    lbl_key = "Key" if is_eng else "Đáp án"

    p_title1 = doc_base.add_paragraph(); p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title1.paragraph_format.space_after = Pt(6)
    format_run(p_title1.add_run(lbl_title), bold=True, size=14, color=COLOR_BLUE)

    hd = config.get('header_data', {})
    ky_thi = hd.get('ky_thi', '...').upper()
    nam_hoc = hd.get('nam_hoc', '...')
    p_title2 = doc_base.add_paragraph(); p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title2.paragraph_format.space_after = Pt(12)
    format_run(p_title2.add_run(f"{ky_thi}, {nam_hoc}"), bold=True, size=13)

    def apply_compact_cell(cell, align_center=True):
        for p in cell.paragraphs:
            force_compact_p(p)
            if align_center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
    def set_col_widths(tbl, col_widths):
        tbl.autofit = False; tbl.allow_autofit = False
        for row in tbl.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells): row.cells[idx].width = width

    def clear_cell_borders_and_shading(cell):
        tcPr = cell._element.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil') 
            tcB.append(border)
        old_borders = tcPr.find(qn('w:tcBorders'))
        if old_borders is not None: tcPr.remove(old_borders)
        tcPr.append(tcB)
        
        old_shd = tcPr.find(qn('w:shd'))
        if old_shd is not None: tcPr.remove(old_shd)

    HEX_BG_TN = "7030A0"  
    HEX_BG_TL = "C65911"     

    import math
    sorted_results = sorted(all_results, key=lambda x: str(x['exam_id']))
    
    for res_idx, res in enumerate(sorted_results):
        exam_id = res['exam_id']
        
        if res_idx > 0:
            p_break = doc_base.add_paragraph()
            p_break.paragraph_format.space_before = Pt(12)
            p_break.paragraph_format.space_after = Pt(12)
            p_break.add_run("=========================================").font.color.rgb = RGBColor(200, 200, 200)
            p_break.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        p_exam = doc_base.add_paragraph()
        p_exam.paragraph_format.space_after = Pt(6)
        format_run(p_exam.add_run(f"{lbl_exam}: {exam_id}"), bold=True, size=14, color=RGBColor(255, 0, 0))

        tn_items = []
        tl_items = []
        for item in res['items']:
            if item['type'] == 'QUESTION':
                if item.get('q_type') == 'TN': tn_items.append(item)
                else: tl_items.append(item)

        if tn_items:
            p_tn = doc_base.add_paragraph()
            p_tn.paragraph_format.space_after = Pt(6)
            format_run(p_tn.add_run(lbl_tn), bold=True, color=COLOR_BLUE)

            chunk_size = 16  
            col_w = Cm(1.18) 

            for chunk_start in range(0, len(tn_items), chunk_size):
                if chunk_start > 0:
                    spacer = doc_base.add_paragraph()
                    spacer.add_run(" ") 
                    spacer.paragraph_format.space_before = Pt(6)
                    spacer.paragraph_format.space_after = Pt(6)
                    spacer.paragraph_format.line_spacing = 1.0
                
                chunk_q = tn_items[chunk_start:chunk_start + chunk_size]
                num_cols = len(chunk_q)
                
                tbl = doc_base.add_table(rows=2, cols=num_cols)
                add_table_borders(tbl)
                from docx.enum.table import WD_TABLE_ALIGNMENT
                tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                
                widths = [col_w] * num_cols
                set_col_widths(tbl, widths)

                r0 = tbl.rows[0]
                r1 = tbl.rows[1]

                for i, q in enumerate(chunk_q):
                    c0 = r0.cells[i]; c0.text = str(q.get('print_id', ''))
                    set_cell_shading(c0, HEX_BG_TN); apply_compact_cell(c0)
                    format_run(c0.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))

                    key_text = str(q.get('key', ''))
                    c1 = r1.cells[i]; c1.text = key_text; apply_compact_cell(c1)
                    if c1.paragraphs[0].runs: format_run(c1.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE, size=11)
                    else: r = c1.paragraphs[0].add_run(key_text); format_run(r, bold=True, color=COLOR_BLUE, size=11)

        if tl_items:
            p_tl = doc_base.add_paragraph()
            p_tl.paragraph_format.space_before = Pt(12)
            p_tl.paragraph_format.space_after = Pt(6)
            format_run(p_tl.add_run(lbl_tl), bold=True, color=COLOR_BLUE)

            n_tl = len(tl_items)
            mid = math.ceil(n_tl / 2)
            left_col = tl_items[:mid]
            right_col = tl_items[mid:]

            tbl = doc_base.add_table(rows=1, cols=5)
            add_table_borders(tbl)
            from docx.enum.table import WD_TABLE_ALIGNMENT
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            widths = [Cm(1.5), Cm(7.2), Cm(1.0), Cm(1.5), Cm(7.2)]
            set_col_widths(tbl, widths)
            
            r0 = tbl.rows[0]
            
            c0 = r0.cells[0]; c0.text = lbl_q; set_cell_shading(c0, HEX_BG_TL); apply_compact_cell(c0)
            format_run(c0.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))
            c1 = r0.cells[1]; c1.text = lbl_key; set_cell_shading(c1, HEX_BG_TL); apply_compact_cell(c1)
            format_run(c1.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))
            
            c2 = r0.cells[2]; c2.text = ""; clear_cell_borders_and_shading(c2); apply_compact_cell(c2)
            
            if right_col:
                c3 = r0.cells[3]; c3.text = lbl_q; set_cell_shading(c3, HEX_BG_TL); apply_compact_cell(c3)
                format_run(c3.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))
                c4 = r0.cells[4]; c4.text = lbl_key; set_cell_shading(c4, HEX_BG_TL); apply_compact_cell(c4)
                format_run(c4.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))
            else:
                c3 = r0.cells[3]; c3.text = ""; clear_cell_borders_and_shading(c3); apply_compact_cell(c3)
                c4 = r0.cells[4]; c4.text = ""; clear_cell_borders_and_shading(c4); apply_compact_cell(c4)

            for i in range(mid):
                row = tbl.add_row(); set_col_widths(tbl, widths)
                
                q_left = left_col[i]
                rc0 = row.cells[0]; rc0.text = str(q_left.get('print_id', '')); apply_compact_cell(rc0)
                format_run(rc0.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE)
                
                key_left = str(q_left.get('key', '')).lower()
                rc1 = row.cells[1]; rc1.text = key_left; apply_compact_cell(rc1, align_center=False)
                if rc1.paragraphs[0].runs: format_run(rc1.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE)
                else: r = rc1.paragraphs[0].add_run(key_left); format_run(r, bold=True, color=COLOR_BLUE)
                
                rc2 = row.cells[2]; rc2.text = ""; clear_cell_borders_and_shading(rc2); apply_compact_cell(rc2)
                
                if i < len(right_col):
                    q_right = right_col[i]
                    rc3 = row.cells[3]; rc3.text = str(q_right.get('print_id', '')); apply_compact_cell(rc3)
                    format_run(rc3.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE)
                    
                    key_right = str(q_right.get('key', '')).lower()
                    rc4 = row.cells[4]; rc4.text = key_right; apply_compact_cell(rc4, align_center=False)
                    if rc4.paragraphs[0].runs: format_run(rc4.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE)
                    else: r = rc4.paragraphs[0].add_run(key_right); format_run(r, bold=True, color=COLOR_BLUE)
                else:
                    rc3 = row.cells[3]; rc3.text = ""; clear_cell_borders_and_shading(rc3); apply_compact_cell(rc3)
                    rc4 = row.cells[4]; rc4.text = ""; clear_cell_borders_and_shading(rc4); apply_compact_cell(rc4)
    # ============================================================================
    # 3. TẠO MÃ QR CHẤM THI (CHỈ TẠO NẾU 100% LÀ TRẮC NGHIỆM)
    # ============================================================================
    has_tl = any(item.get('q_type') == 'TL' for res in sorted_results for item in res['items'] if item['type'] == 'QUESTION')
    
    if not has_tl:
        try:
            tnmaker_data = {}; qm_data = []
            
            for res in sorted_results:
                exam_id = res['exam_id']
                # Rút trích đáp án Trắc nghiệm
                tn_keys = [str(item.get('key', '')).strip().upper() for item in res['items'] if item['type'] == 'QUESTION' and item.get('q_type') == 'TN']
                tn_str = "".join(tn_keys)
                
                # Môn AV/DGNL không có Đúng/Sai và Trả lời ngắn, nên đuôi để trống "##"
                tnmaker_data[str(exam_id)] = f"{tn_str}##"
                qm_data.append([str(exam_id)] + list(tn_str))

            tnmaker_data["success"] = True; tnmaker_data["type"] = 5
            tnmaker_json = json.dumps(tnmaker_data, ensure_ascii=False, separators=(',', ':'))
            qm_json = json.dumps(qm_data, ensure_ascii=False, separators=(',', ':'))

            if qm_data:
                p_qr_title = doc_base.add_paragraph()
                p_qr_title.paragraph_format.space_before = Pt(18); p_qr_title.paragraph_format.space_after = Pt(6)
                format_run(p_qr_title.add_run("MÃ QR NHẬP ĐÁP ÁN TRẮC NGHIỆM"), bold=True, size=12, color=COLOR_BLUE)
                
                tbl_qr = doc_base.add_table(rows=1, cols=2)
                tbl_qr.autofit = False
                tbl_qr.columns[0].width = Cm(9.48); tbl_qr.columns[1].width = Cm(9.48)
                
                # Cấu hình bảng ẩn (xóa viền) - ĐÃ SỬA LỖI CHẾT FILE
                for row in tbl_qr.rows:
                    for cell in row.cells:
                        tcPr = cell._element.get_or_add_tcPr(); tcB = OxmlElement('w:tcBorders')
                        # CHỈ DÙNG 4 CẠNH CHO Ô (CELL). Bỏ insideH và insideV
                        for b in ['top', 'left', 'bottom', 'right']:
                            e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'nil'); tcB.append(e)
                        tcPr.append(tcB)

                def add_qr_to_cell(cell, data_str, title):
                    qr = qrcode.QRCode(version=1, box_size=4, border=2); qr.add_data(data_str); qr.make(fit=True)
                    img = qr.make_image(fill_color="black", back_color="white"); img_io = io.BytesIO(); img.save(img_io, format='PNG'); img_io.seek(0)
                    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    format_run(p.add_run(f"App: {title}\n"), bold=True, size=12, color=COLOR_PURPLE)
                    p.add_run().add_picture(img_io, width=Cm(4.5))

                add_qr_to_cell(tbl_qr.cell(0,0), tnmaker_json, "TNMaker")
                add_qr_to_cell(tbl_qr.cell(0,1), qm_json, "QM")
        except Exception as e:
            doc_base.add_paragraph().add_run(f"[LỖI TẠO QR CODE: {str(e)}]").font.color.rgb = RGBColor(255, 0, 0)

def generate_dapan_tonghop_av(doc_base, all_results, config):
    if not all_results: return
    
    clear_body_preserve_section(doc_base)
    style = doc_base.styles['Normal']; font = style.font
    font.name = FONT_NAME; font.size = Pt(FONT_SIZE)

    p_title1 = doc_base.add_paragraph(); p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title1.paragraph_format.space_after = Pt(6)
    format_run(p_title1.add_run("BẢNG ĐÁP ÁN TỔNG HỢP"), bold=True, size=14, color=COLOR_BLUE)

    hd = config.get('header_data', {})
    ky_thi = hd.get('ky_thi', '...').upper()
    nam_hoc = hd.get('nam_hoc', '...')
    p_title2 = doc_base.add_paragraph(); p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title2.paragraph_format.space_after = Pt(12)
    format_run(p_title2.add_run(f"{ky_thi}, {nam_hoc}"), bold=True, size=13)

    # --- CÁC HÀM HỖ TRỢ ĐỊNH DẠNG ---
    def apply_compact_cell(cell, align_center=True):
        for p in cell.paragraphs:
            force_compact_p(p)
            if align_center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
    def set_col_widths(tbl, col_widths):
        tbl.autofit = False; tbl.allow_autofit = False
        for row in tbl.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells): row.cells[idx].width = width

    def clear_cell_borders_and_shading(cell):
        tcPr = cell._element.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil') 
            tcB.append(border)
        old_borders = tcPr.find(qn('w:tcBorders'))
        if old_borders is not None: tcPr.remove(old_borders)
        tcPr.append(tcB)
        
        old_shd = tcPr.find(qn('w:shd'))
        if old_shd is not None: tcPr.remove(old_shd)

    HEX_BG_TN = "7030A0"  
    HEX_BG_TL = "C65911"     

    import math
    import json
    import qrcode
    
    sorted_results = sorted(all_results, key=lambda x: str(x['exam_id']))
    
    for res_idx, res in enumerate(sorted_results):
        exam_id = res['exam_id']
        
        # --- [TUYỆT KỸ CHỐNG WORD GỘP BẢNG CỦA 2 MÃ ĐỀ] ---
        if res_idx > 0:
            p_break = doc_base.add_paragraph()
            p_break.paragraph_format.space_before = Pt(12)
            p_break.paragraph_format.space_after = Pt(12)
            p_break.add_run("=========================================").font.color.rgb = RGBColor(200, 200, 200)
            p_break.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        p_exam = doc_base.add_paragraph()
        p_exam.paragraph_format.space_after = Pt(6)
        format_run(p_exam.add_run(f"MÃ ĐỀ: {exam_id}"), bold=True, size=14, color=RGBColor(255, 0, 0))

        tn_items = []
        tl_items = []
        for item in res['items']:
            if item['type'] == 'QUESTION':
                if item.get('q_type') == 'TN': tn_items.append(item)
                else: tl_items.append(item)

        # ----------------------------------------------------
        # 1. BẢNG TRẮC NGHIỆM (NẰM NGANG)
        # ----------------------------------------------------
        if tn_items:
            p_tn = doc_base.add_paragraph()
            p_tn.paragraph_format.space_after = Pt(6)
            format_run(p_tn.add_run("PHẦN TRẮC NGHIỆM"), bold=True, color=COLOR_BLUE)

            chunk_size = 16  
            col_w = Cm(1.18) 

            for chunk_start in range(0, len(tn_items), chunk_size):
                if chunk_start > 0:
                    spacer = doc_base.add_paragraph()
                    spacer.add_run(" ") 
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.line_spacing = 1.0
                
                chunk_q = tn_items[chunk_start:chunk_start + chunk_size]
                num_cols = len(chunk_q)
                
                tbl = doc_base.add_table(rows=2, cols=num_cols)
                add_table_borders(tbl)
                from docx.enum.table import WD_TABLE_ALIGNMENT
                tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
                
                widths = [col_w] * num_cols
                set_col_widths(tbl, widths)

                r0 = tbl.rows[0]
                r1 = tbl.rows[1]

                for i, q in enumerate(chunk_q):
                    c0 = r0.cells[i]; c0.text = str(q.get('print_id', ''))
                    set_cell_shading(c0, HEX_BG_TN); apply_compact_cell(c0)
                    format_run(c0.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))

                    key_text = str(q.get('key', ''))
                    c1 = r1.cells[i]; c1.text = key_text; apply_compact_cell(c1)
                    if c1.paragraphs[0].runs: format_run(c1.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE, size=11)
                    else: r = c1.paragraphs[0].add_run(key_text); format_run(r, bold=True, color=COLOR_BLUE, size=11)

        # ----------------------------------------------------
        # 2. BẢNG TỰ LUẬN (NẰM DỌC & GẤP ĐÔI CỘT)
        # ----------------------------------------------------
        if tl_items:
            p_tl = doc_base.add_paragraph()
            p_tl.paragraph_format.space_before = Pt(12)
            p_tl.paragraph_format.space_after = Pt(6)
            format_run(p_tl.add_run("PHẦN TỰ LUẬN / TRẢ LỜI NGẮN"), bold=True, color=COLOR_BLUE)

            n_tl = len(tl_items)
            mid = math.ceil(n_tl / 2)
            left_col = tl_items[:mid]
            right_col = tl_items[mid:]

            tbl = doc_base.add_table(rows=1, cols=5)
            add_table_borders(tbl)
            from docx.enum.table import WD_TABLE_ALIGNMENT
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            widths = [Cm(1.5), Cm(7.2), Cm(1.0), Cm(1.5), Cm(7.2)]
            set_col_widths(tbl, widths)
            
            r0 = tbl.rows[0]
            
            c0 = r0.cells[0]; c0.text = "Câu"; set_cell_shading(c0, HEX_BG_TL); apply_compact_cell(c0)
            format_run(c0.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))
            c1 = r0.cells[1]; c1.text = "Đáp án"; set_cell_shading(c1, HEX_BG_TL); apply_compact_cell(c1)
            format_run(c1.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))
            
            c2 = r0.cells[2]; c2.text = ""; clear_cell_borders_and_shading(c2); apply_compact_cell(c2)
            
            if right_col:
                c3 = r0.cells[3]; c3.text = "Câu"; set_cell_shading(c3, HEX_BG_TL); apply_compact_cell(c3)
                format_run(c3.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))
                c4 = r0.cells[4]; c4.text = "Đáp án"; set_cell_shading(c4, HEX_BG_TL); apply_compact_cell(c4)
                format_run(c4.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))
            else:
                c3 = r0.cells[3]; c3.text = ""; clear_cell_borders_and_shading(c3); apply_compact_cell(c3)
                c4 = r0.cells[4]; c4.text = ""; clear_cell_borders_and_shading(c4); apply_compact_cell(c4)

            for i in range(mid):
                row = tbl.add_row(); set_col_widths(tbl, widths)
                
                q_left = left_col[i]
                rc0 = row.cells[0]; rc0.text = str(q_left.get('print_id', '')); apply_compact_cell(rc0)
                format_run(rc0.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE)
                
                key_left = str(q_left.get('key', '')).lower()
                rc1 = row.cells[1]; rc1.text = key_left; apply_compact_cell(rc1, align_center=False)
                if rc1.paragraphs[0].runs: format_run(rc1.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE)
                else: r = rc1.paragraphs[0].add_run(key_left); format_run(r, bold=True, color=COLOR_BLUE)
                
                rc2 = row.cells[2]; rc2.text = ""; clear_cell_borders_and_shading(rc2); apply_compact_cell(rc2)
                
                if i < len(right_col):
                    q_right = right_col[i]
                    rc3 = row.cells[3]; rc3.text = str(q_right.get('print_id', '')); apply_compact_cell(rc3)
                    format_run(rc3.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE)
                    
                    key_right = str(q_right.get('key', '')).lower()
                    rc4 = row.cells[4]; rc4.text = key_right; apply_compact_cell(rc4, align_center=False)
                    if rc4.paragraphs[0].runs: format_run(rc4.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE)
                    else: r = rc4.paragraphs[0].add_run(key_right); format_run(r, bold=True, color=COLOR_BLUE)
                else:
                    rc3 = row.cells[3]; rc3.text = ""; clear_cell_borders_and_shading(rc3); apply_compact_cell(rc3)
                    rc4 = row.cells[4]; rc4.text = ""; clear_cell_borders_and_shading(rc4); apply_compact_cell(rc4)

    # ============================================================================
    # 3. TẠO MÃ QR CHẤM THI (CHỈ TẠO NẾU 100% LÀ TRẮC NGHIỆM)
    # ============================================================================
    has_tl = any(item.get('q_type') == 'TL' for res in sorted_results for item in res['items'] if item['type'] == 'QUESTION')
    
    if not has_tl:
        try:
            tnmaker_data = {}; qm_data = []
            
            for res in sorted_results:
                exam_id = res['exam_id']
                # Rút trích đáp án Trắc nghiệm
                tn_keys = [str(item.get('key', '')).strip().upper() for item in res['items'] if item['type'] == 'QUESTION' and item.get('q_type') == 'TN']
                tn_str = "".join(tn_keys)
                
                # Môn AV/DGNL không có Đúng/Sai và Trả lời ngắn, nên đuôi để trống "##"
                tnmaker_data[str(exam_id)] = f"{tn_str}##"
                qm_data.append([str(exam_id)] + list(tn_str))

            tnmaker_data["success"] = True; tnmaker_data["type"] = 5
            tnmaker_json = json.dumps(tnmaker_data, ensure_ascii=False, separators=(',', ':'))
            qm_json = json.dumps(qm_data, ensure_ascii=False, separators=(',', ':'))

            if qm_data:
                p_qr_title = doc_base.add_paragraph()
                p_qr_title.paragraph_format.space_before = Pt(18); p_qr_title.paragraph_format.space_after = Pt(6)
                format_run(p_qr_title.add_run("MÃ QR NHẬP ĐÁP ÁN TRẮC NGHIỆM"), bold=True, size=12, color=COLOR_BLUE)
                
                tbl_qr = doc_base.add_table(rows=1, cols=2)
                tbl_qr.autofit = False
                tbl_qr.columns[0].width = Cm(9.48); tbl_qr.columns[1].width = Cm(9.48)
                
                # Cấu hình bảng ẩn (xóa viền)
                for row in tbl_qr.rows:
                    for cell in row.cells:
                        tcPr = cell._element.get_or_add_tcPr(); tcB = OxmlElement('w:tcBorders')
                        for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                            e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'nil'); tcB.append(e)
                        tcPr.append(tcB)

                def add_qr_to_cell(cell, data_str, title):
                    qr = qrcode.QRCode(version=1, box_size=4, border=2); qr.add_data(data_str); qr.make(fit=True)
                    img = qr.make_image(fill_color="black", back_color="white"); img_io = io.BytesIO(); img.save(img_io, format='PNG'); img_io.seek(0)
                    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    format_run(p.add_run(f"App: {title}\n"), bold=True, size=12, color=COLOR_PURPLE)
                    p.add_run().add_picture(img_io, width=Cm(4.5))

                add_qr_to_cell(tbl_qr.cell(0,0), tnmaker_json, "TNMaker")
                add_qr_to_cell(tbl_qr.cell(0,1), qm_json, "QM")
        except Exception as e:
            doc_base.add_paragraph().add_run(f"[LỖI TẠO QR CODE: {str(e)}]").font.color.rgb = RGBColor(255, 0, 0)

# ============================================================================
# HÀM CHÍNH (ĐIỂM VÀO TỪ APP.PY)
# ============================================================================
def generate_loigiai_file_av(doc, items, config, exam_id):
    insert_header_and_student_info(doc, config, exam_id)
    insert_quy_uoc(doc, config)
    
    # --- CÔNG TẮC SONG NGỮ ---
    is_eng = config.get('loai_mon') in ['ENG', 'AV'] 
    lbl_key = "Key" if is_eng else "Đáp án"
    lbl_sol = "Solution" if is_eng else "Lời giải"
    lbl_end = "---------------THE END-----------------" if is_eng else "---------------HẾT-----------------"
    
    for item in items:
        if item['type'] == 'TEXT':
            new_element = deepcopy(item['obj'])
            safe_insert_element(doc, new_element)
            if get_tag(new_element) == 'p':
                p = Paragraph(new_element, doc)
                force_compact_p(p)
                for run in p.runs: run.font.name = FONT_NAME; run.font.size = Pt(FONT_SIZE)
                
        elif item['type'] == 'CONTEXT':
            render_cell_content(doc, item['cell'])

        elif item['type'] == 'QUESTION':
            q_label = f"Question {item['print_id']}. " if is_eng else f"Câu {item['print_id']}. "
            render_cell_content(doc, item['stem'], first_line_prefix=q_label, first_line_color=COLOR_BLUE)
            render_options_with_tabs(doc, item['options'])
            
            p_k = doc.add_paragraph(); p_k.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.element.body.remove(p_k._element); doc.element.body.insert(-1, p_k._element)
            
            raw_key = str(item.get('key', ''))
            display_key = raw_key.lower() if item.get('q_type') == 'TL' else raw_key
            
            format_run(p_k.add_run(f"{lbl_key}: {display_key}"), bold=True, color=COLOR_BLUE)
                
            # [CẬP NHẬT MỚI] - Chỉ in Lời giải nếu KHÔNG PHẢI môn Tiếng Anh
            if item.get('solution') and not is_eng: 
                p_l = doc.add_paragraph(); p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.element.body.remove(p_l._element); doc.element.body.insert(-1, p_l._element)
                
                format_run(p_l.add_run(lbl_sol), bold=True, color=COLOR_BLUE)
                render_cell_content(doc, item['solution'])
            
    p_end = doc.add_paragraph(lbl_end)
    doc.element.body.remove(p_end._element); doc.element.body.insert(-1, p_end._element)
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER; format_run(p_end.runs[0], bold=True)
    p_end.paragraph_format.space_after = Pt(6)
    insert_signature_footer(doc, config); insert_common_footer(doc, exam_id)

import pandas as pd
from openpyxl.styles import Font

def generate_excel_dapan_av(all_results):
    # AV/DGNL: Chống chỉ định các mode khác, ÉP BUỘC dùng Dọc nối tiếp (Mode 1)
    data = []
    sorted_results = sorted(all_results, key=lambda x: str(x['exam_id']))
    for res in sorted_results:
        exam_id = res['exam_id']
        for item in res['items']:
            if item['type'] == 'QUESTION':
                raw_key = str(item.get('key', ''))
                display_key = raw_key.lower() if item.get('q_type') == 'TL' else raw_key
                data.append({
                    'Mã đề': exam_id,
                    'Câu': item.get('print_id', ''),
                    'Đáp án': display_key
                })
    df = pd.DataFrame(data)
    
    # Xuất file & Ép Font Times New Roman 12
    excel_buffer = io.BytesIO()
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False)
        worksheet = writer.sheets['Sheet1']
        font = Font(name='Times New Roman', size=12)
        for row in worksheet.iter_rows():
            for cell in row:
                cell.font = font
                
    excel_buffer.seek(0)
    return excel_buffer
    
def xuat_ket_qua(data_mixed, config, output_folder):
    zip_buffer = io.BytesIO()
    all_results = []
    
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, d in enumerate(data_mixed):
            exam_id = d.get('exam_id', '001')
            tmp = io.BytesIO(); d.get('file_content').save(tmp); raw_bytes = tmp.getvalue()
            
            doc_scan = Document(io.BytesIO(raw_bytes))
            items = extract_data_from_raw_av(doc_scan)
            
            result_entry = {'exam_id': exam_id, 'items': items}
            if i == 0: result_entry['raw_bytes'] = raw_bytes
            all_results.append(result_entry)
            
            # --- TẠO FILE MÃ ĐỀ ---
            doc_made = Document(io.BytesIO(raw_bytes))
            style_made = doc_made.styles['Normal']
            style_made.font.name = FONT_NAME; style_made.font.size = Pt(FONT_SIZE)
            
            clear_body_preserve_section(doc_made); set_narrow_layout(doc_made)
            generate_made_file_av(doc_made, items, config, exam_id)
            out_made = io.BytesIO(); doc_made.save(out_made)
            zf.writestr(f"made_{exam_id}.docx", out_made.getvalue())
            
            # --- TẠO FILE LỜI GIẢI ---
            doc_sol = Document(io.BytesIO(raw_bytes))
            style_sol = doc_sol.styles['Normal']
            style_sol.font.name = FONT_NAME; style_sol.font.size = Pt(FONT_SIZE)
            
            clear_body_preserve_section(doc_sol); set_narrow_layout(doc_sol)
            generate_loigiai_file_av(doc_sol, items, config, exam_id)
            out_sol = io.BytesIO(); doc_sol.save(out_sol)
            zf.writestr(f"loigiai_{exam_id}.docx", out_sol.getvalue())

        all_results.sort(key=lambda x: str(x['exam_id']))
        
        # --- TẠO FILE ĐÁP ÁN TỔNG HỢP ---
        if all_results and 'raw_bytes' in all_results[0]:
            doc_sum = Document(io.BytesIO(all_results[0]['raw_bytes']))
        else: doc_sum = Document()
            
        set_narrow_layout(doc_sum)
        generate_dapan_tonghop_av(doc_sum, all_results, config)
        
        out_sum = io.BytesIO(); doc_sum.save(out_sum)
        zf.writestr("dapan_tonghop.docx", out_sum.getvalue())
        
        # [TÍNH NĂNG MỚI] - XUẤT EXCEL CHO AV/DGNL (Chỉ Mode 1)
        try:
            excel_buf = generate_excel_dapan_av(all_results)
            zf.writestr("dapan_tonghop.xlsx", excel_buf.getvalue())
        except Exception as e:
            import streamlit as st
            st.error(f"⚠️ Lỗi xuất Excel: {e}. Vui lòng kiểm tra xem đã chạy lệnh 'pip install openpyxl' chưa!")    
            
    zip_buffer.seek(0)
    return zip_buffer