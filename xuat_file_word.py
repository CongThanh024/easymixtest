import io
import zipfile
import json
import qrcode
import re
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table

# ============================================================================
# KHU VỰC 0: CẤU HÌNH & HÀM DÙNG CHUNG (KHÓA - CẤM THAY ĐỔI)
# ============================================================================
FONT_NAME = "Times New Roman"
FONT_SIZE = 12
COEFF_TEXT = 0.55
PIXELS_TO_CM = 360000
PAGE_WIDTH_CM = 18.96 
WIDTH_COL_MADE = 2.0
THRESH_MODE_4 = 4.2  
THRESH_MODE_2 = 8.9  
SAFETY_FACTOR = 1.0  
COLOR_BLUE = RGBColor(0, 0, 255)
COLOR_PURPLE = RGBColor(112, 48, 160)
COLOR_PURPLE_TIT = RGBColor(93, 99, 211)
HEX_BG_PART1 = "7030A0"
HEX_BG_PART2 = "385723"
HEX_BG_PART3 = "C65911"
NSMAP = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main', 
         'm': 'http://schemas.openxmlformats.org/officeDocument/2006/math', 
         'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing', 
         'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
         'v': 'urn:schemas-microsoft-com:vml'}

def get_tag(element): return element.tag.split('}')[-1]
def measure_text_cm(text): return len(text) * COEFF_TEXT * FONT_SIZE / 28.35 if text else 0
def count_math_chars_recursive(element, in_fraction=False):
    count = 0.0
    tag = get_tag(element)
    current_in_fraction = in_fraction or (tag == 'f')
    if tag == 't':
        text = element.text if element.text else ""
        length = len(text)
        if current_in_fraction: count += length / 2.0
        else: count += length
    for child in element: count += count_math_chars_recursive(child, current_in_fraction)
    return count
def measure_math_cm(math_element): return count_math_chars_recursive(math_element) * COEFF_TEXT * FONT_SIZE / 28.35
def measure_ole_object_cm(object_element):
    shapes = object_element.findall('.//v:shape', NSMAP)
    total_w = 0.0
    for shape in shapes:
        style = shape.get('style', '')
        if 'width:' in style:
            try:
                width_part = style.split('width:')[1].split(';')[0].strip()
                if 'pt' in width_part: val = float(width_part.replace('pt', '')); total_w += val / 28.35
                elif 'in' in width_part: val = float(width_part.replace('in', '')); total_w += val * 2.54
                elif 'cm' in width_part: val = float(width_part.replace('cm', '')); total_w += val
            except: pass
    return total_w
def measure_image_cm(drawing_element):
    extents = drawing_element.findall('.//wp:extent', NSMAP)
    max_w = 0
    for ext in extents:
        try:
            cx = int(ext.get('cx', 0)); w = cx / PIXELS_TO_CM
            if w > max_w: max_w = w
        except: pass
    return max_w
def calculate_cell_width(cell):
    total = 0.0
    for p in cell.paragraphs:
        for child in p._element:
            tag = get_tag(child)
            if tag == 'r':
                for t in child.findall('.//w:t', NSMAP): total += measure_text_cm(t.text)
                for d in child.findall('.//w:drawing', NSMAP): total += measure_image_cm(d)
                for obj in child.findall('.//w:object', NSMAP): total += measure_ole_object_cm(obj)
            elif tag in ['oMath', 'oMathPara']: total += measure_math_cm(child)
    return total
def calculate_option_width(cell): return calculate_cell_width(cell)
def format_run(run, bold=False, italic=False, color=None, size=None):
    run.font.name = FONT_NAME
    run.font.size = Pt(size if size else FONT_SIZE)
    run.font.bold = bold; run.font.italic = italic
    if color: run.font.color.rgb = color
def set_cell_shading(cell, hex_str):
    if not hex_str: return
    tcPr = cell._element.get_or_add_tcPr()
    old_shd = tcPr.find(qn('w:shd'))
    if old_shd is not None: tcPr.remove(old_shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), str(hex_str)) 
    
    # [BẢO MẬT XML TẦNG 1] Dò tìm thứ tự để chèn Màu Nền (shd) đúng chuẩn
    insert_idx = len(tcPr)
    for i, child in enumerate(tcPr):
        if child.tag in [qn('w:noWrap'), qn('w:tcMar'), qn('w:textDirection'), qn('w:tcFitText'), qn('w:vAlign'), qn('w:hideMark')]:
            insert_idx = i
            break
    tcPr.insert(insert_idx, shd)

def add_table_borders(table):
    tbl = table._tbl; tblPr = tbl.tblPr
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None: tblPr.remove(old_borders)
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4'); border.set(qn('w:space'), '0'); border.set(qn('w:color'), 'auto')
        borders.append(border)
        
    # [BẢO MẬT XML TẦNG 2] Dò tìm thứ tự để chèn Viền Bảng (tblBorders) đúng chuẩn
    insert_idx = len(tblPr)
    for i, child in enumerate(tblPr):
        if child.tag in [qn('w:shd'), qn('w:tblLayout'), qn('w:tblCellMar'), qn('w:tblLook'), qn('w:tblCaption')]:
            insert_idx = i
            break
    tblPr.insert(insert_idx, borders)

def add_field(p, code, italic=False):
    r = p.add_run(); f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin'); r._r.append(f1)
    if italic: r.font.italic = True
    r = p.add_run(); i = OxmlElement('w:instrText'); i.text = code; r._r.append(i)
    if italic: r.font.italic = True
    r = p.add_run(); f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end'); r._r.append(f2)
    if italic: r.font.italic = True
def set_narrow_layout(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(1.02)
        s.page_width, s.page_height = Cm(21), Cm(29.7)
def safe_insert_element(doc, element): doc.element.body.insert(-1, deepcopy(element))
def copy_paragraph_content(source_p, target_p):
    if not source_p: return
    # 1. Bốc nguyên vẹn toàn bộ dữ liệu, công thức, hình ảnh sang
    for child in source_p._element:
        if get_tag(child) in ['r', 'oMath', 'oMathPara', 'hyperlink', 'drawing', 'object']: 
            target_p._element.append(deepcopy(child))
            
    # 2. [AN TOÀN] Chỉ đổi tên Font và Size, không đụng vào in đậm/in nghiêng/chỉ số/công thức
    for run in target_p.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE)

def is_p_empty(p):
    """Máy hút bụi an toàn: Nhận diện chính xác dòng rỗng."""
    if p.text.strip(): return False
    # Soi thật kỹ các thẻ XML, chỉ cần có 1 cấu trúc hình ảnh/toán học là không được xóa
    for child in p._element.iter():
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ['oMath', 'oMathPara', 'drawing', 'object', 'pict', 'AlternateContent', 'shape']:
            return False 
    return True

def render_cell_content(doc, cell, first_line_prefix=None, first_line_color=None):
    if not cell: return
    
    elements = []
    for child in cell._element.iterchildren():
        tag = get_tag(child)
        if tag in ['p', 'tbl']: elements.append(child)
        
    last_valid_idx = -1
    for i in range(len(elements) - 1, -1, -1):
        if get_tag(elements[i]) == 'tbl':
            last_valid_idx = i; break
        elif get_tag(elements[i]) == 'p':
            p_obj = Paragraph(elements[i], cell)
            if not is_p_empty(p_obj):
                last_valid_idx = i; break
                
    valid_elements = elements[:last_valid_idx + 1]

    if not valid_elements:
        if first_line_prefix:
            new_p = doc.add_paragraph(); force_compact_p(new_p)
            doc.element.body.remove(new_p._element); doc.element.body.insert(-1, new_p._element)
            new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
            r = new_p.add_run(first_line_prefix); format_run(r, bold=True, color=first_line_color)
        return

    for i, elem in enumerate(valid_elements):
        tag = get_tag(elem)
        if tag == 'p':
            p = Paragraph(elem, cell)
            new_p = doc.add_paragraph(); force_compact_p(new_p)
            doc.element.body.remove(new_p._element); doc.element.body.insert(-1, new_p._element)
            if i == 0 and first_line_prefix:
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY 
                r = new_p.add_run(first_line_prefix); format_run(r, bold=True, color=first_line_color)
            elif p.alignment: new_p.alignment = p.alignment
            copy_paragraph_content(p, new_p)
        elif tag == 'tbl':
            new_tbl = deepcopy(elem)
            doc.element.body.insert(-1, new_tbl)
            if i == 0 and first_line_prefix:
                new_p = doc.add_paragraph(); force_compact_p(new_p)
                new_p._element.addprevious(new_p._element)
                r = new_p.add_run(first_line_prefix); format_run(r, bold=True, color=first_line_color)
def clear_body_preserve_section(doc):
    body = doc.element.body
    for child in list(body):
        if not child.tag.endswith('sectPr'): body.remove(child)

# --- Header & Footer (KHÓA - CẤM THAY ĐỔI) ---
def force_compact_p(p_obj):
    """Hàm hỗ trợ: Ép khít tuyệt đối khoảng cách đoạn văn do Python sinh ra."""
    p_obj.paragraph_format.space_before = Pt(0)
    p_obj.paragraph_format.space_after = Pt(0)
    p_obj.paragraph_format.line_spacing = 1.0

def insert_header_and_student_info(doc, config, exam_id):
    if config.get('co_header', True):
        hd = config.get('header_data', {})
        tbl = doc.add_table(rows=2, cols=2); tbl.autofit = False
        tbl.columns[0].width, tbl.columns[1].width = Cm(8.0), Cm(10.96)
        
        c = tbl.cell(0,0); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; force_compact_p(p)
        format_run(p.add_run(f"{hd.get('so','').upper()}\n"), size=12, color=COLOR_BLUE)
        format_run(p.add_run(f"{hd.get('truong','').upper()}\n"), bold=True, size=12, color=COLOR_BLUE)
        format_run(p.add_run(hd.get('to_chuyen_mon','').upper()), bold=True, size=12, color=COLOR_BLUE)
        
        c = tbl.cell(0,1); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; force_compact_p(p)
        format_run(p.add_run(f"{hd.get('ky_thi','').upper()}\n"), bold=True, size=12, color=COLOR_BLUE)
        format_run(p.add_run(f"{hd.get('nam_hoc','')}\n"), size=12, color=COLOR_BLUE)
        format_run(p.add_run(f"Môn: {hd.get('mon','')}"), bold=True, color=COLOR_PURPLE_TIT)
        
        c = tbl.cell(1,0); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; force_compact_p(p)
        format_run(p.add_run("ĐỀ CHÍNH THỨC\n"), bold=True, color=COLOR_PURPLE_TIT)
        format_run(p.add_run("(Đề thi có "), italic=True); add_field(p, "NUMPAGES", italic=True); format_run(p.add_run(" trang)"), italic=True)
        
        c = tbl.cell(1,1); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; force_compact_p(p)
        t_val = "".join(re.findall(r'\d+', str(hd.get('thoi_gian','')))) or "..."
        p.add_run(f"Thời gian làm bài: {t_val} phút\n"); format_run(p.add_run("(Không kể thời gian phát đề)"), italic=True, size=12)
        
        for r in tbl.rows:
            for cl in r.cells: 
                tcPr = cl._element.get_or_add_tcPr(); tcB = OxmlElement('w:tcBorders')
                for b in ['top', 'left', 'bottom', 'right']: e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'nil'); tcB.append(e)
                tcPr.append(tcB)
                
        # [SỬA LỖI 1]: Thụt lề dòng này chui vào trong lệnh IF để biến mất khi bạn tắt Tiêu đề
        p_blank = doc.add_paragraph(); force_compact_p(p_blank)
    
    tbl_i = doc.add_table(rows=1, cols=2); tbl_i.autofit = False; tbl_i.allow_autofit = False
    tbl_i.columns[0].width = Cm(15.46); tbl_i.columns[1].width = Cm(3.5)
    tbl_i.cell(0,0).width = Cm(15.46); tbl_i.cell(0,1).width = Cm(3.5)
    
    p = tbl_i.cell(0,0).paragraphs[0]; force_compact_p(p)
    format_run(p.add_run("Họ và tên: ...................................................................... SBD: .........................................."), bold=True, color=COLOR_BLUE)
    
    c = tbl_i.cell(0,1); p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; force_compact_p(p)
    format_run(p.add_run(f"Mã đề thi: {exam_id}"), bold=True, color=COLOR_PURPLE)
    
    tcPr0 = tbl_i.cell(0,0)._element.get_or_add_tcPr(); tcB0 = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']: e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'nil'); tcB0.append(e)
    tcPr0.append(tcB0)
    tcPr1 = c._element.get_or_add_tcPr(); tcB1 = OxmlElement('w:tcBorders')
    for b in ['top', 'left', 'bottom', 'right']: e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'double'); e.set(qn('w:sz'), '6'); tcB1.append(e)
    tcPr1.append(tcB1)
    
def insert_common_footer(doc, exam_id):
    sect = doc.sections[0]; sect.footer.is_linked_to_previous = False
    for p in sect.footer.paragraphs: p._element.getparent().remove(p._element)
    p = sect.footer.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    force_compact_p(p) # Ép khít Chân trang
    r = p.add_run("Trang "); format_run(r, bold=True, color=COLOR_PURPLE)
    add_field(p, "PAGE"); r = p.add_run("/"); format_run(r, bold=True, color=COLOR_PURPLE)
    add_field(p, "NUMPAGES"); r = p.add_run(f" - Mã đề {exam_id}"); format_run(r, bold=True, color=COLOR_PURPLE)

def insert_signature_footer(doc, config):
    if not config.get('co_footer', True): return
    def add_sig_line(title, is_first=False):
        p = doc.add_paragraph(); doc.element.body.remove(p._element); doc.element.body.insert(-1, p._element)
        force_compact_p(p) # Ép khít dòng Chữ ký
        if is_first: p.paragraph_format.space_before = Pt(6) 
        else: p.paragraph_format.space_before = Pt(6)
        tabs = p.paragraph_format.tab_stops
        tabs.add_tab_stop(Cm(9.48), WD_TAB_ALIGNMENT.LEFT, WD_TAB_LEADER.DOTS)
        tabs.add_tab_stop(Cm(18.96), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        r = p.add_run(f"{title}\t Ký tên:\t"); format_run(r, size=12)
    add_sig_line("Giám thị 1:", is_first=True); add_sig_line("Giám thị 2:", is_first=False)

# ============================================================================
# KHU VỰC 1: FILE MADE & LOIGIAI (KHÓA - CẤM THAY ĐỔI)
# ============================================================================
from docx.text.paragraph import Paragraph

def insert_images(doc, config):
    img_tn = config.get('img_phieu_to')
    img_tl = config.get('img_tu_luan')
    co_header = config.get('co_header', True)
    
    def add_pic(img, is_page_1=True):
        if not img: return
        try:
            from PIL import Image as PILImage
            img.seek(0)
            with PILImage.open(img) as pi:
                w_orig, h_orig = pi.size
            
            max_w = 18.96
            
            # Tính toán không gian bề dọc: Trừ hao an toàn thêm 0.2cm để chống rớt trang lặt vặt
            if is_page_1:
                max_h = 22.4 if co_header else 25.8
            else:
                max_h = 27.4 # Phiếu tự luận chiếm toàn bộ trang 2
                
            ratio_w = max_w / w_orig
            ratio_h = max_h / h_orig
            ratio = min(ratio_w, ratio_h) 
            
            final_w = w_orig * ratio
            final_h = h_orig * ratio
            
            img.seek(0); p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # XỬ LÝ NGẮT TRANG VÀ LỀ TRÊN
            from docx.enum.text import WD_BREAK
            r = p.add_run()
            if not is_page_1:
                r.add_break(WD_BREAK.PAGE) # Ép nhảy sang trang 2
                p.paragraph_format.space_before = Pt(0)
            else:
                p.paragraph_format.space_before = Pt(12) # Cách bảng Info 12pt cho trang 1
                
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r.add_picture(img, width=Cm(final_w), height=Cm(final_h))
            
        except Exception as e: 
            try:
                img.seek(0); p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                from docx.enum.text import WD_BREAK
                r = p.add_run()
                if not is_page_1: r.add_break(WD_BREAK.PAGE)
                p.paragraph_format.space_before = Pt(6) if is_page_1 else Pt(0)
                r.add_picture(img, width=Cm(18.96))
            except: pass

    # Chỉ gọi hàm khi có ảnh tải lên
    if img_tn: add_pic(img_tn, is_page_1=True)
    if img_tl: add_pic(img_tl, is_page_1=False)

# ============================================================================
# HÀM BỐC QUY ƯỚC MÔN (BẢN VÁ LỖI TÀNG HÌNH & HỆ THỐNG BÁO LỖI TRỰC TIẾP)
# ============================================================================
def insert_quy_uoc(doc, config):
    file_qu = config.get('file_quy_uoc')
    if not file_qu: return
    try:
        from docx import Document
        from docx.table import Table
        import io
        
        file_qu.seek(0)
        doc_qu = Document(io.BytesIO(file_qu.read()))
        
        for element in doc_qu.element.body.iterchildren():
            tag = get_tag(element)
            if tag in ['p', 'tbl']:
                # 1. Bốc nguyên xi cấu trúc XML
                new_element = deepcopy(element)
                
                # 2. Gọt bỏ mầm mống gây hỏng file (Hình ảnh, MathType cũ)
                bad_tags = ['drawing', 'object', 'hyperlink', 'pict']
                for bad_tag in bad_tags:
                    for bad_elem in new_element.findall(f'.//w:{bad_tag}', NSMAP):
                        parent = bad_elem.getparent()
                        if parent is not None:
                            parent.remove(bad_elem)
                            
                # 3. CHÈN TRỰC TIẾP VÀO BODY - Chìa khóa giúp hiện chữ (Không qua copy trung gian)
                doc.element.body.insert(-1, new_element)
                
                # 4. Ép format
                try:
                    if tag == 'p':
                        p = Paragraph(new_element, doc)
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        p.paragraph_format.line_spacing = 1.0
                        for run in p.runs:
                            run.font.name = FONT_NAME
                            run.font.size = Pt(FONT_SIZE)
                    elif tag == 'tbl':
                        tbl = Table(new_element, doc)
                        for row in tbl.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    p.paragraph_format.space_before = Pt(0)
                                    p.paragraph_format.space_after = Pt(0)
                                    p.paragraph_format.line_spacing = 1.0
                                    for run in p.runs:
                                        run.font.name = FONT_NAME
                                        run.font.size = Pt(FONT_SIZE)
                except:
                    pass # Bỏ qua lỗi ép font (nếu có), miễn là giữ được chữ
    except Exception as e:
        # Nếu có lỗi phần mềm, in thẳng cảnh báo màu đỏ lên Đề thi để nhận diện
        p_err = doc.add_paragraph()
        r_err = p_err.add_run(f"[HỆ THỐNG BÁO LỖI CHÈN QUY ƯỚC: {str(e)}]")
        r_err.font.color.rgb = RGBColor(255, 0, 0)

def generate_made_file(doc, items, config, exam_id):
    insert_header_and_student_info(doc, config, exam_id); insert_images(doc, config)
    insert_quy_uoc(doc, config)
    for item in items:
        if item['type'] == 'PART':
            p_part = deepcopy(item['obj']); pPr = p_part.get_or_add_pPr()
            spacing = OxmlElement('w:spacing'); spacing.set(qn('w:before'), '120'); spacing.set(qn('w:after'), '120')
            old = pPr.find(qn('w:spacing'))
            if old is not None: pPr.remove(old)
            pPr.append(spacing); safe_insert_element(doc, p_part)
            
        elif item['type'] == 'TEXT':
            new_element = deepcopy(item['obj'])
            safe_insert_element(doc, new_element)
            if get_tag(new_element) == 'p':
                p = Paragraph(new_element, doc)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(FONT_SIZE)
            elif get_tag(new_element) == 'tbl':
                tbl = Table(new_element, doc)
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                run.font.name = FONT_NAME
                                run.font.size = Pt(FONT_SIZE)
                                
        elif item['type'] == 'TEXT_CELL':
            render_cell_content(doc, item['obj'])
            
        elif item['type'] == 'QUESTION':
            q_num = item.get('num', item['id']) 
            render_cell_content(doc, item['stem'], first_line_prefix=f"Câu {q_num}. ", first_line_color=COLOR_BLUE)
            render_options_with_tabs(doc, item['options'])
            
    p_end = doc.add_paragraph("---------------HẾT-----------------")
    doc.element.body.remove(p_end._element); doc.element.body.insert(-1, p_end._element)
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER; format_run(p_end.runs[0], bold=True)
    p_end.paragraph_format.space_after = Pt(6)
    insert_signature_footer(doc, config); insert_common_footer(doc, exam_id)

def generate_loigiai_file(doc, items, config, exam_id):
    insert_header_and_student_info(doc, config, exam_id)
    insert_quy_uoc(doc, config)
    for item in items:
        if item['type'] == 'PART':
            p_part = deepcopy(item['obj']); pPr = p_part.get_or_add_pPr()
            spacing = OxmlElement('w:spacing'); spacing.set(qn('w:before'), '120'); spacing.set(qn('w:after'), '120')
            old = pPr.find(qn('w:spacing'))
            if old is not None: pPr.remove(old)
            pPr.append(spacing); safe_insert_element(doc, p_part)
            
        elif item['type'] == 'TEXT':
            new_element = deepcopy(item['obj'])
            safe_insert_element(doc, new_element)
            if get_tag(new_element) == 'p':
                p = Paragraph(new_element, doc)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(FONT_SIZE)
            elif get_tag(new_element) == 'tbl':
                tbl = Table(new_element, doc)
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(0)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                run.font.name = FONT_NAME
                                run.font.size = Pt(FONT_SIZE)
                                
        elif item['type'] == 'TEXT_CELL':
            render_cell_content(doc, item['obj'])

        elif item['type'] == 'QUESTION':
            q_num = item.get('num', item['id'])
            render_cell_content(doc, item['stem'], first_line_prefix=f"Câu {q_num}. ", first_line_color=COLOR_BLUE)
            render_options_with_tabs(doc, item['options'])
            p_l = doc.add_paragraph(); p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.element.body.remove(p_l._element); doc.element.body.insert(-1, p_l._element)
            format_run(p_l.add_run("Lời giải"), bold=True, color=COLOR_BLUE)
            if item['q_type'] != 'TL':
                p_k = doc.add_paragraph(); doc.element.body.remove(p_k._element); doc.element.body.insert(-1, p_k._element)
                raw_k = item['key']
                if item['q_type'] == 'DS': raw_k = raw_k.upper().replace('T','Đ').replace('F','S')
                format_run(p_k.add_run(f"Đáp án: {raw_k}"), bold=True, color=COLOR_BLUE)
            if item['solution']: render_cell_content(doc, item['solution'])
            
    p_end = doc.add_paragraph("---------------HẾT-----------------")
    doc.element.body.remove(p_end._element); doc.element.body.insert(-1, p_end._element)
    p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER; format_run(p_end.runs[0], bold=True)
    p_end.paragraph_format.space_after = Pt(6)
    insert_signature_footer(doc, config); insert_common_footer(doc, exam_id)

# ============================================================================
# KHU VỰC 3: FILE DAPAN_TONGHOP (ĐÃ SỬA LỖI DÍNH BẢNG)
# ============================================================================
def generate_dapan_tonghop(doc_base, all_results, config):
    if not all_results: return
    
    # =========================================================================
    # ĐOẠN 1: SETUP VÀ BẢO TỒN TỰ LUẬN (THUẬT TOÁN GIÁM ĐỐC)
    # =========================================================================
    def local_strict_detect_temp(key_val):
        k_up = str(key_val).strip().upper()
        if k_up == 'TL' or 'TỰ LUẬN' in k_up: return 'TL'
        clean_tn = k_up.replace(".", "").replace(")", "").strip()
        if clean_tn in ['A', 'B', 'C', 'D']: return 'TN'
        clean_ds = re.sub(r'[^TFĐS]', '', k_up)
        if len(clean_ds) == 4: return 'DS'
        return 'TLN'

    body = doc_base.element.body
    all_blocks = list(body)
    
    # Đếm chính xác số lượng câu Tự luận
    num_tl = 0
    for item in all_results[0]['items']:
        if item['type'] == 'QUESTION' and local_strict_detect_temp(item.get('key', '')) == 'TL':
            num_tl += 1
            
    question_tables = []
    for block in all_blocks:
        if block.tag.endswith('tbl'):
            tbl = Table(block, doc_base)
            try:
                if "CÂU" in tbl.rows[0].cells[0].text.strip().upper() and len(tbl.rows[0].cells) >= 2:
                    question_tables.append(block)
            except: pass

    tl_blocks = []
    # Cất giữ nguyên khối các Bảng Tự luận (Giữ nguyên Căn cước ID gốc)
    if num_tl > 0 and len(question_tables) >= num_tl:
        first_tl_tbl_block = question_tables[-num_tl]
        idx_start_tl = all_blocks.index(first_tl_tbl_block)
        for b in all_blocks[idx_start_tl:]:
            if not b.tag.endswith('sectPr'):
                tl_blocks.append(b)
                body.remove(b)

    # Dọn dẹp mặt bằng để vẽ Bảng QR, TN, DS...
    clear_body_preserve_section(doc_base)

    style = doc_base.styles['Normal']; font = style.font
    font.name = FONT_NAME; font.size = Pt(FONT_SIZE)

    p_title1 = doc_base.add_paragraph(); p_title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title1.paragraph_format.space_after = Pt(6)
    format_run(p_title1.add_run("BẢNG ĐÁP ÁN TỔNG HỢP"), bold=True, size=14, color=COLOR_BLUE)

    hd = config.get('header_data', {})
    ky_thi = hd.get('ky_thi', '...').upper()
    nam_hoc = hd.get('nam_hoc', '...')
    p_title2 = doc_base.add_paragraph(); p_title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title2.paragraph_format.space_after = Pt(12)
    format_run(p_title2.add_run(f"{ky_thi}, {nam_hoc}"), bold=True, size=13)

    # 2. HELPER FUNCTIONS
    def get_score_val(k):
        try: return float(config.get(k, 0))
        except: return 0.0
    def fmt_float(val): return str(round(val, 3)).replace('.', ',')
    def apply_compact_cell(cell, align_center=True):
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if align_center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def set_col_widths(tbl, col_widths):
        tbl.autofit = False; tbl.allow_autofit = False
        for row in tbl.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells): row.cells[idx].width = width

    # --- LOGIC STRICT: CLEAN & DETECT ---
    def local_clean_key(raw_text):
        text = re.sub(r"^(KEY|ĐÁP ÁN|LỜI GIẢI|HƯỚNG DẪN GIẢI)[:\s.]*", "", raw_text, flags=re.IGNORECASE).strip()
        text = text.strip('"').strip()
        return text

    def local_strict_detect(key_val):
        k_up = str(key_val).strip().upper()
        if k_up == 'TL' or 'TỰ LUẬN' in k_up: return 'TL'
        clean_tn = k_up.replace(".", "").replace(")", "").strip()
        if clean_tn in ['A', 'B', 'C', 'D']: return 'TN'
        clean_ds = re.sub(r'[^TFĐS]', '', k_up)
        if len(clean_ds) == 4: return 'DS'
        return 'TLN'

    # 3. GLOBAL CENSUS
    mapped_data = {}
    for res in all_results:
        exam_id = res['exam_id']
        buckets = {'TN': [], 'DS': [], 'TLN': [], 'TL': []}
        for item in res['items']:
            if item['type'] == 'QUESTION':
                raw_k = item['key']
                real_type = local_strict_detect(raw_k)
                buckets[real_type].append(item)
        mapped_data[exam_id] = buckets

    first_exam_id = all_results[0]['exam_id']
    ref_buckets = mapped_data[first_exam_id]

    # 4. VẼ BẢNG THEO THỨ TỰ
    draw_order = [('TN', 'TRẮC NGHIỆM', 'diem_p1'), 
                  ('DS', 'TRẮC NGHIỆM ĐÚNG SAI', 'diem_p2'), 
                  ('TLN', 'TRẮC NGHIỆM TRẢ LỜI NGẮN', 'diem_p3'),
                  ('TL', 'TỰ LUẬN', 'diem_p4')]
    
    part_roman_map = {1: 'I', 2: 'II', 3: 'III', 4: 'IV'}
    current_part_idx = 0 

    for q_type, base_title, score_key in draw_order:
        ref_items = ref_buckets[q_type]
        count = len(ref_items)
        # --- BẮT ĐẦU CHÈN BẢNG MÃ QR (CHỈ VẼ 1 LẦN TRƯỚC PHẦN TỰ LUẬN) ---
        if q_type == 'TL':
            try:
                tnmaker_data = {}; qm_data = []
                sorted_exam_ids = sorted(mapped_data.keys(), key=lambda x: str(x))
                for exam_id in sorted_exam_ids:
                    my_bucket = mapped_data[exam_id]
                    tn_str = "".join([str(it.get('key','')).strip().upper() for it in my_bucket['TN']])
                    
                    ds_chunks_tnmaker = []; ds_qm_chars = []
                    for it in my_bucket['DS']:
                        clean_ds = re.sub(r'[^TFtfDdSsĐđ]', '', str(it.get('key','')))
                        fmt_tnmaker = ""; fmt_qm = []
                        for char in clean_ds[:4]:
                            is_true = char.upper() in ['T', 'D', 'Đ']
                            fmt_tnmaker += 'Đ' if is_true else 'S'
                            fmt_qm.append('D' if is_true else 'S')
                        ds_chunks_tnmaker.append(fmt_tnmaker.ljust(4, 'S'))
                        while len(fmt_qm) < 4: fmt_qm.append('S')
                        ds_qm_chars.extend(fmt_qm)

                    tln_chunks_tnmaker = []; tln_qm_vals = []
                    for it in my_bucket['TLN']:
                        raw_tln = str(it.get('key','')).strip()
                        tln_chunks_tnmaker.append(raw_tln.replace('.', ','))
                        clean_qm = raw_tln.replace(',', '.')
                        try: tln_qm_vals.append(float(clean_qm) if '.' in clean_qm else int(clean_qm))
                        except: tln_qm_vals.append(raw_tln)

                    tnmaker_data[str(exam_id)] = f"{tn_str}#{'_'.join(ds_chunks_tnmaker)}#{'_'.join(tln_chunks_tnmaker)}"
                    qm_data.append([str(exam_id)] + list(tn_str) + ds_qm_chars + tln_qm_vals)

                tnmaker_data["success"] = True; tnmaker_data["type"] = 5
                tnmaker_json = json.dumps(tnmaker_data, ensure_ascii=False, separators=(',', ':'))
                qm_json = json.dumps(qm_data, ensure_ascii=False, separators=(',', ':'))

                if qm_data:
                    p_qr_title = doc_base.add_paragraph()
                    p_qr_title.paragraph_format.space_before = Pt(18); p_qr_title.paragraph_format.space_after = Pt(6)
                    format_run(p_qr_title.add_run("MÃ QR NHẬP ĐÁP ÁN TRẮC NGHIỆM"), bold=True, size=12, color=COLOR_BLUE)
                    
                    tbl_qr = doc_base.add_table(rows=1, cols=2)
                    tbl_qr.autofit = False
                    tbl_qr.columns[0].width = Cm(9.48); tbl_qr.columns[1].width = Cm(9.48)
                    
                    # Cấu hình bảng ẩn (xóa viền)
                    # Cấu hình bảng ẩn (xóa viền) - ĐÃ SỬA LỖI CHẾT FILE
                    for row in tbl_qr.rows:
                        for cell in row.cells:
                            tcPr = cell._element.get_or_add_tcPr(); tcB = OxmlElement('w:tcBorders')
                            # CHỈ DÙNG 4 CẠNH CHO Ô (CELL). Bỏ insideH và insideV
                            for b in ['top', 'left', 'bottom', 'right']:
                                e = OxmlElement(f'w:{b}'); e.set(qn('w:val'), 'nil'); tcB.append(e)
                            tcPr.append(tcB)

                    def add_qr_to_cell(cell, data_str, title):
                        qr = qrcode.QRCode(version=1, box_size=4, border=2); qr.add_data(data_str); qr.make(fit=True)
                        img = qr.make_image(fill_color="black", back_color="white"); img_io = io.BytesIO(); img.save(img_io, format='PNG'); img_io.seek(0)
                        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        format_run(p.add_run(f"App: {title}\n"), bold=True, size=12, color=COLOR_PURPLE)
                        p.add_run().add_picture(img_io, width=Cm(4.5))

                    add_qr_to_cell(tbl_qr.cell(0,0), tnmaker_json, "TNMaker")
                    add_qr_to_cell(tbl_qr.cell(0,1), qm_json, "QM")
            except Exception as e:
                doc_base.add_paragraph().add_run(f"[LỖI TẠO QR CODE: {str(e)}]").font.color.rgb = RGBColor(255, 0, 0)
        # --- KẾT THÚC CHÈN BẢNG MÃ QR ---
        if count == 0: continue 

        current_part_idx += 1
        part_roman = part_roman_map.get(current_part_idx, str(current_part_idx))
        full_title = f"PHẦN {part_roman}. {base_title}"

        # TIÊU ĐỀ PHẦN
        p_head = doc_base.add_paragraph()
        p_head.paragraph_format.space_before = Pt(6)
        p_head.paragraph_format.space_after = Pt(0)
        format_run(p_head.add_run(full_title), bold=True, color=COLOR_BLUE)

        total_score = get_score_val(score_key)
        avg_score = total_score / count if count > 0 else 0
        
        # --- [MỚI] DÒNG TỔNG ĐIỂM PHẦN TÁCH BIỆT ---
        p_note_total = doc_base.add_paragraph()
        p_note_total.paragraph_format.left_indent = Cm(0.5)
        p_note_total.paragraph_format.space_before = Pt(0)
        p_note_total.paragraph_format.space_after = Pt(0)
        p_note_total.paragraph_format.line_spacing = 1.0
        
        r_prefix = p_note_total.add_run(f"- Tổng điểm phần {base_title.lower()}: ")
        format_run(r_prefix, italic=True)
        r_score = p_note_total.add_run(fmt_float(total_score))
        format_run(r_score, italic=True, bold=True, color=RGBColor(255, 0, 0)) # Số điểm in đậm, màu đỏ
        r_suffix = p_note_total.add_run(" điểm.")
        format_run(r_suffix, italic=True)
        
        # --- GHI CHÚ BAREM ĐIỂM CHI TIẾT ---
        if q_type != 'TL': # Tự luận chỉ cần hiện tổng điểm ở trên là đủ
            p_note = doc_base.add_paragraph()
            p_note.paragraph_format.left_indent = Cm(0.5)
            p_note.paragraph_format.space_before = Pt(0)
            p_note.paragraph_format.space_after = Pt(6) 
            p_note.paragraph_format.line_spacing = 1.0
            
            def add_score_text(text, is_number=False):
                r = p_note.add_run(text)
                if is_number:
                    format_run(r, italic=True, bold=True, color=RGBColor(255, 0, 0))
                else:
                    format_run(r, italic=True)

            if q_type == 'TN': 
                add_score_text("- Mỗi câu đúng được ")
                add_score_text(fmt_float(avg_score), is_number=True)
                add_score_text(" điểm.")
            elif q_type == 'TLN': 
                add_score_text("- Mỗi câu trả lời đúng được ")
                add_score_text(fmt_float(avg_score), is_number=True)
                add_score_text(" điểm.")
            elif q_type == 'DS':
                s1, s2, s3, s4 = avg_score*0.1, avg_score*0.25, avg_score*0.5, avg_score
                add_score_text("- Điểm tối đa mỗi câu: ")
                add_score_text(fmt_float(avg_score), is_number=True)
                add_score_text(" điểm.\n- Đúng 1 ý: ")  
                add_score_text(fmt_float(s1), is_number=True)
                add_score_text(" điểm; Đúng 2 ý: ")
                add_score_text(fmt_float(s2), is_number=True)
                add_score_text(" điểm; Đúng 3 ý: ")
                add_score_text(fmt_float(s3), is_number=True)
                add_score_text(" điểm; Đúng 4 ý: ")
                add_score_text(fmt_float(s4), is_number=True)
                add_score_text(" điểm.")

        if q_type == 'TL':
            # =========================================================================
            # ĐOẠN 2: TRẢ LẠI TỰ LUẬN VÀ PHÁ BẢNG
            # =========================================================================
            # Trả khối Tự luận về đáy file (Ngay trên định dạng trang sectPr)
            sect_pr = doc_base.element.body.find(qn('w:sectPr'))
            for b in tl_blocks:
                if sect_pr is not None:
                    sect_pr.addprevious(b)
                else:
                    doc_base.element.body.append(b)
                
            # Duyệt qua các khối vừa trả lại để đập vỡ khung bảng
            cau_idx = 1
            for block in tl_blocks:
                if block.tag.endswith('tbl'):
                    tbl = Table(block, doc_base)
                    try:
                        # [SỬA LỖI 1]: Quét TẤT CẢ các hàng trong bảng để không sót Lời Giải
                        for row in tbl.rows:
                            if len(row.cells) == 0: continue
                            text0 = row.cells[0].text.strip().upper()
                            
                            # --- XỬ LÝ Ô ĐỀ BÀI ---
                            if "CÂU" in text0 and len(row.cells) >= 2:
                                c_stem = row.cells[1]
                                
                                # [SỬA LỖI 2]: Nhét chữ "Câu X. " chui vào cùng dòng với nội dung
                                if c_stem.paragraphs:
                                    p0 = c_stem.paragraphs[0]
                                    r_new = p0.add_run(f"Câu {cau_idx}. ")
                                    format_run(r_new, bold=True, color=COLOR_BLUE)
                                    
                                    # Lách luật XML: Tìm vị trí an toàn để nhét chữ lên đầu mà không làm hỏng file
                                    insert_idx = 0
                                    for i, child in enumerate(p0._element):
                                        if child.tag.endswith('pPr'):
                                            insert_idx = i + 1
                                            break
                                    p0._element.insert(insert_idx, r_new._element)
                                else:
                                    p_title = doc_base.add_paragraph()
                                    format_run(p_title.add_run(f"Câu {cau_idx}. "), bold=True, color=COLOR_BLUE)
                                    block.addprevious(p_title._element)
                                    
                                cau_idx += 1
                                
                                # Kéo nội dung ra ngoài
                                for child in list(c_stem._element):
                                    if child.tag.endswith('p') or child.tag.endswith('tbl'):
                                        block.addprevious(child)
                                        
                            # --- XỬ LÝ Ô LỜI GIẢI ---
                            elif "LỜI GIẢI" in text0 or "KEY" in text0 or "ĐÁP ÁN" in text0:
                                p_sol = doc_base.add_paragraph()
                                p_sol.alignment = WD_ALIGN_PARAGRAPH.CENTER # Canh giữa, bỏ in nghiêng
                                format_run(p_sol.add_run("Lời giải"), bold=True, color=COLOR_PURPLE)
                                block.addprevious(p_sol._element)
                                
                                c_sol = row.cells[1] if len(row.cells) >= 2 else row.cells[0]
                                for child in list(c_sol._element):
                                    if child.tag.endswith('p') or child.tag.endswith('tbl'):
                                        block.addprevious(child)
                                        
                        # Đập vỡ khung bảng SAU KHI đã vét sạch Đề và Lời Giải
                        block.getparent().remove(block) 
                        
                    except Exception as e:
                        pass
            continue

        chunk_size = 15 if q_type == 'TN' else (5 if q_type == 'DS' else 8)
        col_w = Cm(1.13) if q_type == 'TN' else (Cm(3.392) if q_type == 'DS' else Cm(2.12)) 
        bg = HEX_BG_PART1 if q_type == 'TN' else (HEX_BG_PART2 if q_type == 'DS' else HEX_BG_PART3)

        for chunk_start in range(0, count, chunk_size):
            
            # [ĐÃ THÊM] Chèn dòng trắng để ngắt bảng, áp dụng chung cho cả Phần I, II, III
            if chunk_start > 0:
                spacer = doc_base.add_paragraph()
                spacer.paragraph_format.space_after = Pt(6)
            
            chunk_end = min(chunk_start + chunk_size, count)
            header_indices = list(range(chunk_start + 1, chunk_end + 1))
            
            num_cols = len(header_indices) + 1
            tbl = doc_base.add_table(rows=1, cols=num_cols)
            
            # [GIẢI PHÁP CHÍNH ĐẠO]: KHÔNG dùng Style của Word để tránh viền lỗi lề âm.
            # Tự kẻ viền và Căn trái chuẩn thư viện python-docx. File an toàn 100%!
            add_table_borders(tbl)
            from docx.enum.table import WD_TABLE_ALIGNMENT
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            widths = [Cm(WIDTH_COL_MADE)] + [col_w] * len(header_indices)
            set_col_widths(tbl, widths)

            r0 = tbl.rows[0]; c0 = r0.cells[0]; c0.text = "Mã đề"
            set_cell_shading(c0, bg); apply_compact_cell(c0)
            format_run(c0.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))
            for i, val in enumerate(header_indices):
                c = r0.cells[i+1]; c.text = str(val)
                set_cell_shading(c, bg); apply_compact_cell(c)
                format_run(c.paragraphs[0].runs[0], bold=True, color=RGBColor(255,255,255))

            sorted_exam_ids = sorted(mapped_data.keys(), key=lambda x: str(x))
            for exam_id in sorted_exam_ids:
                my_bucket = mapped_data[exam_id][q_type]
                row = tbl.add_row(); set_col_widths(tbl, widths)
                c_id = row.cells[0]; c_id.text = str(exam_id); apply_compact_cell(c_id)
                format_run(c_id.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE)
                for i, relative_idx in enumerate(range(chunk_start, chunk_end)):
                    key_text = ""
                    if relative_idx < len(my_bucket):
                        item = my_bucket[relative_idx]
                        raw_k = item['key']
                        if q_type == 'DS':
                            clean_k = re.sub(r'[^TFtfDdSs]', '', str(raw_k))
                            lbls = ['a', 'b', 'c', 'd']; formatted = []
                            for k_i, char in enumerate(clean_k[:4]):
                                is_true = char.upper() in ['T', 'D', 'Đ']
                                formatted.append(f"{lbls[k_i]}){'Đ' if is_true else 'S'}")
                            key_text = "-".join(formatted)
                        else: key_text = str(raw_k)
                    c_val = row.cells[i+1]; c_val.text = key_text; apply_compact_cell(c_val)
                    if c_val.paragraphs[0].runs: format_run(c_val.paragraphs[0].runs[0], bold=True, color=COLOR_BLUE, size=11)
                    else: r = c_val.paragraphs[0].add_run(key_text); format_run(r, bold=True, color=COLOR_BLUE, size=11)
        
# ============================================================================
# KHU VỰC 4: RENDER OPTION (KHÓA - CẤM THAY ĐỔI)
# ============================================================================
def determine_mode(options_data):
    for opt in options_data:
        cell = opt['cell']
        if not cell: continue
        valid_count = 0
        for child in cell._element.iterchildren():
            tag = get_tag(child)
            if tag == 'tbl':
                return 1
            elif tag == 'p':
                p = Paragraph(child, cell)
                if not is_p_empty(p): valid_count += 1
        if valid_count > 1: return 1 

    max_w = 0
    for opt in options_data:
        w = calculate_option_width(opt['cell']) + 0.1 
        if w > max_w: max_w = w
    final_w = max_w * SAFETY_FACTOR
    if final_w <= THRESH_MODE_4: return 4
    if final_w <= THRESH_MODE_2: return 2
    return 1

def render_options_with_tabs(doc, options_data):
    if not options_data: return
    mode = determine_mode(options_data)
    
    def create_p():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(0.5) 
        return p
        
    def copy_cell_to_paragraph_inline(source_cell, target_p):
        if not source_cell: return
        for child in source_cell._element.iterchildren():
            if get_tag(child) == 'p':
                p = Paragraph(child, source_cell)
                if not is_p_empty(p):
                    copy_paragraph_content(p, target_p); break
                
    if mode == 4:
        p = create_p(); ts = p.paragraph_format.tab_stops
        ts.add_tab_stop(Cm(4.74)); ts.add_tab_stop(Cm(9.48)); ts.add_tab_stop(Cm(14.22))
        for i in range(4):
            if i > 0: p.add_run("\t")
            r = p.add_run(options_data[i]['lbl'] + " "); format_run(r, bold=True, color=COLOR_BLUE)
            copy_cell_to_paragraph_inline(options_data[i]['cell'], p)
    elif mode == 2:
        p1 = create_p(); p1.paragraph_format.tab_stops.add_tab_stop(Cm(9.48)) 
        r = p1.add_run(options_data[0]['lbl'] + " "); format_run(r, bold=True, color=COLOR_BLUE)
        copy_cell_to_paragraph_inline(options_data[0]['cell'], p1); p1.add_run("\t")
        r = p1.add_run(options_data[1]['lbl'] + " "); format_run(r, bold=True, color=COLOR_BLUE)
        copy_cell_to_paragraph_inline(options_data[1]['cell'], p1)
        
        p2 = create_p(); p2.paragraph_format.tab_stops.add_tab_stop(Cm(9.48)) 
        r = p2.add_run(options_data[2]['lbl'] + " "); format_run(r, bold=True, color=COLOR_BLUE)
        copy_cell_to_paragraph_inline(options_data[2]['cell'], p2); p2.add_run("\t")
        r = p2.add_run(options_data[3]['lbl'] + " "); format_run(r, bold=True, color=COLOR_BLUE)
        copy_cell_to_paragraph_inline(options_data[3]['cell'], p2)
    else:
        for opt in options_data:
            start_p_count = len(doc.paragraphs)
            render_cell_content(doc, opt['cell'], first_line_prefix=opt['lbl'] + " ", first_line_color=COLOR_BLUE)
            for i in range(start_p_count, len(doc.paragraphs)):
                doc.paragraphs[i].paragraph_format.left_indent = Cm(0.5)

# ============================================================================
# KHU VỰC 5: XỬ LÝ DỮ LIỆU INPUT (KHÓA - CẤM THAY ĐỔI)
# ============================================================================
def identify_q_type(key_val):
    k = str(key_val).strip().upper().replace(".", "").replace(")", "")
    if not k: return 'TN'
    if k in ['A', 'B', 'C', 'D']: return 'TN'
    if len(k) == 4 and all(c in 'TFĐS' for c in k): return 'DS'
    return 'TLN'

def clean_key_string(raw_key_cell_text):
    text = re.sub(r"^(KEY|ĐÁP ÁN|LỜI GIẢI)[:\s]*", "", raw_key_cell_text, flags=re.IGNORECASE).strip()
    text = text.strip('"')
    if text.upper().startswith("TL"): return "TL"
    match = re.match(r"^([-]?\d+,\d+|[^,]+)", text)
    if match: return match.group(1).strip()
    return text

def extract_data_from_raw(doc_raw):
    items = []; idx = 1; part_idx = 1
    PART_HEADERS = ["PHẦN I", "PHẦN II", "PHẦN III", "PHẦN IV"]
    current_q = None
    
    def flush():
        nonlocal current_q, idx, part_idx
        if current_q: 
            current_q['q_type'] = identify_q_type(current_q['key'])
            items.append(current_q); idx += 1; part_idx += 1
            current_q = None
            
    for block in doc_raw.element.body.iterchildren():
        tag = get_tag(block)
        if tag == 'p':
            txt = "".join([t.text for t in block.findall('.//w:t', NSMAP) if t.text]).strip().upper()
            if any(h in txt for h in PART_HEADERS):
                flush(); part_idx = 1; items.append({'type': 'PART', 'obj': deepcopy(block)})
            else:
                clean_txt = re.sub(r'\s+', '', txt)
                if ("@" in clean_txt and "#" in clean_txt) or ("*" in clean_txt and "#" in clean_txt):
                    continue
                    
                has_content = False
                if txt: has_content = True
                else:
                    for child in block.iter():
                        c_tag = get_tag(child)
                        if c_tag in ['oMath', 'oMathPara', 'drawing', 'object', 'pict', 'AlternateContent', 'shape']:
                            has_content = True; break
                
                if has_content:
                    flush()
                    items.append({'type': 'TEXT', 'obj': deepcopy(block)}) 
                    
        elif tag == 'tbl':
            tbl = Table(block, doc_raw)
            is_question_tbl = False
            is_option_tbl = False
            
            for row in tbl.rows:
                try:
                    c0_raw = row.cells[0].text.strip()
                    c0_up = c0_raw.upper()
                    if "CÂU" in c0_up and len(c0_up) < 50 and len(row.cells) >= 2:
                        is_question_tbl = True; break
                    if c0_up in ['A', 'B', 'C', 'D', 'A.', 'B.', 'C.', 'D.', 'A)', 'B)', 'C)', 'D)']:
                        is_option_tbl = True
                    if "KEY" in c0_up or "ĐÁP ÁN" in c0_up or "LỜI GIẢI" in c0_up:
                        is_option_tbl = True
                except: pass
                
            if is_question_tbl:
                flush()
                current_q = {'id': str(idx), 'num': str(part_idx), 'stem': None, 'options': [], 'key': '', 'solution': None, 'type': 'QUESTION'}
                for row in tbl.rows:
                    try:
                        c0_raw = row.cells[0].text.strip()
                        c0_up = c0_raw.upper()
                        if "CÂU" in c0_up and len(c0_up) < 50 and len(row.cells) >= 2:
                            current_q['stem'] = row.cells[1]
                        elif c0_up in ['A', 'B', 'C', 'D', 'A.', 'B.', 'C.', 'D.']:
                            lbl = c0_up.replace('.', '') + "." 
                            current_q['options'].append({'lbl': lbl, 'cell': row.cells[1]})
                        elif c0_up in ['A)', 'B)', 'C)', 'D)']:
                            lbl = c0_raw[0].lower() + ")" 
                            current_q['options'].append({'lbl': lbl, 'cell': row.cells[1]})
                        elif "KEY" in c0_up or "ĐÁP ÁN" in c0_up: 
                            current_q['key'] = clean_key_string(c0_raw)
                            if len(row.cells) >= 2: current_q['solution'] = row.cells[1]
                            else: current_q['solution'] = row.cells[0]
                        elif "LỜI GIẢI" in c0_up: 
                            current_q['solution'] = row.cells[1] if len(row.cells) >= 2 else row.cells[0]
                    except: pass
            elif is_option_tbl and current_q:
                for row in tbl.rows:
                    try:
                        c0_raw = row.cells[0].text.strip()
                        c0_up = c0_raw.upper()
                        if c0_up in ['A', 'B', 'C', 'D', 'A.', 'B.', 'C.', 'D.']:
                            lbl = c0_up.replace('.', '') + "." 
                            current_q['options'].append({'lbl': lbl, 'cell': row.cells[1]})
                        elif c0_up in ['A)', 'B)', 'C)', 'D)']:
                            lbl = c0_raw[0].lower() + ")" 
                            current_q['options'].append({'lbl': lbl, 'cell': row.cells[1]})
                        elif "KEY" in c0_up or "ĐÁP ÁN" in c0_up: 
                            current_q['key'] = clean_key_string(c0_raw)
                            if len(row.cells) >= 2: current_q['solution'] = row.cells[1]
                            else: current_q['solution'] = row.cells[0]
                        elif "LỜI GIẢI" in c0_up: 
                            current_q['solution'] = row.cells[1] if len(row.cells) >= 2 else row.cells[0]
                    except: pass
            else:
                flush()
                try:
                    c0_text = tbl.rows[0].cells[0].text.strip().lower()
                    # [ĐÃ SỬA]: Nếu là bảng marker 'p', 'ht', '*', '@' -> Chỉ lấy ô nội dung bên phải
                    if len(tbl.rows[0].cells) >= 2 and c0_text in ['p', 'ht', '@', '*', '']:
                        items.append({'type': 'TEXT_CELL', 'obj': tbl.rows[0].cells[1]})
                    else:
                        # Nếu là bảng dữ liệu thật của bài toán -> Lấy nguyên bảng
                        items.append({'type': 'TEXT', 'obj': deepcopy(block)})
                except:
                    items.append({'type': 'TEXT', 'obj': deepcopy(block)})
    flush(); return items
# ============================================================================
# KHU VỰC 6: ĐÁNH LẠI THỨ TỰ PHẦN (NEW - SAFE MODE)
# ============================================================================
def renumber_sections_in_doc(doc):
    romans = ["I", "II", "III", "IV", "V", "VI"]
    part_count = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.upper().startswith("PHẦN"):
            match = re.search(r"^(PHẦN\s+)([0-9IVX]+)(.*)", text, re.IGNORECASE)
            if match:
                if part_count < len(romans): new_roman = romans[part_count]
                else: new_roman = str(part_count + 1)
                part_count += 1
                p.text = f"{match.group(1)}{new_roman}{match.group(3)}"
                if p.runs:
                    run = p.runs[0]
                    run.font.name = FONT_NAME
                    run.font.size = Pt(FONT_SIZE) 
                    run.font.bold = True
                    run.font.color.rgb = COLOR_BLUE

# ============================================================================
# HÀM CHÍNH (FINAL)
# ============================================================================
import pandas as pd
from openpyxl.styles import Font

def generate_excel_dapan(all_results, mode=1):
    def local_strict_detect(key_val):
        k_up = str(key_val).strip().upper()
        if k_up == 'TL' or 'TỰ LUẬN' in k_up: return 'TL'
        clean_tn = k_up.replace(".", "").replace(")", "").strip()
        if clean_tn in ['A', 'B', 'C', 'D']: return 'TN'
        clean_ds = re.sub(r'[^TFĐS]', '', k_up)
        if len(clean_ds) == 4: return 'DS'
        return 'TLN'

    sorted_results = sorted(all_results, key=lambda x: str(x['exam_id']))
    
    # 1. Gom nhóm câu hỏi, LOẠI BỎ Tự Luận và ĐÁNH LẠI số thứ tự từ 1 cho từng phần
    ref_items = sorted_results[0]['items']
    parts = {'TN': [], 'DS': [], 'TLN': [], 'TL': []}
    for item in ref_items:
        if item['type'] == 'QUESTION':
            q_type = local_strict_detect(item.get('key', ''))
            parts[q_type].append(item)
            
    ref_seq = []
    for p_type in ['TN', 'DS', 'TLN']: # Cố tình bỏ qua 'TL'
        for idx, item in enumerate(parts[p_type]):
            ref_seq.append({'id': item['id'], 'new_num': idx + 1, 'type': p_type})
    
    # 2. Xử lý format Đáp án (ĐĐSS)
    exam_answers = {}
    for res in sorted_results:
        exam_id = res['exam_id']
        ans_dict = {}
        for item in res['items']:
            if item['type'] == 'QUESTION':
                q_type = local_strict_detect(item.get('key', ''))
                raw_k = item.get('key', '')
                if q_type == 'DS':
                    clean_k = re.sub(r'[^TFtfDdSs]', '', str(raw_k))
                    formatted = []
                    for char in clean_k[:4]:
                        is_true = char.upper() in ['T', 'D', 'Đ']
                        formatted.append('Đ' if is_true else 'S')
                    display_key = "".join(formatted) # Gộp thành ĐĐSS
                else:
                    display_key = str(raw_k)
                ans_dict[item['id']] = display_key
        exam_answers[exam_id] = ans_dict

    # 3. Dàn trang theo Mode
    data = []
    if mode == 1:
        for res in sorted_results:
            exam_id = res['exam_id']
            for q_info in ref_seq:
                q_id = q_info['id']
                data.append({'Mã đề': exam_id, 'Câu': q_info['new_num'], 'Đáp án': exam_answers[exam_id].get(q_id, '')})
        df = pd.DataFrame(data)
        
    elif mode == 2:
        for q_info in ref_seq:
            q_id = q_info['id']
            row = {'Câu': q_info['new_num']}
            for exam_id in exam_answers: row[str(exam_id)] = exam_answers[exam_id].get(q_id, '')
            data.append(row)
        df = pd.DataFrame(data)

    elif mode == 3:
        q_nums = [q_info['new_num'] for q_info in ref_seq]
        for exam_id in exam_answers:
            row1 = [f"Mã đề: {exam_id}"] + q_nums
            row2 = ["Đáp án"] + [exam_answers[exam_id].get(q_info['id'], '') for q_info in ref_seq]
            data.append(row1); data.append(row2); data.append([''] * len(row1)) 
        df = pd.DataFrame(data)

    elif mode == 4:
        q_nums = [q_info['new_num'] for q_info in ref_seq]
        for exam_id in exam_answers:
            row = [exam_id] + [exam_answers[exam_id].get(q_info['id'], '') for q_info in ref_seq]
            data.append(row)
        columns = ['Mã đề'] + q_nums
        df = pd.DataFrame(data, columns=columns)
    else:
        df = pd.DataFrame()

    # 4. Xuất file & Ép Font Times New Roman 12
    excel_buffer = io.BytesIO()
    with pd.ExcelWriter(excel_buffer, engine='openpyxl') as writer:
        if mode == 3: df.to_excel(writer, index=False, header=False)
        else: df.to_excel(writer, index=False)
        
        worksheet = writer.sheets['Sheet1']
        font = Font(name='Times New Roman', size=12)
        for row in worksheet.iter_rows():
            for cell in row:
                cell.font = font
                
    excel_buffer.seek(0)
    return excel_buffer
    
def xuat_ket_qua(data_mixed, config, output_folder, stats):
    zip_buffer = io.BytesIO()
    all_results = []
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, d in enumerate(data_mixed):
            exam_id = d.get('exam_id', '001')
            tmp = io.BytesIO(); d.get('file_content').save(tmp); raw_bytes = tmp.getvalue()
            doc_scan = Document(io.BytesIO(raw_bytes))
            items = extract_data_from_raw(doc_scan)
            result_entry = {'exam_id': exam_id, 'items': items}
            if i == 0: result_entry['raw_bytes'] = raw_bytes
            all_results.append(result_entry)
            
          # Made (CÓ CHẠY AREA 6)
            doc_made = Document(io.BytesIO(raw_bytes))
            
            # --- Ép font mặc định cho file Mã Đề ---
            style_made = doc_made.styles['Normal']
            style_made.font.name = FONT_NAME
            style_made.font.size = Pt(FONT_SIZE)
            
            clear_body_preserve_section(doc_made); set_narrow_layout(doc_made)
            generate_made_file(doc_made, items, config, exam_id)
            renumber_sections_in_doc(doc_made)
            out_made = io.BytesIO(); doc_made.save(out_made)
            zf.writestr(f"made_{exam_id}.docx", out_made.getvalue())
            
            # Loigiai (CÓ CHẠY AREA 6)
            doc_sol = Document(io.BytesIO(raw_bytes))
            
            # --- Ép font mặc định cho file Lời Giải ---
            style_sol = doc_sol.styles['Normal']
            style_sol.font.name = FONT_NAME
            style_sol.font.size = Pt(FONT_SIZE)
            
            clear_body_preserve_section(doc_sol); set_narrow_layout(doc_sol)
            generate_loigiai_file(doc_sol, items, config, exam_id)
            renumber_sections_in_doc(doc_sol)
            out_sol = io.BytesIO(); doc_sol.save(out_sol)
            zf.writestr(f"loigiai_{exam_id}.docx", out_sol.getvalue())

        all_results.sort(key=lambda x: str(x['exam_id']))
        
        # Dapan Tonghop (KHÔNG CHẠY AREA 6)
        if all_results and 'raw_bytes' in all_results[0]:
            doc_sum = Document(io.BytesIO(all_results[0]['raw_bytes']))
        else:
            doc_sum = Document()
            
        set_narrow_layout(doc_sum)
        generate_dapan_tonghop(doc_sum, all_results, config)
        
        # Chỉ lưu 1 lần duy nhất
        out_sum = io.BytesIO(); doc_sum.save(out_sum)
        zf.writestr("dapan_tonghop.docx", out_sum.getvalue())
        
        # [TÍNH NĂNG MỚI] - XUẤT EXCEL CHO KHTN
        try:
            excel_mode = int(config.get('excel_mode', 1))
            excel_buf = generate_excel_dapan(all_results, mode=excel_mode)
            zf.writestr("dapan_tonghop.xlsx", excel_buf.getvalue())
        except Exception as e:
            import streamlit as st
            st.error(f"⚠️ Lỗi xuất Excel: {e}. Vui lòng chạy lệnh 'pip install openpyxl' trong Terminal!")
            
    zip_buffer.seek(0)
    return zip_buffer