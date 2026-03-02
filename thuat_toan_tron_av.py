import random
import re
import time
import io
from copy import deepcopy
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml import OxmlElement
from docx.shared import RGBColor

# Kích hoạt bộ sinh số ngẫu nhiên theo thời gian thực
random.seed(time.time())

# =============================================================================
# KHU VỰC 1: CÁC HÀM HỖ TRỢ VÀ XỬ LÝ BỘ NHỚ RAM
# =============================================================================
def tao_dong_trang():
    """Tạo dòng trắng chuẩn XML."""
    return OxmlElement('w:p')

def doc_file_tu_ram(doc_input):
    """Đọc an toàn từ luồng RAM (BytesIO) hoặc file gốc, chống văng lỗi (Crash)."""
    if isinstance(doc_input, io.BytesIO):
        doc_input.seek(0)
        return Document(io.BytesIO(doc_input.read()))
    elif hasattr(doc_input, 'save'):
        stream = io.BytesIO()
        doc_input.save(stream)
        stream.seek(0)
        return Document(stream)
    elif isinstance(doc_input, str):
        return Document(doc_input)

def lay_chu_cot_trai(block, doc):
    """Lấy danh sách text của tất cả các dòng trong Cột Trái của bảng."""
    if isinstance(block, CT_Tbl):
        tbl = Table(block, doc)
        texts = []
        try:
            for row in tbl.rows:
                if row.cells:
                    txt = row.cells[0].text.strip()
                    if txt: texts.append(txt)
        except: pass
        return texts
    if isinstance(block, CT_P):
        return [Paragraph(block, doc).text.strip()]
    return []

def nhan_dien_the_nhom(block, doc):
    """Nhận diện các thẻ nhóm (#*#, #@#...) từ text thô."""
    text = ""
    if isinstance(block, CT_Tbl):
        tbl = Table(block, doc)
        try:
            for c in tbl.rows[0].cells: text += c.text
        except: pass
    else:
        text = Paragraph(block, doc).text

    clean = re.sub(r'\s+', '', text).upper()
    if "@" in clean and "#" in clean: return "BAT_DAU_GHIM"
    if "**" in clean and "#" in clean: return "KET_THUC"
    if "*" in clean and "#" in clean: return "BAT_DAU_THUONG"
    return None

def kiem_tra_ghim_vi_tri(item, doc):
    """Kiểm tra xem câu hỏi có chứa ký tự @ để ghim lên đầu không."""
    if item['type'] == 'GROUP' and item.get('fixed', False): return False
    obj = None
    if item['type'] == 'SINGLE': obj = item['obj']
    elif item['type'] == 'GROUP': obj = item['header']
    elif item['type'] == 'PARA': obj = item['obj']
    
    if obj:
        texts = lay_chu_cot_trai(obj, doc)
        for t in texts:
            if '@' in t: return True
    return False

def phan_loai_bang(block, doc):
    """
    Phân loại bảng: 'TRAC_NGHIEM', 'TU_LUAN', hoặc 'DOAN_VAN'.
    Quét sâu mọi dòng trong cột trái để nhận diện.
    """
    if not isinstance(block, CT_Tbl): return 'DOAN_VAN'
    col0_texts = lay_chu_cot_trai(block, doc)
    if not col0_texts: return 'DOAN_VAN'
    if col0_texts[0].upper() == 'P': return 'DOAN_VAN'
    
    is_question = False
    pattern = r'^\s*([!@])?\s*(CÂU|QUESTION)\s+\d+'
    
    for txt in col0_texts:
        if re.match(pattern, txt.upper()):
            is_question = True; break
            
    if is_question:
        has_A = False; has_B = False; has_C = False; has_D = False
        for t in col0_texts:
            t_up = t.upper().replace(".", "").replace(")", "").strip()
            if t_up == 'A': has_A = True
            elif t_up == 'B': has_B = True
            elif t_up == 'C': has_C = True
            elif t_up == 'D': has_D = True
            
        if has_A and has_B and has_C and has_D: return 'TRAC_NGHIEM'
        return 'TU_LUAN'
    return 'DOAN_VAN'

# =============================================================================
# KHU VỰC 2: TẦNG 1 - ĐẢO PHƯƠNG ÁN (CHỈ DÀNH CHO TRẮC NGHIỆM)
# =============================================================================
def dao_phuong_an_trac_nghiem(table, config):
    rows = table.rows
    if len(rows) < 2: return
    try:
        # Nếu có dấu chấm than (!) thì không đảo
        for r in rows:
            if r.cells and '!' in r.cells[0].text: return
    except: return

    map_lbl = {}
    key_row = None; key_origin = ""
    
    for i, r in enumerate(rows):
        txt = r.cells[0].text.strip().upper().replace(".", "").replace(")", "")
        if txt.startswith("KEY"):
            key_row = r
            parts = r.cells[0].text.split(":")
            raw_k = parts[-1].strip() if len(parts) > 1 else r.cells[0].text.upper().replace("KEY", "").strip()
            key_origin = raw_k
            continue
        if txt in ['A', 'B', 'C', 'D']:
            map_lbl[txt] = i
    
    if len(map_lbl) < 4 or not config.get('tron_mcq', True): return

    labels = ['A', 'B', 'C', 'D']
    permuted = labels[:]
    random.shuffle(permuted)

    xml_cache = {}
    for lbl in labels:
        r_idx = map_lbl[lbl]
        xml_cache[lbl] = deepcopy(rows[r_idx].cells[1]._element)

    for i, target_lbl in enumerate(labels):
        src_lbl = permuted[i]
        r_idx = map_lbl[target_lbl]
        cell = rows[r_idx].cells[1]._element
        cell.clear_content()
        for child in xml_cache[src_lbl]:
            cell.append(deepcopy(child))
            
    if key_row and key_origin:
        clean_k = key_origin.replace(".", "").strip()
        if clean_k in permuted:
            idx = permuted.index(clean_k)
            new_key = labels[idx]
            
            # Sửa lỗi mất màu đỏ của chữ KEY khi đảo đáp án
            cell_key = key_row.cells[0]
            for p in cell_key.paragraphs:
                p._element.getparent().remove(p._element)
            
            p_new = cell_key.add_paragraph()
            r = p_new.add_run(f"KEY: {new_key}")
            r.bold = True
            r.font.color.rgb = RGBColor(255, 0, 0)

def tang1_dao_noi_dung(items, config, doc):
    for item in items:
        if item['type'] == 'SINGLE':
            if isinstance(item['obj'], CT_Tbl) and phan_loai_bang(item['obj'], doc) == 'TRAC_NGHIEM':
                dao_phuong_an_trac_nghiem(Table(item['obj'], doc), config)
        elif item['type'] == 'GROUP':
            for q_blk in item['qs']:
                if isinstance(q_blk, CT_Tbl) and phan_loai_bang(q_blk, doc) == 'TRAC_NGHIEM':
                    dao_phuong_an_trac_nghiem(Table(q_blk, doc), config)
    return items

# =============================================================================
# KHU VỰC 3: TẦNG 2 VÀ TẦNG 3 - ĐẢO TRONG NHÓM VÀ ĐẢO TOÀN ĐỀ
# =============================================================================
def tang2_dao_trong_nhom(items, config, doc):
    for item in items:
        if item['type'] == 'GROUP':
            if item.get('fixed', False): continue
            qs = item['qs']
            doan_van_chung = []; cac_cau_hoi = []  
            for block in qs:
                kind = phan_loai_bang(block, doc)
                if kind in ['TRAC_NGHIEM', 'TU_LUAN']: cac_cau_hoi.append(block)
                else: doan_van_chung.append(block)
            
            ghim_vi_tri = []; khong_ghim = []
            for b in cac_cau_hoi:
                is_p = False
                try:
                    texts = lay_chu_cot_trai(b, doc)
                    for t in texts:
                        if '@' in t: is_p = True; break
                except: pass
                if is_p: ghim_vi_tri.append(b)
                else: khong_ghim.append(b)
            
            random.shuffle(khong_ghim)
            item['qs'] = doan_van_chung + ghim_vi_tri + khong_ghim
    return items

def xoc_danh_sach(sub_items, doc):
    ghim_vi_tri = []; khong_ghim = []
    for item in sub_items:
        if kiem_tra_ghim_vi_tri(item, doc): ghim_vi_tri.append(item)
        else: khong_ghim.append(item)
    random.shuffle(khong_ghim)
    return ghim_vi_tri + khong_ghim

def tang3_dao_toan_de(items, config, doc):
    if config.get('tron_nhom', False):
        return xoc_danh_sach(items, doc)
    else:
        final_list = []; buffer = []
        for item in items:
            if item['type'] == 'GROUP' or item['type'] == 'PARA':
                if buffer:
                    final_list.extend(xoc_danh_sach(buffer, doc))
                    buffer = []
                final_list.append(item)
            else:
                buffer.append(item)
        if buffer: final_list.extend(xoc_danh_sach(buffer, doc))
        return final_list

# =============================================================================
# KHU VỰC 4: ĐIỀN SỐ THỨ TỰ (CÂU, PHẦN LA MÃ, SỐ TỰ NHIÊN)
# =============================================================================
def thay_the_chu_trong_paragraph(paragraph, replacements):
    if not paragraph.runs: return
    for run in paragraph.runs:
        for key, val in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, str(val))

def dien_so_vao_doan_van_nhom(header, doan_van_chung, start_num, end_num, doc):
    se_replacements = {'<S>': f"{start_num}", '<E>': f"{end_num}"}
    if header:
        if isinstance(header, CT_P): thay_the_chu_trong_paragraph(Paragraph(header, doc), se_replacements)
        elif isinstance(header, CT_Tbl):
            tbl = Table(header, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: thay_the_chu_trong_paragraph(p, se_replacements)
    
    current_fill_idx = start_num
    for block in doan_van_chung:
        target_paragraphs = []
        if isinstance(block, CT_P): target_paragraphs.append(Paragraph(block, doc))
        elif isinstance(block, CT_Tbl):
            tbl = Table(block, doc)
            for row in tbl.rows:
                for cell in row.cells: target_paragraphs.extend(cell.paragraphs)
        
        for p in target_paragraphs:
            thay_the_chu_trong_paragraph(p, se_replacements)
            if '[#]' in p.text:
                for run in p.runs:
                    while '[#]' in run.text:
                        val = f"({current_fill_idx})"
                        run.text = run.text.replace('[#]', val, 1)
                        current_fill_idx += 1

def so_nguyen_sang_la_ma(num):
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syb = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num

def dien_so_la_ma_va_tu_nhien(doc):
    """Quét toàn bộ văn bản để biến <<x>> thành I, II... và <<y>> thành 1, 2..."""
    x_counter = 1
    y_counter = 1
    
    def thay_the(p):
        nonlocal x_counter, y_counter
        # Chống đứt gãy text trong thẻ XML (an toàn tuyệt đối)
        while '<<x>>' in p.text:
            da_thay_the = False
            for run in p.runs:
                if '<<x>>' in run.text:
                    run.text = run.text.replace('<<x>>', so_nguyen_sang_la_ma(x_counter), 1)
                    x_counter += 1; da_thay_the = True; break
            if not da_thay_the: # Nếu bị Word cắt vụn, ép thay thế thô
                p.text = p.text.replace('<<x>>', so_nguyen_sang_la_ma(x_counter), 1)
                x_counter += 1

        while '<<y>>' in p.text:
            da_thay_the = False
            for run in p.runs:
                if '<<y>>' in run.text:
                    run.text = run.text.replace('<<y>>', str(y_counter), 1)
                    y_counter += 1; da_thay_the = True; break
            if not da_thay_the:
                p.text = p.text.replace('<<y>>', str(y_counter), 1)
                y_counter += 1

    for p in doc.paragraphs: thay_the(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs: thay_the(p)

# =============================================================================
# KHU VỰC 5: BỘ NHẬN DIỆN VÀ ĐÓNG GÓI DỮ LIỆU
# =============================================================================
def dong_goi_du_lieu(doc):
    sections = []; cac_thanh_phan_chinh = []; nhom_hien_tai = None 
    all_blocks = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P) or isinstance(child, CT_Tbl):
            all_blocks.append(child)

    for block in all_blocks:
        if isinstance(block, CT_P):
            if not Paragraph(block, doc).text.strip(): continue

        m_type = nhan_dien_the_nhom(block, doc)
        if m_type in ["BAT_DAU_THUONG", "BAT_DAU_GHIM"]:
            if nhom_hien_tai: cac_thanh_phan_chinh.append(nhom_hien_tai)
            nhom_hien_tai = {'type': 'GROUP', 'header': block, 'qs': [], 'footer': None, 'fixed': (m_type == "BAT_DAU_GHIM")}
            continue
        
        if m_type == "KET_THUC":
            if nhom_hien_tai:
                nhom_hien_tai['footer'] = block
                cac_thanh_phan_chinh.append(nhom_hien_tai)
                nhom_hien_tai = None
            continue
        
        if nhom_hien_tai: nhom_hien_tai['qs'].append(block)
        else:
            kind = phan_loai_bang(block, doc)
            if kind in ['TRAC_NGHIEM', 'TU_LUAN']: cac_thanh_phan_chinh.append({'type': 'SINGLE', 'obj': block})
            else: cac_thanh_phan_chinh.append({'type': 'PARA', 'obj': block})

    if nhom_hien_tai: cac_thanh_phan_chinh.append(nhom_hien_tai)
    sections.append({'header': [], 'items': cac_thanh_phan_chinh})
    return sections

# =============================================================================
# KHU VỰC 6: LUỒNG CHẠY CHÍNH (HÀM TRỘN ĐỀ)
# =============================================================================
def tron_de(input_source, so_luong_de, config=None):
    if config is None: config = {}
    list_result = []
    
    random.seed(time.time())
    start_id = config.get('ma_de_start', 101)
    kieu_ma = config.get('kieu_ma_de', 'SEQUENTIAL')
    generated_ids = set()
    
    for i in range(so_luong_de):
        if kieu_ma == 'SEQUENTIAL': exam_id = str(start_id + i)
        else:
            while True:
                rid = random.randint(100, 999) if kieu_ma == 'RANDOM_3' else random.randint(1000, 9999)
                if str(rid) not in generated_ids: generated_ids.add(str(rid)); exam_id = str(rid); break
        
        # Đọc dữ liệu 100% trên RAM
        doc = doc_file_tu_ram(input_source)
        sections = dong_goi_du_lieu(doc)
        
        for section in sections:
            if section['items']:
                items = section['items']
                
                # --- BƯỚC 1: TRỘN ĐỀ ---
                items = tang1_dao_noi_dung(items, config, doc)
                items = tang2_dao_trong_nhom(items, config, doc)
                # Gán lại mảng sau khi xóc (Lỗi cũ đã được sửa)
                section['items'] = tang3_dao_toan_de(items, config, doc)
                
                # --- BƯỚC 2: ĐÁNH SỐ CÂU ---
                current_idx = 1
                detected_prefix = "Câu"
                prefix_locked = False
                
                for item in section['items']:
                    list_qs = []
                    if item['type'] == 'SINGLE': list_qs.append(item['obj'])
                    elif item['type'] == 'GROUP':
                        for b in item['qs']:
                            if phan_loai_bang(b, doc) in ['TRAC_NGHIEM', 'TU_LUAN']: list_qs.append(b)
                    
                    for q_block in list_qs:
                        tbl = Table(q_block, doc)
                        try:
                            target_row = None
                            for row in tbl.rows:
                                if row.cells and row.cells[0].text.strip():
                                    if re.match(r'^\s*([!@])?\s*(CÂU|QUESTION)\s+\d+', row.cells[0].text.strip().upper()):
                                        target_row = row; break
                            
                            if target_row:
                                c0 = target_row.cells[0].text.strip()
                                if not prefix_locked and ("Question" in c0 or "Câu" in c0):
                                    detected_prefix = "Question" if "Question" in c0 else "Câu"
                                    prefix_locked = True
                                
                                match = re.match(r'^\s*([!@])?', c0)
                                prefix = match.group(1) if match and match.group(1) else ""
                                
                                target_row.cells[0].text = f"{prefix}{detected_prefix} {current_idx}"
                                current_idx += 1
                        except: pass
                
                # --- BƯỚC 3: LẤP LỖ TRỐNG (PLACEHOLDERS) ---
                for item in section['items']:
                    if item['type'] == 'GROUP':
                        doan_van_chung = []; cau_hoi_trong_nhom = []
                        for b in item['qs']:
                            if phan_loai_bang(b, doc) in ['TRAC_NGHIEM', 'TU_LUAN']: cau_hoi_trong_nhom.append(b)
                            else: doan_van_chung.append(b)
                        
                        if cau_hoi_trong_nhom:
                            def get_q_num(blk):
                                t = Table(blk, doc)
                                for r in t.rows:
                                    if r.cells:
                                        txt = r.cells[0].text.strip()
                                        nums = re.findall(r'\d+', txt)
                                        if nums: return int(nums[-1])
                                return None
                                
                            start_num = get_q_num(cau_hoi_trong_nhom[0])
                            end_num = get_q_num(cau_hoi_trong_nhom[-1])
                            
                            if start_num is not None and end_num is not None:
                                dien_so_vao_doan_van_nhom(item['header'], doan_van_chung, start_num, end_num, doc)

        # --- BƯỚC 4: KẾT XUẤT RA LUỒNG ---
        doc.element.body.clear_content()
        doc.element.body.append(tao_dong_trang())
        
        for section in sections:
            final_stream = []
            final_stream.extend(section['header'])
            
            if section['items']:
                for item in section['items']:
                    if item['type'] == 'PARA': final_stream.append(item['obj'])
                    elif item['type'] == 'SINGLE': final_stream.append(item['obj'])
                    elif item['type'] == 'GROUP':
                        final_stream.append(item['header'])
                        final_stream.extend(item['qs'])
                        if item['footer']: final_stream.append(item['footer'])
            
            for block in final_stream:
                doc.element.body.append(block)
                if isinstance(block, CT_Tbl):
                    doc.element.body.append(tao_dong_trang())

        # Chạy thuật toán dịch <<x>> thành số La Mã và <<y>> thành số tự nhiên cho toàn bộ văn bản
        dien_so_la_ma_va_tu_nhien(doc)

        list_result.append({'exam_id': exam_id, 'file_content': doc})

    return list_result