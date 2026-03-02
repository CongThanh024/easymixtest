import random
import re
import io 
from copy import deepcopy
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml import OxmlElement

# =============================================================================
# A. CÁC HÀM HỖ TRỢ (UTILS)
# =============================================================================
def create_empty_p():
    """Tạo dòng trắng chuẩn XML."""
    return OxmlElement('w:p')

def get_block_text(block, doc):
    """Lấy text từ Table hoặc Paragraph."""
    if isinstance(block, CT_P):
        return Paragraph(block, doc).text.strip()
    elif isinstance(block, CT_Tbl):
        tbl = Table(block, doc)
        try:
            full_text = ""
            for row in tbl.rows:
                for cell in row.cells:
                    full_text += cell.text
            return full_text.strip()
        except: return ""
    return ""

def identify_marker(block, doc):
    """Xác định xem block này có phải là Marker không."""
    text = get_block_text(block, doc)
    clean = re.sub(r'\s+', '', text).upper()
    if "@" in clean and "#" in clean: return "START_FIXED"
    if "**" in clean and "#" in clean: return "END"
    if "*" in clean and "#" in clean: return "START_NORMAL"
    return None

def check_ghim_vi_tri(item, doc):
    """Kiểm tra xem item này có ghim @ không (để đưa lên đầu)."""
    if item['type'] == 'GROUP' and item.get('fixed', False): return False
    obj = None
    if item['type'] == 'SINGLE': obj = item['obj']
    elif item['type'] == 'GROUP': obj = item['header']
    elif item['type'] == 'PARA': obj = item['obj']
    
    if obj:
        txt = get_block_text(obj, doc)
        return '@' in txt
    return False

def _clone_document_in_ram(doc_input):
    """Nhân bản vô tính (Deepcopy) an toàn 100% trên RAM (Hỗ trợ đa định dạng đầu vào)."""
    # 1. Nếu đầu vào là luồng RAM thô (BytesIO)
    if isinstance(doc_input, io.BytesIO):
        doc_input.seek(0)
        new_stream = io.BytesIO(doc_input.read())
        return Document(new_stream)
        
    # 2. Nếu đầu vào đã là Đối tượng Document
    elif hasattr(doc_input, 'save'):
        stream = io.BytesIO()
        doc_input.save(stream)
        stream.seek(0)
        return Document(stream)
        
    # 3. Nếu đầu vào là đường dẫn file cứng (String)
    elif isinstance(doc_input, str):
        return Document(doc_input)

def _replace_text_in_p(paragraph, old_text, new_text):
    """Thay thế văn bản (dành cho biến <S> và <E>) mà không làm hỏng XML."""
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
    if old_text in paragraph.text:
        paragraph.text = paragraph.text.replace(old_text, new_text)

# =============================================================================
# B. LỚP 1: TRỘN PHƯƠNG ÁN (A,B,C,D)
# =============================================================================
def _dao_phuong_an_table(table, config):
    rows = table.rows
    if not rows: return
    
    try:
        if '!' in rows[0].cells[0].text: return
    except: return

    map_lbl = {}
    key_row = None; key_origin = ""
    is_mcq = False; is_ds = False

    for i, r in enumerate(rows):
        txt = r.cells[0].text.strip()
        if txt.upper().startswith("KEY"):
            key_row = r
            if ":" in txt: key_origin = txt.split(":", 1)[1].strip()
            else: key_origin = txt.replace("KEY", "").strip()
            if "," in key_origin: key_origin_clean = key_origin.split(",")[0].strip()
            else: key_origin_clean = key_origin
            continue
        
        lbl = txt.replace(".", "").replace(")", "").strip()
        if lbl in ['A', 'B', 'C', 'D']: map_lbl[lbl] = i; is_mcq = True
        elif lbl in ['a', 'b', 'c', 'd']: map_lbl[lbl] = i; is_ds = True

    if is_mcq and not config.get('tron_mcq', True): return
    if is_ds and not config.get('tron_ds', True): return

    labels = sorted(list(map_lbl.keys()))
    if len(labels) < 2: return

    permuted = labels[:]
    random.shuffle(permuted)

    xml_cache = {}
    for lbl in labels:
        xml_cache[lbl] = deepcopy(rows[map_lbl[lbl]].cells[1]._element)

    for i, target_lbl in enumerate(labels):
        src_lbl = permuted[i]
        r_idx = map_lbl[target_lbl]
        cell = rows[r_idx].cells[1]._element
        cell.clear_content()
        for child in xml_cache[src_lbl]:
            cell.append(deepcopy(child))

    if key_row and key_origin:
        if is_mcq:
            clean_k = key_origin.replace(".", "").replace(")", "").split(",")[0].strip()
            if clean_k in permuted:
                idx = permuted.index(clean_k)
                new_key = labels[idx]
                old_full = key_row.cells[0].text
                suffix = ""
                if "," in old_full: 
                    parts = old_full.split(",", 1)
                    if len(parts) > 1: suffix = "," + parts[1]
                key_row.cells[0].text = f"KEY: {new_key}{suffix}"
        elif is_ds:
            truth_chars = re.findall(r'[TFĐS]', key_origin.upper())
            if len(truth_chars) == len(labels):
                origin_map = {lbl: truth_chars[idx] for idx, lbl in enumerate(labels)}
                new_values = [origin_map[src] for src in permuted]
                new_key_str = "".join(new_values)
                key_row.cells[0].text = f"KEY: {new_key_str}"

def layer1_dao_noi_dung(items, config, doc):
    for item in items:
        if item['type'] == 'SINGLE':
            if isinstance(item['obj'], CT_Tbl):
                _dao_phuong_an_table(Table(item['obj'], doc), config)
        elif item['type'] == 'GROUP':
            for q_blk in item['qs']:
                if isinstance(q_blk, CT_Tbl):
                    _dao_phuong_an_table(Table(q_blk, doc), config)
    return items

# =============================================================================
# C. LỚP 2: TRỘN TRONG NHÓM (INTRA-GROUP)
# =============================================================================
def layer2_tron_trong_nhom(items, config, doc):
    for item in items:
        if item['type'] == 'GROUP':
            if item.get('fixed', False): continue
            
            qs = item['qs']
            pinned_q = []
            normal_q = []
            
            for b in qs:
                is_p = False
                if isinstance(b, CT_Tbl):
                    try:
                        if '@' in Table(b, doc).rows[0].cells[0].text: is_p = True
                    except: pass
                
                if is_p: pinned_q.append(b)
                else: normal_q.append(b)
            
            random.shuffle(normal_q)
            item['qs'] = pinned_q + normal_q
    return items

# =============================================================================
# D. LỚP 3: TRỘN CẤU TRÚC LỚN (GLOBAL STRUCTURE)
# =============================================================================
def _xoc_list(sub_items, doc):
    pinned = []
    normal = []
    for item in sub_items:
        if check_ghim_vi_tri(item, doc): pinned.append(item)
        else: normal.append(item)
    random.shuffle(normal)
    return pinned + normal

def layer3_tron_cau_truc(items, config, doc):
    if config.get('tron_nhom', False):
        return _xoc_list(items, doc)
    else:
        final_list = []
        buffer = [] 
        for item in items:
            if item['type'] == 'GROUP' or item['type'] == 'PARA':
                if buffer:
                    shuffled_buf = _xoc_list(buffer, doc)
                    final_list.extend(shuffled_buf)
                    buffer = []
                final_list.append(item)
            else:
                buffer.append(item)
        if buffer:
            shuffled_buf = _xoc_list(buffer, doc)
            final_list.extend(shuffled_buf)
        return final_list

# =============================================================================
# E. PARSER: ĐÓNG GÓI CHẶT (FIXED LOGIC)
# =============================================================================
def parser_dong_goi(doc):
    sections = []
    curr_header = []; curr_items = []
    state = "OUT" 
    active_group = None 
    
    all_blocks = []
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P) or isinstance(child, CT_Tbl):
            all_blocks.append(child)

    for block in all_blocks:
        txt = get_block_text(block, doc).upper()
        if "PHẦN" in txt and ("TRẮC NGHIỆM" in txt or "TỰ LUẬN" in txt):
            if curr_header or curr_items:
                if active_group:
                    curr_items.append(active_group)
                    active_group = None
                sections.append({'header': curr_header, 'items': curr_items})
            curr_header = [block]; curr_items = []; state = "HEADER"
            continue
        
        if state == "HEADER": state = "BODY"
        if state == "OUT":
            if not sections: sections.append({'header': [], 'items': []})
            sections[0]['items'].append({'type': 'PARA', 'obj': block})
            continue

        if isinstance(block, CT_P):
            if not Paragraph(block, doc).text.strip(): continue

        m_type = identify_marker(block, doc)
        
        if m_type in ["START_NORMAL", "START_FIXED"]:
            if active_group: curr_items.append(active_group) 
            active_group = {
                'type': 'GROUP',
                'header': block,
                'group_header': [],
                'qs': [],
                'footer': None,
                'fixed': (m_type == "START_FIXED"),
                'has_seen_q': False
            }
            continue 
        
        if m_type == "END":
            if active_group:
                active_group['footer'] = block
                curr_items.append(active_group)
                active_group = None
            continue
            
        if active_group:
            is_q = False
            if isinstance(block, CT_Tbl):
                try:
                    c0_text = Table(block, doc).rows[0].cells[0].text.strip().upper()
                    if c0_text.startswith("CÂU") or c0_text.startswith("!") or c0_text.startswith("@") or c0_text.startswith("QUESTION"):
                        is_q = True
                except: pass
            
            if is_q:
                active_group['has_seen_q'] = True
                active_group['qs'].append(block)
            else:
                if not active_group['has_seen_q']:
                    active_group['group_header'].append(block)
                else:
                    active_group['qs'].append(block)
        else:
            if isinstance(block, CT_P):
                curr_items.append({'type': 'PARA', 'obj': block})
            else:
                curr_items.append({'type': 'SINGLE', 'obj': block})

    if active_group: curr_items.append(active_group)
    if curr_header or curr_items:
        sections.append({'header': curr_header, 'items': curr_items})
        
    return sections

# =============================================================================
# F. MAIN FLOW
# =============================================================================
def tron_de(doc_chuan_hoa, so_luong_de, config=None):
    if config is None: config = {}
    list_result = []
    
    start_id = config.get('ma_de_start', 101)
    kieu_ma = config.get('kieu_ma_de', 'SEQUENTIAL')
    generated_ids = set()
    
    for i in range(so_luong_de):
        if kieu_ma == 'SEQUENTIAL':
            exam_id = str(start_id + i)
        else:
            while True:
                rid = random.randint(100, 999) if kieu_ma == 'RANDOM_3' else random.randint(1000, 9999)
                if str(rid) not in generated_ids:
                    generated_ids.add(str(rid)); exam_id = str(rid); break
        
        doc = _clone_document_in_ram(doc_chuan_hoa)
        sections = parser_dong_goi(doc)
        
        for section in sections:
            if section['items']:
                items = section['items']
                items = layer1_dao_noi_dung(items, config, doc)
                items = layer2_tron_trong_nhom(items, config, doc)
                # [ĐÃ SỬA]: Cập nhật danh sách mới vào lại mảng gốc
                section['items'] = layer3_tron_cau_truc(items, config, doc)

        current_idx = 1
        for section in sections:
            for block in section['header']:
                txt = get_block_text(block, doc).upper()
                if "PHẦN" in txt and ("TRẮC NGHIỆM" in txt or "TỰ LUẬN" in txt):
                    current_idx = 1

            for item in section['items']:
                if item['type'] in ['SINGLE', 'PARA']:
                    block = item['obj']
                    if isinstance(block, CT_Tbl):
                        tbl = Table(block, doc)
                        try:
                            c0 = tbl.rows[0].cells[0].text.strip()
                            if c0.startswith("Câu") or c0.startswith("Question") or c0.startswith("!") or c0.startswith("@"):
                                prefix = ""
                                if "!" in c0: prefix += "!"
                                if "@" in c0: prefix += "@"
                                tbl.rows[0].cells[0].text = f"{prefix}Câu {current_idx}"
                                current_idx += 1
                        except: pass
                        
                elif item['type'] == 'GROUP':
                    start_idx = current_idx
                    for block in item['qs']:
                        if isinstance(block, CT_Tbl):
                            tbl = Table(block, doc)
                            try:
                                c0 = tbl.rows[0].cells[0].text.strip()
                                if c0.startswith("Câu") or c0.startswith("Question") or c0.startswith("!") or c0.startswith("@"):
                                    prefix = ""
                                    if "!" in c0: prefix += "!"
                                    if "@" in c0: prefix += "@"
                                    tbl.rows[0].cells[0].text = f"{prefix}Câu {current_idx}"
                                    current_idx += 1
                            except: pass
                    end_idx = current_idx - 1 if current_idx > start_idx else start_idx

                    for g_block in item['group_header']:
                        if isinstance(g_block, CT_Tbl):
                            g_tbl = Table(g_block, doc)
                            for row in g_tbl.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        if "<S>" in p.text or "<E>" in p.text:
                                            _replace_text_in_p(p, "<S>", str(start_idx))
                                            _replace_text_in_p(p, "<E>", str(end_idx))
                        elif isinstance(g_block, CT_P):
                            p = Paragraph(g_block, doc)
                            if "<S>" in p.text or "<E>" in p.text:
                                _replace_text_in_p(p, "<S>", str(start_idx))
                                _replace_text_in_p(p, "<E>", str(end_idx))

        final_stream = []
        for section in sections:
            final_stream.extend(section['header'])
            if section['items']:
                for item in section['items']:
                    if item['type'] in ['PARA', 'SINGLE']:
                        final_stream.append(item['obj'])
                    elif item['type'] == 'GROUP':
                        final_stream.append(item['header'])
                        final_stream.extend(item['group_header']) 
                        final_stream.extend(item['qs'])           
                        if item['footer']: final_stream.append(item['footer'])
        
        doc.element.body.clear_content()
        doc.element.body.append(create_empty_p())
        
        for block in final_stream:
            doc.element.body.append(block)
            if isinstance(block, CT_Tbl): doc.element.body.append(create_empty_p())

        list_result.append({'exam_id': exam_id, 'file_content': doc})

    return list_result