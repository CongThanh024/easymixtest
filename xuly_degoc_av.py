import re
import io 
from copy import deepcopy 
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

# ============================================================================
# 🔒 KHU VỰC 1: CÔNG CỤ LÕI & CẤU HÌNH
# ============================================================================
COLOR_BLUE = RGBColor(0, 0, 255) 
COLOR_RED = RGBColor(255, 0, 0)

# [BỘ LỌC CHUẨN MỰC] - TUYỆT ĐỐI CHỈ BẮT A. B. C. D. (KHÔNG BẮT a) b) c) d) NỮA)
REGEX_CAU = re.compile(r"^[\s\u00A0]*([@!]?[\s\u00A0]*(?:Câu|Question)[\s\u00A0]+(\d+))([:.])", re.IGNORECASE) 
REGEX_PA_ANCHOR = re.compile(r"(?<![a-zA-Z0-9])([A-D])\.") 
REGEX_KEY_LINE = re.compile(r"^[\s\u00A0]*(LỜI GIẢI|HƯỚNG DẪN GIẢI|SOL)[\s\u00A0]*[:\.]?[\s\u00A0]*", re.IGNORECASE) 
REGEX_TRASH = re.compile(r"(<Tự luận>|<TỰ LUẬN>|<key[^>]*>|TLN)", re.IGNORECASE)
REGEX_OLD_CHOICE = re.compile(r"^\s*(Chọn|Đáp án|Phương án|Chon)\s+([A-D])\s*[.:]?\s*$", re.IGNORECASE)

class XuLyDeChuanHoaAV:
    def __init__(self, mode='ENG'):
        self.doc = None
        self.mode = mode 
        self.stats = {"GROUPS": 0, "QUESTIONS": 0, "ERRORS": 0}
        self.error_log = []

    def _clone_paragraph(self, paragraph):
        p_element = deepcopy(paragraph._element)
        return Paragraph(p_element, paragraph._parent)

    def _setup_page_layout(self):
        try:
            for section in self.doc.sections:
                section.page_width, section.page_height = Cm(21.0), Cm(29.7)
                section.left_margin = section.right_margin = Cm(1.27)
                section.top_margin = section.bottom_margin = Cm(1.27)
        except: pass

    def _clear_footers(self):
        self.doc.settings.odd_and_even_pages_header_footer = False
        for section in self.doc.sections:
            section.different_first_page_header_footer = False
            for el in [section.footer, section.header, section.even_page_header, section.even_page_footer, section.first_page_header, section.first_page_footer]:
                if el: self._wipe(el._element)

    def _wipe(self, element):
        for child in list(element): element.remove(child)

    def _sanitize_element(self, element):
        ns_sect = qn("w:sectPr"); ns_tabs = qn("w:tabs"); ns_ind = qn("w:ind")
        for sect in element.findall(".//" + ns_sect): sect.getparent().remove(sect)
        for tabs in element.findall(".//" + ns_tabs): tabs.getparent().remove(tabs)
        for ind in element.findall(".//" + ns_ind): ind.getparent().remove(ind)
        if element.tag == ns_sect: return None
        return element

    # [BẢO VỆ XML]: Vẽ viền bảng bọc thép chống Corrupt file (ĐÃ KIỂM TRA CHUẨN)
    def _set_table_grid(self, table):
        try: table.style = 'Table Grid'
        except KeyError:
            tblPr = table._tbl.tblPr
            if tblPr is None: tblPr = OxmlElement('w:tblPr'); table._tbl.insert(0, tblPr)
            borders = OxmlElement('w:tblBorders')
            for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{b}')
                border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4'); border.set(qn('w:space'), '0'); border.set(qn('w:color'), 'auto')
                borders.append(border)
            
            insert_idx = len(tblPr)
            for i, child in enumerate(tblPr):
                if child.tag in [qn('w:shd'), qn('w:tblLayout'), qn('w:tblCellMar'), qn('w:tblLook'), qn('w:tblCaption')]:
                    insert_idx = i; break
            tblPr.insert(insert_idx, borders)

    def _set_col_width(self, table):
        table.autofit = False; table.allow_autofit = False
        tblPr = table._tbl.tblPr
        if tblPr is not None:
            for w in tblPr.xpath("w:tblW"): tblPr.remove(w)
            tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), '10466'); tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)
        tblGrid = table._tbl.tblGrid
        if tblGrid is None: tblGrid = OxmlElement('w:tblGrid'); table._tbl.insert(0, tblGrid)
        else: self._wipe(tblGrid)
        gridCol1 = OxmlElement('w:gridCol'); gridCol1.set(qn('w:w'), '1701') 
        gridCol2 = OxmlElement('w:gridCol'); gridCol2.set(qn('w:w'), '8765') 
        tblGrid.append(gridCol1); tblGrid.append(gridCol2)

    def _format_label_cell(self, cell, text, color=COLOR_BLUE):
        self._wipe(cell._element)
        tcPr = cell._tc.get_or_add_tcPr()
        for w in tcPr.xpath("w:tcW"): tcPr.remove(w)
        tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), '1701'); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
        p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text); run.bold = True; run.font.color.rgb = color
        run.font.name = "Times New Roman"; run.font.size = Pt(11)

    def _check_label_highlight(self, paragraph, label_char, match_start_index):
        text = paragraph.text
        label_index = text.find(label_char, match_start_index)
        if label_index == -1: return False
        current_pos = 0
        for run in paragraph.runs:
            run_len = len(run.text)
            if current_pos <= label_index < current_pos + run_len:
                if (run.font.color and run.font.color.rgb == RGBColor(255, 0, 0)) or run.font.underline: return True
                return False
            current_pos += run_len
        return False

    def _has_special(self, element):
        for node in element.iter():
            tag = node.tag.split('}')[-1]
            if tag in ['drawing', 'oMath', 'oMathPara', 'shape', 'object', 'pict', 'AlternateContent']: return True
        return False

    def _is_followed_by_special(self, p, end_idx):
        current_idx = 0
        for node in p._element.iterchildren():
            tag = node.tag.split('}')[-1]
            if current_idx == end_idx:
                if tag in ['drawing', 'object', 'pict', 'oMath', 'oMathPara']: return True
            if tag == 'r':
                for child in node.iter():
                    ctag = child.tag.split('}')[-1]
                    if current_idx == end_idx:
                         if ctag in ['drawing', 'object', 'pict', 'oMath']: return True
                    if ctag == 't' and child.text: current_idx += len(child.text)
                    elif ctag == 'tab': current_idx += 1
                    elif ctag in ['br', 'cr']: current_idx += 1
        return False

    def _is_content_paragraph(self, p):
        if p.text.strip(): return True
        return self._has_special(p._element)

    def _run_lstrip(self, p):
        for run in p.runs:
            if self._has_special(run._element): return
            text = run.text
            if not text.strip(): run.text = ""
            else: run.text = text.lstrip(" \t\u00A0"); return

    def _run_rstrip(self, p):
        for run in reversed(p.runs):
            if self._has_special(run._element): return
            text = run.text
            if not text.strip(): run.text = ""
            else: run.text = text.rstrip(" \t\u00A0"); return

    def _crop_paragraph(self, p, start_index, end_index):
        full_text = p.text
        if end_index is not None and end_index < len(full_text):
            curr_idx = 0
            for run in list(p.runs):
                run_len = len(run.text)
                xml_str = run._element.xml
                is_special = 'w:drawing' in xml_str or 'm:oMath' in xml_str or 'v:shape' in xml_str or 'w:object' in xml_str or 'w:pict' in xml_str or 'mc:AlternateContent' in xml_str
                if curr_idx >= end_index:
                    if is_special and curr_idx == end_index: pass
                    else: run._element.getparent().remove(run._element)
                elif curr_idx < end_index < curr_idx + run_len:
                    run.text = run.text[:end_index - curr_idx]
                curr_idx += run_len

        if start_index is not None and start_index > 0:
            curr_idx = 0
            for run in list(p.runs):
                run_len = len(run.text)
                run_end = curr_idx + run_len
                xml_str = run._element.xml
                is_special = 'w:drawing' in xml_str or 'm:oMath' in xml_str or 'v:shape' in xml_str or 'w:object' in xml_str or 'w:pict' in xml_str or 'mc:AlternateContent' in xml_str
                if run_end <= start_index:
                    if is_special and run_end == start_index: pass
                    else: run._element.getparent().remove(run._element)
                elif curr_idx < start_index < run_end:
                    run.text = run.text[start_index - curr_idx:]
                curr_idx += run_len

    def _remove_text_range(self, p, start_index, end_index):
        if start_index >= end_index: return
        current_pos = 0
        for run in p.runs:
            run_len = len(run.text)
            run_end = current_pos + run_len
            intersect_start = max(current_pos, start_index)
            intersect_end = min(run_end, end_index)
            if intersect_start < intersect_end:
                cut_start = intersect_start - current_pos
                cut_end = intersect_end - current_pos
                head = run.text[:cut_start]
                tail = run.text[cut_end:]
                run.text = head + tail
            current_pos += run_len

    def _trim_cell_safe(self, cell, regex_prefix=None):
        if not cell.paragraphs: return
        for p in cell.paragraphs:
            prev_text = None
            while True:
                match = REGEX_TRASH.search(p.text)
                if not match: break
                self._remove_text_range(p, match.start(), match.end())
                if p.text == prev_text: break 
                prev_text = p.text

        safe_count = 0
        while len(cell.paragraphs) > 1 and safe_count < 20:
            safe_count += 1
            if not self._is_content_paragraph(cell.paragraphs[0]): 
                try: cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
                except: break
            else: break
            
        safe_count = 0
        while len(cell.paragraphs) > 1 and safe_count < 20:
            safe_count += 1
            if not self._is_content_paragraph(cell.paragraphs[-1]): 
                try: cell.paragraphs[-1]._element.getparent().remove(cell.paragraphs[-1]._element)
                except: break
            else: break
            
        if not cell.paragraphs: 
            cell.add_paragraph(""); return
            
        if regex_prefix:
            p = cell.paragraphs[0]
            match = regex_prefix.match(p.text.strip())
            if match:
                trash = match.group(0)
                start_idx = p.text.find(trash)
                if start_idx != -1: self._crop_paragraph(p, start_idx + len(trash), None)
                    
        safe_count = 0
        while len(cell.paragraphs) > 0 and safe_count < 20:
            safe_count += 1
            if len(cell.paragraphs) > 1:
                if not self._is_content_paragraph(cell.paragraphs[0]): 
                    try: cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
                    except: break
                else: break
            else:
                if not self._is_content_paragraph(cell.paragraphs[0]):
                    for run in cell.paragraphs[0].runs: run.text = ""
                break
                
        if cell.paragraphs:
            self._run_lstrip(cell.paragraphs[0])
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            self._run_rstrip(cell.paragraphs[-1])

    def _clean_solution_cell(self, cell):
        if not cell.paragraphs: return
        for i in range(len(cell.paragraphs) - 1, -1, -1):
            text = cell.paragraphs[i].text.strip()
            if REGEX_OLD_CHOICE.match(text): 
                try: cell.paragraphs[i]._element.getparent().remove(cell.paragraphs[i]._element)
                except: pass
                continue
        self._trim_cell_safe(cell, REGEX_KEY_LINE)

# ============================================================================
# 🔄 KHU VỰC 2: BỘ NHẬN DẠNG & ĐÓNG HỘP
# ============================================================================
    def _khu_vuc_2_nhan_dang_du_lieu(self):
        all_blocks = []
        in_group = False
        current_group = None
        current_cau = None 
        absolute_q_idx = 0 

        def flush_cau():
            nonlocal current_cau
            if current_cau:
                if in_group and current_group is not None: current_group['questions'].append(current_cau)
                else: all_blocks.append({'type': 'SINGLE', 'data': current_cau})
                current_cau = None

        for child in self.doc._body._element.iterchildren():
            if child.tag.endswith('tbl'):
                if current_cau:
                    current_cau['raw_xml'].append(deepcopy(child))
                    class TableWrapper:
                        def __init__(self, element):
                            self._element = element; self.text = ""; self.runs = []
                    tbl_wrap = TableWrapper(deepcopy(child))
                    if current_cau['curr_part'] == 'solution': current_cau['solution'].append(tbl_wrap)
                    else:
                        if current_cau['last_marker'] is None: current_cau['stem'].append(tbl_wrap)
                        else: current_cau['options'][current_cau['last_marker']].append(tbl_wrap)
                elif in_group and current_group is not None and current_cau is None:
                    class TableWrapper:
                        def __init__(self, element):
                            self._element = element; self.text = ""; self.runs = []
                    current_group['context'].append(TableWrapper(deepcopy(child)))
                continue
            
            if not child.tag.endswith('p'): continue
            
            p = Paragraph(child, self.doc)
            text = p.text; clean_text = text.strip()
            if re.match(r"^[\-\s]*(HẾT|HẾT\.)[\-\s]*$", clean_text, re.IGNORECASE): continue

            clean_nospace = re.sub(r'\s+', '', clean_text)
            if "#*#" in clean_nospace or "#@#" in clean_nospace:
                flush_cau()
                if in_group:
                    self.error_log.append(("Hệ thống", "Lỗi cấu trúc nhóm", "Phát hiện thẻ mở nhóm mới nhưng chưa đóng nhóm trước đó. Tự động chốt nhóm cũ."))
                    all_blocks.append({'type': 'GROUP', 'data': current_group})
                in_group = True
                current_group = {'start_marker': p, 'context': [], 'questions': [], 'end_marker': None}
                continue
                
            if "#**#" in clean_nospace:
                flush_cau()
                if not in_group:
                    self.error_log.append(("Hệ thống", "Lỗi cấu trúc nhóm", "Phát hiện thẻ đóng nhóm nhưng không có thẻ mở nhóm trước đó. Hệ thống bỏ qua thẻ này."))
                    continue
                in_group = False
                current_group['end_marker'] = p
                all_blocks.append({'type': 'GROUP', 'data': current_group})
                current_group = None
                continue

            m_cau = REGEX_CAU.match(clean_text)
            is_new_q_line = False
            
            if m_cau:
                flush_cau()
                absolute_q_idx += 1 
                raw_label = m_cau.group(1).strip()
                
                current_cau = {
                    'label': raw_label, 'stem': [], 
                    'options': {'A': [], 'B': [], 'C': [], 'D': []},
                    'solution': [], 'key_value': "", 'curr_part': 'stem', 'last_marker': None,
                    'absolute_id': absolute_q_idx,       
                    'raw_xml': [deepcopy(p._element)],
                    'has_floating': False,
                    'marker_counts': {'A':0, 'B':0, 'C':0, 'D':0},
                    'q_type': None,
                    'tl_key_paras': []
                }
                
                if "<wp:anchor" in p._element.xml: current_cau['has_floating'] = True
                is_new_q_line = True

            if in_group and current_group is not None and current_cau is None:
                current_group['context'].append(p)
                continue

            if current_cau:
                if not is_new_q_line:
                    if "<wp:anchor" in p._element.xml: current_cau['has_floating'] = True
                    current_cau['raw_xml'].append(deepcopy(p._element)) 
                
                # [BẮT KEY TỰ LUẬN]: Trích xuất không rớt 1 cái MathType nào
                # [BỐC KEY CHÍNH XÁC KHÔNG MẤT MATHTYPE]
                text_lower = text.lower()
                if "<key=" in text_lower and ">" in text_lower:
                    if not current_cau.get('tl_key_paras'):
                        p_key = self._clone_paragraph(p)
                        start_tag = text_lower.find("<key=")
                        end_tag = text_lower.find(">", start_tag)
                        
                        if start_tag != -1 and end_tag != -1:
                            # Cắt chính xác từ sau dấu = đến trước dấu >
                            self._crop_paragraph(p_key, start_tag + 5, end_tag)
                            self._run_lstrip(p_key)
                            self._run_rstrip(p_key)
                            current_cau['tl_key_paras'] = [p_key]
                    current_cau['q_type'] = 'TL'

                is_solution_boundary = False
                if current_cau['curr_part'] == 'solution': is_solution_boundary = True
                else:
                    m_sol = REGEX_KEY_LINE.match(text)
                    if m_sol:
                        matched_str = m_sol.group(0)
                        remaining_text = text[len(matched_str):].strip()
                        if not remaining_text: is_solution_boundary = True
                        else:
                            for run in p.runs:
                                if run.text.strip():
                                    if run.bold or (run._element.rPr is not None and run._element.rPr.b is not None): is_solution_boundary = True
                                    break 
                                    
                if is_solution_boundary:
                    current_cau['curr_part'] = 'solution'
                    current_cau['solution'].append(self._clone_paragraph(p))
                    continue
                
                raw_matches = list(REGEX_PA_ANCHOR.finditer(text))
                raw_matches.sort(key=lambda x: x.start())

                valid_matches = []
                for m in raw_matches:
                    end_idx = m.end()
                    is_valid = False
                    if end_idx == len(text): is_valid = True
                    else:
                        next_char = text[end_idx]
                        if next_char in [' ', '\t', '\n', '\r', '\u00A0']: is_valid = True
                        else:
                            curr_pos = 0; is_valid_special = False
                            for run in p.runs:
                                run_len = len(run.text)
                                if curr_pos == end_idx:
                                    xml_str = run._element.xml
                                    if 'w:drawing' in xml_str or 'm:oMath' in xml_str or 'v:shape' in xml_str or 'w:object' in xml_str or 'w:pict' in xml_str or 'mc:AlternateContent' in xml_str:
                                        is_valid_special = True; break
                                    elif run_len > 0: break 
                                elif curr_pos > end_idx: break
                                curr_pos += run_len
                            if is_valid_special or self._is_followed_by_special(p, end_idx):
                                is_valid = True
                    
                    if is_valid:
                        valid_matches.append(m)
                        lbl = m.group(1).upper()
                        current_cau['marker_counts'][lbl] += 1
                        if self._check_label_highlight(p, m.group(1), m.start()):
                            current_cau['key_value'] += lbl # Ghép vào để check nếu người dùng tô đỏ nhiều đáp án

                if not valid_matches:
                    if is_new_q_line:
                        p_clone = self._clone_paragraph(p)
                        self._crop_paragraph(p_clone, m_cau.end(), None)
                        if self._is_content_paragraph(p_clone): current_cau['stem'].append(p_clone)
                    else:
                        if current_cau['last_marker'] is None: current_cau['stem'].append(self._clone_paragraph(p))
                        else: current_cau['options'][current_cau['last_marker']].append(self._clone_paragraph(p))
                    continue
                
                current_idx = m_cau.end() if is_new_q_line else 0 
                for m in valid_matches:
                    new_marker = m.group(1).upper()
                    start_marker, end_marker = m.start(), m.end()
                    if start_marker > current_idx:
                        segment = self._clone_paragraph(p) 
                        self._crop_paragraph(segment, current_idx, start_marker)
                        if self._is_content_paragraph(segment):
                            if current_cau['last_marker'] is None: current_cau['stem'].append(segment)
                            else: current_cau['options'][current_cau['last_marker']].append(segment)
                    current_cau['last_marker'] = new_marker
                    current_idx = end_marker 
                
                segment = self._clone_paragraph(p) 
                self._crop_paragraph(segment, current_idx, None)
                if self._is_content_paragraph(segment):
                    current_cau['options'][current_cau['last_marker']].append(segment)

        flush_cau()
        if in_group:
            all_blocks.append({'type': 'GROUP', 'data': current_group})
            self.error_log.append(("Hệ thống", "Lỗi cấu trúc nhóm", "Chạm đáy file nhưng chưa đóng nhóm. Tự động chốt nhóm."))
            
        return all_blocks

# ============================================================================
# 🔄 KHU VỰC 3: TRẠM KIỂM ĐỊNH LỖI ĐA TẦNG (CÁCH LY SẮT ĐÁ)
# ============================================================================
    def _check_dinh_chu(self, opt_box):
        if not opt_box: return False
        try:
            p = opt_box[0]
            if not isinstance(p, Paragraph): return False 
            for node in p._element.iterchildren():
                tag = node.tag.split('}')[-1]
                if tag in ['drawing', 'oMath', 'oMathPara', 'shape', 'object', 'pict', 'AlternateContent']: return False 
                if tag == 'r':
                    for child in node.iter():
                        ctag = child.tag.split('}')[-1]
                        if ctag in ['drawing', 'object', 'pict', 'AlternateContent', 'tab', 'br', 'cr']: return False 
                        if ctag == 't' and child.text:
                            first_char = child.text[0]
                            if first_char in [' ', '\t', '\n', '\r', '\u00A0']: return False 
                            else: return True 
        except: pass
        return False

    def _khu_vuc_3_kiem_dinh_loi(self, all_blocks):
        all_qs_for_check = []
        for block in all_blocks:
            if block['type'] == 'SINGLE': all_qs_for_check.append(block['data'])
            else: all_qs_for_check.extend(block['data']['questions'])

        for q in all_qs_for_check:
            abs_id = q.get('absolute_id', '?')
            q_name = q.get('label', f"Câu {abs_id}")
            error_list = [] 
            is_isolated = False

            has_tn_markers = all(q['marker_counts'][lbl] == 1 for lbl in ['A', 'B', 'C', 'D'])
            is_tl = (q.get('q_type') == 'TL')

            # TÔN TRỌNG TÁC GIẢ TỰ LUẬN: Đã tự luận là không check A B C D, khôi phục nguyên bản gốc!
            if is_tl:
                q['detected_type'] = 'PHAN_TL'
                q['stem'] = []; q['solution'] = []
                curr_part = 'stem'
                for i, xml_elem in enumerate(q['raw_xml']):
                    if xml_elem.tag.endswith('tbl'):
                        class DummyBlock:
                            def __init__(self, el): self._element = el
                        if curr_part == 'stem': q['stem'].append(DummyBlock(deepcopy(xml_elem)))
                        else: q['solution'].append(DummyBlock(deepcopy(xml_elem)))
                        continue
                    
                    p_temp = Paragraph(deepcopy(xml_elem), self.doc)
                    text = p_temp.text
                    
                    if curr_part == 'stem':
                        m_sol = REGEX_KEY_LINE.match(text)
                        is_sol = False
                        if m_sol:
                            if not text[len(m_sol.group(0)):].strip(): is_sol = True
                            else:
                                for run in p_temp.runs:
                                    if run.text.strip():
                                        if run.bold or (run._element.rPr is not None and run._element.rPr.b is not None): is_sol = True
                                        break
                        if is_sol:
                            curr_part = 'solution'
                            q['solution'].append(p_temp)
                            continue

                    if curr_part == 'stem':
                        if i == 0:
                            m_cau = REGEX_CAU.match(text)
                            if m_cau: self._crop_paragraph(p_temp, m_cau.end(), None)
                        
                        # [BẢO TOÀN NỘI DUNG CÂU 48]: Chỉ tỉa đúng vỏ <key=...>, giữ nguyên nội dung gốc
                        text_tmp_lower = p_temp.text.lower()
                        start_tag = text_tmp_lower.find("<key=")
                        while start_tag != -1:
                            end_tag = text_tmp_lower.find(">", start_tag)
                            if end_tag != -1:
                                self._remove_text_range(p_temp, start_tag, end_tag + 1)
                                text_tmp_lower = p_temp.text.lower()
                                start_tag = text_tmp_lower.find("<key=")
                            else:
                                break

                        if self._is_content_paragraph(p_temp): q['stem'].append(p_temp)
                    else:
                        if "<key=" in p_temp.text.lower() and ">" in p_temp.text.lower():
                            if not REGEX_KEY_LINE.match(p_temp.text): continue
                        q['solution'].append(p_temp)
                
                for lbl in ['A', 'B', 'C', 'D']: q['options'][lbl] = []
                
            else:
                # KIỂM ĐỊNH TRẮC NGHIỆM SẮT ĐÁ
                total_markers = sum(q['marker_counts'][lbl] for lbl in ['A', 'B', 'C', 'D'])
                
                # NẾU KHÔNG CÓ BẤT KỲ MỐC NÀO -> BÁO LỖI THIẾU KEY/SAI ĐỊNH DẠNG NGAY
                if total_markers == 0:
                    error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Thiếu KEY / Sai định dạng", "Câu hỏi không có thẻ <key=...> (Tự luận) và cũng không có các phương án A., B., C., D. (Trắc nghiệm). Hệ thống từ chối nhận diện."))
                    is_isolated = True
                else:
                    if not has_tn_markers:
                        missing = [lbl for lbl in ['A', 'B', 'C', 'D'] if q['marker_counts'][lbl] == 0]
                        duplicates = [lbl for lbl in ['A', 'B', 'C', 'D'] if q['marker_counts'][lbl] > 1]
                        if missing:
                            error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Thiếu phương án", f"Không tìm thấy phương án {', '.join(missing)}."))
                            is_isolated = True
                        if duplicates:
                            for lbl in duplicates:
                                error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Lỗi lặp phương án", f"Phương án {lbl} xuất hiện {q['marker_counts'][lbl]} lần."))
                            is_isolated = True

                if not is_isolated:
                    if not q['key_value']:
                        error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Thiếu KEY đáp án", "Không tìm thấy phương án nào được bôi đỏ/gạch chân."))
                        is_isolated = True
                    elif len(q['key_value']) > 1:
                        error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Nhiều KEY đáp án", f"Phát hiện nhiều phương án ({q['key_value']}) được bôi đỏ/gạch chân."))
                        is_isolated = True

                if not is_isolated:
                    for lbl in ['A', 'B', 'C', 'D']:
                        opt_box = q['options'][lbl]
                        if self._check_dinh_chu(opt_box):
                            error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Lỗi dính chữ phương án", f"Ký tự dính sát nhãn {lbl}."))
                            is_isolated = True
                        is_empty = True
                        for p_box in opt_box:
                            if hasattr(p_box, 'text'):
                                if self._is_content_paragraph(p_box): is_empty = False; break
                            else:
                                is_empty = False; break 
                        if is_empty:
                            error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Lỗi phương án trống", f"Nội dung của phương án {lbl} trống rỗng."))
                            is_isolated = True

                if not is_isolated:
                    q['detected_type'] = 'PHAN_TN'

            if q.get('has_floating'): 
                error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Lỗi chèn ảnh", "Câu này chứa hình ảnh ở chế độ trôi nổi. Vui lòng chọn Wrap Text -> 'In Line with Text'."))
                is_isolated = True
            
            for err in error_list: self.error_log.append(err)
            q['is_isolated'] = is_isolated

            # Trả lại nguyên bản cho câu lỗi
            if is_isolated and not is_tl:
                q['stem'] = [] 
                for xml_elem in q['raw_xml']:
                    if xml_elem.tag.endswith('tbl'):
                        class DummyBlock:
                            def __init__(self, el): self._element = el
                        q['stem'].append(DummyBlock(deepcopy(xml_elem)))
                        continue
                    
                    p_temp = Paragraph(xml_elem, self.doc)
                    if REGEX_KEY_LINE.match(p_temp.text.strip()): break 
                    q['stem'].append(Paragraph(deepcopy(xml_elem), self.doc))
                for lbl in ['A', 'B', 'C', 'D']: q['options'][lbl] = []

        buckets = {'VALID': [], 'ERROR': []}
        for block in all_blocks:
            if block['type'] == 'SINGLE':
                q = block['data']
                if q['is_isolated']: buckets['ERROR'].append(block)
                else: buckets['VALID'].append(block)
            elif block['type'] == 'GROUP':
                valid_qs = []; faulty_qs = []
                for q in block['data']['questions']:
                    if q['is_isolated']: faulty_qs.append(q)
                    else: valid_qs.append(q)
                
                if valid_qs:
                    block['data']['questions'] = valid_qs
                    buckets['VALID'].append(block)
                for fq in faulty_qs:
                    buckets['ERROR'].append({'type': 'SINGLE', 'data': fq})
                    
        return buckets

# ============================================================================
# 🔒 KHU VỰC 4: VẼ BẢNG CHUẨN HÓA (ĐÃ GỘP BẢNG TỰ LUẬN LÀM 1)
# ============================================================================
    def _khu_vuc_4_ve_bang_chuan_hoa(self, buckets):
        body = self.doc._body._element
        sectPr = None
        for child in list(body):
            if child.tag.endswith('sectPr'): sectPr = child
            else: body.remove(child)

        blocks_in_phan = buckets['VALID']
        if not blocks_in_phan: return 
        
        self.stats["GROUPS"] = sum(1 for b in blocks_in_phan if b['type'] == 'GROUP')
        local_count = 0
        for b in blocks_in_phan:
            if b['type'] == 'SINGLE': local_count += 1
            else: local_count += len(b['data']['questions'])
        self.stats["QUESTIONS"] = local_count

        label_sol = "SOL" if self.mode == 'ENG' else "LỜI GIẢI"
        
        for block in blocks_in_phan:
            if block['type'] == 'GROUP':
                self.doc.add_paragraph(block['data']['start_marker'].text)
                
                # Lọc bỏ các dòng trắng, chỉ vẽ ô P nếu thực sự có chữ hoặc hình
                valid_context = [p for p in block['data'].get('context', []) if self._is_content_paragraph(p)]
                
                if valid_context:
                    tbl_p = self.doc.add_table(rows=0, cols=2)
                    self._set_table_grid(tbl_p); self._set_col_width(tbl_p)
                    row = tbl_p.add_row()
                    self._format_label_cell(row.cells[0], "P", COLOR_BLUE)
                    cell_content = row.cells[1]; self._wipe(cell_content._element)
                    for c_block in valid_context: 
                        cell_content._element.append(deepcopy(c_block._element))
                    self._trim_cell_safe(cell_content)
                    
                    # Ép khít dòng trắng giữa P và Câu hỏi thành 1 dòng siêu mỏng
                    spacer = self.doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(0)
                    spacer.paragraph_format.space_after = Pt(0)
                    spacer.paragraph_format.line_spacing = 1.0
            
            questions = [block['data']] if block['type'] == 'SINGLE' else block['data']['questions']
            for q in questions:
                current_table = self.doc.add_table(rows=0, cols=2)
                self._set_table_grid(current_table); self._set_col_width(current_table)
                
                if q['detected_type'] == 'PHAN_TN':
                    r1 = current_table.add_row(); self._format_label_cell(r1.cells[0], q['label'])
                    for p in q['stem']: 
                        new_p = deepcopy(p._element)
                        if self._sanitize_element(new_p) is not None: r1.cells[1]._element.append(new_p) 
                    self._trim_cell_safe(r1.cells[1], REGEX_CAU)
                    
                    for lbl in ['A', 'B', 'C', 'D']:
                        r = current_table.add_row(); self._format_label_cell(r.cells[0], lbl)
                        for p in q['options'][lbl]: 
                            new_p = deepcopy(p._element)
                            if self._sanitize_element(new_p) is not None: r.cells[1]._element.append(new_p) 
                        self._trim_cell_safe(r.cells[1])
                    rk = current_table.add_row(); k_txt = f"KEY: {q['key_value']}"; self._format_label_cell(rk.cells[0], k_txt, COLOR_RED)
                    for p in q['solution']: 
                        new_p = deepcopy(p._element)
                        if self._sanitize_element(new_p) is not None: rk.cells[1]._element.append(new_p) 
                    self._clean_solution_cell(rk.cells[1])
                    
                elif q['detected_type'] == 'PHAN_TL':
                    # [ĐẠI PHẪU BẢNG TL]: 2 dòng x 2 cột. Gộp KEY và LỜI GIẢI
                    r1 = current_table.add_row(); self._format_label_cell(r1.cells[0], q['label'])
                    for p in q['stem']: 
                        new_p = deepcopy(p._element)
                        if self._sanitize_element(new_p) is not None: r1.cells[1]._element.append(new_p) 
                    self._trim_cell_safe(r1.cells[1], REGEX_CAU)
                    
                    r2 = current_table.add_row()
                    c_key = r2.cells[0]
                    self._wipe(c_key._element)
                    tcPr = c_key._tc.get_or_add_tcPr()
                    tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), '1701'); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
                    
                    p_key = c_key.add_paragraph()
                    p_key.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run_k = p_key.add_run("KEY: ")
                    # Sửa lỗi format_run ở đây
                    run_k.bold = True
                    run_k.font.color.rgb = COLOR_RED
                    run_k.font.name = "Times New Roman"
                    
                    if q.get('tl_key_paras'):
                        for kp in q['tl_key_paras']:
                            for child in kp._element.iterchildren():
                                tag = child.tag.split('}')[-1]
                                if tag in ['r', 'oMath', 'oMathPara', 'hyperlink', 'drawing', 'object', 'pict', 'AlternateContent']:
                                    cloned = deepcopy(child)
                                    p_key._element.append(cloned)
                                    if tag == 'r':
                                        from docx.text.run import Run
                                        r_obj = Run(cloned, p_key)
                                        # Sửa lỗi format_run ở đây
                                        r_obj.bold = True
                                        r_obj.font.color.rgb = COLOR_RED
                    
                    self._trim_cell_safe(c_key)
                    
                    c_sol = r2.cells[1]
                    for p in q['solution']: 
                        new_p = deepcopy(p._element)
                        if self._sanitize_element(new_p) is not None: c_sol._element.append(new_p) 
                    self._clean_solution_cell(c_sol)
                    
                spacer = self.doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(1)
                
            if block['type'] == 'GROUP' and block['data']['end_marker']:
                self.doc.add_paragraph(block['data']['end_marker'].text)

        if sectPr: self.doc._body._element.append(sectPr)

    def _khu_vuc_4_ve_bang_loi_va_cach_ly(self, buckets):
        unique_err_qs = set([err[0] for err in self.error_log])
        self.stats["LOI"] = len(unique_err_qs) 
        
        if self.error_log:
            try:
                self.doc.add_page_break()
                p_title = self.doc.add_paragraph(); p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_title.add_run("BẢNG TỔNG HỢP LỖI CẤU TRÚC ĐỀ GỐC"); run.bold = True; run.font.color.rgb = COLOR_RED; run.font.size = Pt(14)
                
                p_note = self.doc.add_paragraph()
                run_note = p_note.add_run("Chú ý: Có thể số thứ tự câu không khớp thông báo do đề gốc chưa được đánh đúng thứ tự. Vui lòng dựa vào nội dung câu để đối chiếu.")
                run_note.italic = True; run_note.font.color.rgb = COLOR_RED
                
                err_table = self.doc.add_table(rows=1, cols=3); self._set_table_grid(err_table)
                hdr_cells = err_table.rows[0].cells
                hdr_cells[0].paragraphs[0].add_run('Vị trí phát hiện').bold = True
                hdr_cells[1].paragraphs[0].add_run('Phân loại lỗi').bold = True
                hdr_cells[2].paragraphs[0].add_run('Hướng dẫn khắc phục').bold = True
                            
                for error in self.error_log:
                    row_cells = err_table.add_row().cells
                    row_cells[0].paragraphs[0].add_run(error[0]).font.color.rgb = COLOR_RED
                    row_cells[1].paragraphs[0].add_run(error[1]).font.color.rgb = COLOR_RED
                    row_cells[2].paragraphs[0].add_run(error[2]).font.color.rgb = COLOR_RED
            except: pass 

        if buckets['ERROR']:
            try:
                self.doc.add_paragraph() 
                p_loi_title = self.doc.add_paragraph(); p_loi_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_loi = p_loi_title.add_run("DANH SÁCH CÁC CÂU LỖI BỊ CÁCH LY"); run_loi.bold = True; run_loi.font.color.rgb = COLOR_RED; run_loi.font.size = Pt(14)
                
                p_note_loi = self.doc.add_paragraph()
                run_note_loi = p_note_loi.add_run("Các câu dưới đây bị thiếu dữ liệu trọng yếu hoặc sai định dạng. Nội dung gốc được giữ nguyên bên dưới để dễ dàng đối chiếu:")
                run_note_loi.italic = True; run_note_loi.font.color.rgb = COLOR_RED
                
                for block in buckets['ERROR']:
                    q = block['data']; p_label = self.doc.add_paragraph()
                    q_label_name = q.get('label', "Câu lỗi"); abs_id_val = q.get('absolute_id', '?')
                    r_label = p_label.add_run(f"--- Nội dung gốc của Câu thứ {abs_id_val} ({q_label_name}) ---"); r_label.bold = True; r_label.font.color.rgb = COLOR_RED
                    
                    # 1. Xác định các phương án bị gõ trùng (Đếm trong sổ > 1)
                    duplicates = [lbl for lbl in ['A', 'B', 'C', 'D'] if q.get('marker_counts', {}).get(lbl, 0) > 1]
                    
                    for xml_elem in q['raw_xml']: 
                        self.doc._body._element.append(deepcopy(xml_elem))
                        
                        # 2. Quét tia laser tô màu xanh chữ cái bị lặp
                        if duplicates:
                            last_elem = self.doc._body._element[-1]
                            paras_to_check = []
                            if last_elem.tag.endswith('p'):
                                paras_to_check.append(Paragraph(last_elem, self.doc))
                            elif last_elem.tag.endswith('tbl'):
                                from docx.table import Table
                                paras_to_check.extend([p for row in Table(last_elem, self.doc).rows for cell in row.cells for p in cell.paragraphs])
                                
                            for new_p in paras_to_check:
                                text = new_p.text
                                for m in REGEX_PA_ANCHOR.finditer(text):
                                    lbl = m.group(1).upper()
                                    if lbl in duplicates:
                                        start_idx = m.start()
                                        end_idx = m.end()
                                        curr_pos = 0
                                        for run in new_p.runs:
                                            run_len = len(run.text)
                                            # Nếu đoạn văn bản nằm khớp tọa độ của chữ bị lặp (Ví dụ: "B.")
                                            if max(curr_pos, start_idx) < min(curr_pos + run_len, end_idx):
                                                run.font.color.rgb = COLOR_BLUE
                                                run.font.bold = True
                                            curr_pos += run_len
                                            
                    spacer = self.doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(1)
            except: pass

# ============================================================================
# 🔄 KHU VỰC 5: TRẠM ĐIỀU KHIỂN & TÁCH LUỒNG (XUẤT FILE)
# ============================================================================
    def xu_ly(self, input_source, output_path=""):
        self.doc = Document(input_source)
        self._setup_page_layout()
        self._clear_footers()

        all_blocks = self._khu_vuc_2_nhan_dang_du_lieu()
        buckets = self._khu_vuc_3_kiem_dinh_loi(all_blocks)
        self._khu_vuc_4_ve_bang_chuan_hoa(buckets)

        stream_sach = io.BytesIO()
        self.doc.save(stream_sach)
        stream_sach.seek(0)

        self._khu_vuc_4_ve_bang_loi_va_cach_ly(buckets)
        stream_vat_ly = io.BytesIO()
        self.doc.save(stream_vat_ly)
        stream_vat_ly.seek(0)
        
        # Đồng bộ từ khóa "LOI" để hiển thị đúng số liệu trên Web
        unique_err_qs = set([err[0] for err in self.error_log])
        self.stats["LOI"] = len(unique_err_qs)

        # Truyền ĐÚNG stream_sach thay vì mảng rỗng []
        return stream_sach, stream_vat_ly, self.stats, self.stats["LOI"]