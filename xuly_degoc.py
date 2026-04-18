import re
import io 
from copy import deepcopy 
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

# ============================================================================
# 🔒 KHU VỰC 1: CÔNG CỤ LÕI & CẤU HÌNH (NIÊM PHONG - CẤM XÂM PHẠM LÀM ẢNH HƯỞNG)
# ============================================================================
COLOR_BLUE = RGBColor(0, 0, 255) 
COLOR_RED = RGBColor(255, 0, 0)

STD_TEMPLATE = {
    "PHAN_I": {"bold": "PHẦN I. Câu trắc nghiệm nhiều phương án lựa chọn.", "normal": " Thí sinh trả lời từ câu {s} đến câu {e}. Mỗi câu hỏi thí sinh chỉ chọn một phương án."},
    "PHAN_II": {"bold": "PHẦN II. Câu trắc nghiệm đúng sai.", "normal": " Thí sinh trả lời từ câu {s} đến câu {e}. Trong mỗi ý a), b), c), d) ở mỗi câu, thí sinh chọn đúng hoặc sai."},
    "PHAN_III": {"bold": "PHẦN III. Câu trắc nghiệm trả lời ngắn.", "normal": " Thí sinh trả lời từ câu {s} đến câu {e}."},
    "PHAN_IV": {"bold": "PHẦN IV. Tự luận.", "normal": " Thí sinh trình bày lời giải chi tiết vào phiếu làm bài từ câu {s} đến câu {e}."}
}

# [BỘ LỌC CHUẨN MỰC]
REGEX_CAU = re.compile(r"^[\s\u00A0]*([@!]?[\s\u00A0]*Câu[\s\u00A0]+(\d+))([:.])", re.IGNORECASE) # Bắt buộc có dấu . hoặc :
REGEX_PA_ANCHOR = re.compile(r"(?<!\w)([A-D])\.") 
REGEX_PA_DS = re.compile(r"(?<![\w\(])([a-d])\)") # Chặn đứng (a)
REGEX_KEY_LINE = re.compile(r"^[\s\u00A0]*(LỜI GIẢI|HƯỚNG DẪN GIẢI)[\s\u00A0]*[:\.]?[\s\u00A0]*", re.IGNORECASE) # Chỉ 2 từ khóa
REGEX_TRASH = re.compile(r"(<Tự luận>|<TỰ LUẬN>|<key[^>]*>)", re.IGNORECASE)
REGEX_OLD_CHOICE = re.compile(r"^\s*(Chọn|Đáp án|Phương án|Chon)\s+([A-D])\s*[.:]?\s*$", re.IGNORECASE)
REGEX_OLD_HEADER = re.compile(r"^PHẦN\s+[0-9IV]+", re.IGNORECASE)
REGEX_DAP_SO = re.compile(r"^\s*Đáp\s*số\s*[:]?\s*(.*?)\s*\.?\s*$", re.IGNORECASE)

class XuLyDeChuanHoa:
    def __init__(self):
        self.doc = None
        self.stats = {"PHAN_I": 0, "PHAN_II": 0, "PHAN_III": 0, "PHAN_IV": 0, "LOI": 0}
        self.error_log = []

    def _clone_paragraph(self, paragraph):
        p_element = deepcopy(paragraph._element)
        return Paragraph(p_element, paragraph._parent)

    def _setup_page_layout(self):
        try:
            for section in self.doc.sections:
                section.page_width, section.page_height = Cm(21.0), Cm(29.7)
                section.left_margin = section.right_margin = Cm(1.27)
                section.top_margin = section.bottom_margin = Cm(1.27)
        except: pass

    def _clear_footers(self):
        self.doc.settings.odd_and_even_pages_header_footer = False
        for section in self.doc.sections:
            section.different_first_page_header_footer = False
            for el in [section.footer, section.header, section.even_page_header, section.even_page_footer, section.first_page_header, section.first_page_footer]:
                if el: self._wipe(el._element)

    def _wipe(self, element):
        for child in list(element): element.remove(child)

    def _sanitize_element(self, element):
        ns_sect = qn("w:sectPr"); ns_tabs = qn("w:tabs"); ns_ind = qn("w:ind")
        for sect in element.findall(".//" + ns_sect): sect.getparent().remove(sect)
        for tabs in element.findall(".//" + ns_tabs): tabs.getparent().remove(tabs)
        for ind in element.findall(".//" + ns_ind): ind.getparent().remove(ind)
        if element.tag == ns_sect: return None
        return element

    # [BẢO VỆ]: Code fallback vẽ viền bảng khi file Word bị lỗi định dạng
    def _set_table_grid(self, table):
        try: table.style = 'Table Grid'
        except KeyError:
            tblPr = table._tbl.tblPr
            if tblPr is None: tblPr = OxmlElement('w:tblPr'); table._tbl.insert(0, tblPr)
            borders = OxmlElement('w:tblBorders')
            for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{b}')
                border.set(qn('w:val'), 'single'); border.set(qn('w:sz'), '4'); border.set(qn('w:space'), '0'); border.set(qn('w:color'), 'auto')
                borders.append(border)
            tblPr.append(borders)

    def _set_col_width(self, table):
        table.autofit = False; table.allow_autofit = False
        tblPr = table._tbl.tblPr
        if tblPr is not None:
            for w in tblPr.xpath("w:tblW"): tblPr.remove(w)
            tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), '10466'); tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)
        tblGrid = table._tbl.tblGrid
        if tblGrid is None: tblGrid = OxmlElement('w:tblGrid'); table._tbl.insert(0, tblGrid)
        else: self._wipe(tblGrid)
        gridCol1 = OxmlElement('w:gridCol'); gridCol1.set(qn('w:w'), '1701') 
        gridCol2 = OxmlElement('w:gridCol'); gridCol2.set(qn('w:w'), '8765') 
        tblGrid.append(gridCol1); tblGrid.append(gridCol2)

    def _format_label_cell(self, cell, text, color=COLOR_BLUE):
        self._wipe(cell._element)
        tcPr = cell._tc.get_or_add_tcPr()
        for w in tcPr.xpath("w:tcW"): tcPr.remove(w)
        tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'), '1701'); tcW.set(qn('w:type'), 'dxa'); tcPr.append(tcW)
        p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text); run.bold = True; run.font.color.rgb = color
        run.font.name = "Times New Roman"; run.font.size = Pt(11)

    def _check_label_highlight(self, paragraph, label_char, match_start_index):
        text = paragraph.text
        label_index = text.find(label_char, match_start_index)
        if label_index == -1: return False
        current_pos = 0
        for run in paragraph.runs:
            run_len = len(run.text)
            if current_pos <= label_index < current_pos + run_len:
                if (run.font.color and run.font.color.rgb == RGBColor(255, 0, 0)) or run.font.underline: return True
                return False
            current_pos += run_len
        return False

    def _check_is_highlighted_general(self, paras):
        for p in paras:
            for run in p.runs:
                if (run.font.color and run.font.color.rgb == RGBColor(255, 0, 0)) or run.font.underline: return True
        return False

    def _generate_tf_key(self, cau_obj):
        labels = ['a', 'b', 'c', 'd']; key_chars = []
        for lbl in labels:
            paras = cau_obj['options'].get(lbl, [])
            if cau_obj['tf_keys'].get(lbl, False) or self._check_is_highlighted_general(paras): key_chars.append("T")
            else: key_chars.append("F")
        return "".join(key_chars)

    def _has_special(self, element):
        for node in element.iter():
            tag = node.tag.split('}')[-1]
            if tag in ['drawing', 'oMath', 'oMathPara', 'shape', 'object', 'pict', 'AlternateContent']: return True
        return False

    def _is_followed_by_special(self, p, end_idx):
        current_idx = 0
        for node in p._element.iterchildren():
            tag = node.tag.split('}')[-1]
            if current_idx == end_idx:
                if tag in ['drawing', 'object', 'pict', 'oMath', 'oMathPara']: return True
            if tag == 'r':
                for child in node.iter():
                    ctag = child.tag.split('}')[-1]
                    if current_idx == end_idx:
                         if ctag in ['drawing', 'object', 'pict', 'oMath']: return True
                    if ctag == 't' and child.text: current_idx += len(child.text)
                    elif ctag == 'tab': current_idx += 1
                    elif ctag in ['br', 'cr']: current_idx += 1
        return False

    def _is_content_paragraph(self, p):
        if p.text.strip(): return True
        return self._has_special(p._element)

    def _run_lstrip(self, p):
        for run in p.runs:
            if self._has_special(run._element): return
            text = run.text
            if not text.strip(): run.text = ""
            else: run.text = text.lstrip(" \t\u00A0"); return

    def _run_rstrip(self, p):
        for run in reversed(p.runs):
            if self._has_special(run._element): return
            text = run.text
            if not text.strip(): run.text = ""
            else: run.text = text.rstrip(" \t\u00A0"); return

    def _crop_paragraph(self, p, start_index, end_index):
        full_text = p.text
        if end_index is not None and end_index < len(full_text):
            curr_idx = 0
            for run in list(p.runs):
                run_len = len(run.text)
                xml_str = run._element.xml
                is_special = 'w:drawing' in xml_str or 'm:oMath' in xml_str or 'v:shape' in xml_str or 'w:object' in xml_str or 'w:pict' in xml_str or 'mc:AlternateContent' in xml_str
                if curr_idx >= end_index:
                    if is_special and curr_idx == end_index: pass
                    else: run._element.getparent().remove(run._element)
                elif curr_idx < end_index < curr_idx + run_len:
                    run.text = run.text[:end_index - curr_idx]
                curr_idx += run_len

        if start_index is not None and start_index > 0:
            curr_idx = 0
            for run in list(p.runs):
                run_len = len(run.text)
                run_end = curr_idx + run_len
                xml_str = run._element.xml
                is_special = 'w:drawing' in xml_str or 'm:oMath' in xml_str or 'v:shape' in xml_str or 'w:object' in xml_str or 'w:pict' in xml_str or 'mc:AlternateContent' in xml_str
                if run_end <= start_index:
                    if is_special and run_end == start_index: pass
                    else: run._element.getparent().remove(run._element)
                elif curr_idx < start_index < run_end:
                    run.text = run.text[start_index - curr_idx:]
                curr_idx += run_len

    def _remove_text_range(self, p, start_index, end_index):
        if start_index >= end_index: return
        current_pos = 0
        for run in p.runs:
            run_len = len(run.text)
            run_end = current_pos + run_len
            intersect_start = max(current_pos, start_index)
            intersect_end = min(run_end, end_index)
            if intersect_start < intersect_end:
                cut_start = intersect_start - current_pos
                cut_end = intersect_end - current_pos
                head = run.text[:cut_start]
                tail = run.text[cut_end:]
                run.text = head + tail
            current_pos += run_len

    def _trim_cell_safe(self, cell, regex_prefix=None):
        if not cell.paragraphs: return
        for p in cell.paragraphs:
            prev_text = None
            while True:
                match = REGEX_TRASH.search(p.text)
                if not match: break
                self._remove_text_range(p, match.start(), match.end())
                if p.text == prev_text: break 
                prev_text = p.text

        safe_count = 0
        while len(cell.paragraphs) > 1 and safe_count < 20:
            safe_count += 1
            if not self._is_content_paragraph(cell.paragraphs[0]): 
                try: cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
                except: break
            else: break
            
        safe_count = 0
        while len(cell.paragraphs) > 1 and safe_count < 20:
            safe_count += 1
            if not self._is_content_paragraph(cell.paragraphs[-1]): 
                try: cell.paragraphs[-1]._element.getparent().remove(cell.paragraphs[-1]._element)
                except: break
            else: break
            
        if not cell.paragraphs: 
            cell.add_paragraph(""); return
            
        if regex_prefix:
            p = cell.paragraphs[0]
            match = regex_prefix.match(p.text.strip())
            if match:
                trash = match.group(0)
                start_idx = p.text.find(trash)
                if start_idx != -1: self._crop_paragraph(p, start_idx + len(trash), None)
                    
        safe_count = 0
        while len(cell.paragraphs) > 0 and safe_count < 20:
            safe_count += 1
            if len(cell.paragraphs) > 1:
                if not self._is_content_paragraph(cell.paragraphs[0]): 
                    try: cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
                    except: break
                else: break
            else:
                if not self._is_content_paragraph(cell.paragraphs[0]):
                    for run in cell.paragraphs[0].runs: run.text = ""
                break
                
        if cell.paragraphs:
            self._run_lstrip(cell.paragraphs[0])
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            self._run_rstrip(cell.paragraphs[-1])

    def _clean_solution_cell(self, cell):
        if not cell.paragraphs: return
        for i in range(len(cell.paragraphs) - 1, -1, -1):
            text = cell.paragraphs[i].text.strip()
            if REGEX_OLD_CHOICE.match(text): 
                try: cell.paragraphs[i]._element.getparent().remove(cell.paragraphs[i]._element)
                except: pass
                continue
            m_ds = REGEX_DAP_SO.match(text)
            if m_ds:
                val = m_ds.group(1).strip().upper()
                val = val.replace(" ", "") 
                is_valid_rubbish = False
                if len(val) == 1 and val in ['A', 'B', 'C', 'D']: is_valid_rubbish = True
                elif len(val) == 4 and all(c in ['T', 'F', 'Đ', 'S'] for c in val): is_valid_rubbish = True
                elif re.match(r"^-?[\d,\.]{1,8}$", val): is_valid_rubbish = True
                if is_valid_rubbish:
                    try: cell.paragraphs[i]._element.getparent().remove(cell.paragraphs[i]._element)
                    except: pass
        self._trim_cell_safe(cell, REGEX_KEY_LINE)
    def _tien_xu_ly_shift_enter(self):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        import re

        # LUẬT ÉP CỨNG KHÔNG KHOAN NHƯỢNG: 
        # Bắt buộc in hoa A., B., C., D. hoặc in thường a), b), c), d)
        regex_marker = re.compile(r"^\s*(A\.|B\.|C\.|D\.|a\)|b\)|c\)|d\))")

        def get_text_after(node, length=5):
            # Hàm nội bộ dò tìm chữ cái phía sau Shift+Enter mà không phá XML
            text = ""
            current = node.getnext()
            while current is not None and len(text) < length:
                for t in current.xpath('.//w:t'):
                    if t.text: text += t.text
                if len(text) >= length: break
                current = current.getnext()

            parent = node.getparent()
            while parent is not None and len(text) < length:
                sibling = parent.getnext()
                while sibling is not None and len(text) < length:
                    for t in sibling.xpath('.//w:t'):
                        if t.text: text += t.text
                    if len(text) >= length: break
                    sibling = sibling.getnext()
                parent = parent.getparent()
                if parent is None or parent.tag.endswith('p'): break 
            return text

        def check_previous_is_special(br_node):
            # Hàm nội bộ lùi lại 1 bước xem có đụng MathType/Hình ảnh không
            current = br_node.getprevious()
            while current is not None:
                t_tags = current.xpath('.//w:t')
                if any(t.text and t.text.strip() for t in t_tags): return False
                if current.xpath('.//w:drawing') or current.xpath('.//w:pict') or current.xpath('.//w:object'): return True
                current = current.getprevious()

            parent_r = br_node.getparent()
            if parent_r is not None and parent_r.tag.endswith('r'):
                prev_r = parent_r.getprevious()
                while prev_r is not None:
                    tag = prev_r.tag.split('}')[-1]
                    if tag in ['oMath', 'oMathPara']: return True
                    t_tags = prev_r.xpath('.//w:t')
                    if any(t.text and t.text.strip() for t in t_tags): return False
                    if prev_r.xpath('.//w:drawing') or prev_r.xpath('.//w:pict') or prev_r.xpath('.//w:object'): return True
                    prev_r = prev_r.getprevious()
            return False

        body = self.doc._body._element
        i = 0
        while i < len(body):
            p_el = body[i]
            if not p_el.tag.endswith('p'):
                i += 1; continue

            # Tìm tất cả thẻ <w:br> (Shift+Enter) trong đoạn
            br_list = p_el.xpath('.//w:br')
            if not br_list:
                i += 1; continue

            split_happened = False
            for br in br_list:
                next_text = get_text_after(br, 5)
                # KHI PHÁT HIỆN TARGET THỎA ĐIỀU KIỆN ÉP CỨNG
                if regex_marker.match(next_text):
                    is_special = check_previous_is_special(br)
                    parent_r = br.getparent()
                    if parent_r is None or not parent_r.tag.endswith('r'): continue

                    # 1. BƠM VÙNG ĐỆM AN TOÀN (Dấu Chấm) nếu đằng trước là MathType
                    if is_special:
                        new_r = OxmlElement('w:r'); new_t = OxmlElement('w:t'); new_t.text = "."
                        new_r.append(new_t)
                        parent_r.addprevious(new_r)

                    # 2. PHẪU THUẬT: Tách Run và Dời XML
                    r_after = OxmlElement('w:r')
                    if parent_r.find(qn('w:rPr')) is not None: r_after.append(deepcopy(parent_r.find(qn('w:rPr'))))
                    
                    curr_node = br.getnext()
                    while curr_node is not None:
                        next_node = curr_node.getnext(); r_after.append(curr_node); curr_node = next_node
                        
                    # TIÊU DIỆT Shift+Enter
                    br.getparent().remove(br)

                    # 3. CHUYỂN SANG ĐOẠN ENTER MỚI
                    new_p = OxmlElement('w:p')
                    if p_el.find(qn('w:pPr')) is not None: new_p.append(deepcopy(p_el.find(qn('w:pPr'))))
                    if len(r_after) > (1 if r_after.find(qn('w:rPr')) is not None else 0): new_p.append(r_after)
                        
                    curr_sibling = parent_r.getnext()
                    while curr_sibling is not None:
                        next_sibling = curr_sibling.getnext(); new_p.append(curr_sibling); curr_sibling = next_sibling

                    # Lắp ráp đoạn mới vào file Word
                    p_el.addnext(new_p)
                    split_happened = True
                    break # Reset vòng lặp để soi tiếp đoạn mới

            if not split_happened: i += 1

# ============================================================================
# 🔄 KHU VỰC 2: BỘ NHẬN DẠNG & ĐÓNG HỘP (ĐÃ CẬP NHẬT ĐẾM MỐC TRƯỚC KHI CẮT)
# ============================================================================
    def _khu_vuc_2_nhan_dang_du_lieu(self):
        all_blocks = []
        in_group = False
        current_group = None
        current_cau = None 
        absolute_q_idx = 0 

        def flush_cau():
            nonlocal current_cau
            if current_cau:
                if in_group and current_group is not None: current_group['questions'].append(current_cau)
                else: all_blocks.append({'type': 'SINGLE', 'data': current_cau})
                current_cau = None

        for child in self.doc._body._element.iterchildren():
            if child.tag.endswith('tbl'):
                if current_cau:
                    current_cau['raw_xml'].append(deepcopy(child))
                    current_cau['has_table'] = True
                    
                    class TableWrapper:
                        def __init__(self, element):
                            self._element = element
                            self.text = "" 
                            self.runs = []
                    
                    tbl_wrap = TableWrapper(deepcopy(child))
                    
                    if current_cau['curr_part'] == 'solution':
                        current_cau['solution'].append(tbl_wrap)
                    else:
                        if current_cau['last_marker'] is None:
                            current_cau['stem'].append(tbl_wrap)
                        else:
                            current_cau['options'][current_cau['last_marker']].append(tbl_wrap)
                continue
            
            if not child.tag.endswith('p'): continue
            
            p = Paragraph(child, self.doc)
            text = p.text
            clean_text = text.strip()
            
            if re.match(r"^[\-\s]*(HẾT|HẾT\.)[\-\s]*$", clean_text, re.IGNORECASE): continue
            if REGEX_OLD_HEADER.match(clean_text): flush_cau(); continue

            clean_nospace = re.sub(r'\s+', '', clean_text)
            if "#*#" in clean_nospace or "#@#" in clean_nospace:
                flush_cau()
                if in_group:
                    self.error_log.append(("Hệ thống", "Lỗi cấu trúc nhóm", "Phát hiện thẻ mở nhóm mới nhưng chưa đóng nhóm trước đó. Tự động chốt nhóm cũ."))
                    all_blocks.append({'type': 'GROUP', 'data': current_group})
                in_group = True
                current_group = {'start_marker': p, 'header_content': [], 'questions': [], 'end_marker': None}
                continue
                
            if "#**#" in clean_nospace:
                flush_cau()
                if not in_group:
                    self.error_log.append(("Hệ thống", "Lỗi cấu trúc nhóm", "Phát hiện thẻ đóng nhóm nhưng không có thẻ mở nhóm trước đó. Hệ thống bỏ qua thẻ này."))
                    continue
                in_group = False
                current_group['end_marker'] = p
                all_blocks.append({'type': 'GROUP', 'data': current_group})
                current_group = None
                continue

            m_cau = REGEX_CAU.match(clean_text)
            if m_cau:
                flush_cau()
                absolute_q_idx += 1 
                raw_label = m_cau.group(1).strip()
                marker = "@" if raw_label.startswith("@") else ("!" if raw_label.startswith("!") else "")
                
                # Bổ sung 'marker_counts' để lưu trữ sổ Nam Tào đếm mốc
                current_cau = {
                    'label': m_cau.group(1), 'stem': [], 
                    'options': {'A': [], 'B': [], 'C': [], 'D': [], 'a': [], 'b': [], 'c': [], 'd': []},
                    'solution': [], 'key_value': "", 'curr_part': 'stem', 'last_marker': None, 'marker': marker,
                    'absolute_id': absolute_q_idx,       
                    'raw_xml': [deepcopy(p._element)],
                    'tf_keys': {'a': False, 'b': False, 'c': False, 'd': False},
                    'has_table': False,     
                    'has_floating': False,
                    'marker_counts': {'A':0, 'B':0, 'C':0, 'D':0, 'a':0, 'b':0, 'c':0, 'd':0}
                }
                
                if "<wp:anchor" in p._element.xml: current_cau['has_floating'] = True
                
                p_clone = self._clone_paragraph(p)
                self._crop_paragraph(p_clone, m_cau.end(), None)
                if self._is_content_paragraph(p_clone): current_cau['stem'].append(p_clone)
                continue

            if in_group and current_group is not None and current_cau is None:
                current_group['header_content'].append(p)
                continue

            if current_cau:
                if "<wp:anchor" in p._element.xml: current_cau['has_floating'] = True
                current_cau['raw_xml'].append(deepcopy(p._element)) 
                
                # --- [BỘ LỌC CHUẨN]: Ranh giới Lời giải / Hướng dẫn giải ---
                is_solution_boundary = False
                if current_cau['curr_part'] == 'solution':
                    is_solution_boundary = True
                else:
                    m_sol = REGEX_KEY_LINE.match(text)
                    if m_sol:
                        matched_str = m_sol.group(0)
                        remaining_text = text[len(matched_str):].strip()
                        if not remaining_text:
                            # Trường hợp 1: Nằm riêng 1 dòng (Hoàn toàn hợp lệ)
                            is_solution_boundary = True
                        else:
                            # Trường hợp 2: Nằm chung dòng văn bản -> Bắt buộc chữ phải IN ĐẬM
                            for run in p.runs:
                                if run.text.strip():
                                    # Kiểm tra thuộc tính in đậm của cụm từ đầu tiên
                                    if run.bold or (run._element.rPr is not None and run._element.rPr.b is not None):
                                        is_solution_boundary = True
                                    break 
                                    
                if is_solution_boundary:
                    current_cau['curr_part'] = 'solution'
                    current_cau['solution'].append(self._clone_paragraph(p))
                    continue
                
                raw_matches_tn = list(REGEX_PA_ANCHOR.finditer(text))
                raw_matches_ds = list(REGEX_PA_DS.finditer(text))
                raw_matches = raw_matches_tn + raw_matches_ds
                raw_matches.sort(key=lambda x: x.start())

                valid_matches = []
                for m in raw_matches:
                    end_idx = m.end()
                    is_valid = False
                    if end_idx == len(text):
                        is_valid = True
                    else:
                        next_char = text[end_idx]
                        if next_char in [' ', '\t', '\n', '\r', '\u00A0']:
                            is_valid = True
                        else:
                            curr_pos = 0; is_valid_special = False
                            for run in p.runs:
                                run_len = len(run.text)
                                if curr_pos == end_idx:
                                    xml_str = run._element.xml
                                    if 'w:drawing' in xml_str or 'm:oMath' in xml_str or 'v:shape' in xml_str or 'w:object' in xml_str or 'w:pict' in xml_str or 'mc:AlternateContent' in xml_str:
                                        is_valid_special = True; break
                                    elif run_len > 0: break 
                                elif curr_pos > end_idx: break
                                curr_pos += run_len
                            if is_valid_special or self._is_followed_by_special(p, end_idx):
                                is_valid = True
                    
                    if is_valid:
                        valid_matches.append(m)
                        lbl = m.group(1)
                        # ĐẾM SỐ LƯỢNG MỐC TRƯỚC KHI CẮT BỎ VÀO HỘP
                        if m in raw_matches_ds: current_cau['marker_counts'][lbl.lower()] += 1
                        else: current_cau['marker_counts'][lbl.upper()] += 1

                if not valid_matches:
                    if current_cau['last_marker'] is None: current_cau['stem'].append(self._clone_paragraph(p))
                    else: current_cau['options'][current_cau['last_marker']].append(self._clone_paragraph(p))
                    continue
                
                current_idx = 0
                for m in valid_matches:
                    new_marker = m.group(1) if m in raw_matches_ds else m.group(1).upper()
                    start_marker, end_marker = m.start(), m.end()
                    
                    if start_marker > current_idx:
                        segment = self._clone_paragraph(p) 
                        self._crop_paragraph(segment, current_idx, start_marker)
                        if self._is_content_paragraph(segment):
                            if current_cau['last_marker'] is None: current_cau['stem'].append(segment)
                            else: current_cau['options'][current_cau['last_marker']].append(segment)
                    
                    current_cau['last_marker'] = new_marker
                    current_idx = end_marker 
                
                segment = self._clone_paragraph(p) 
                self._crop_paragraph(segment, current_idx, None)
                if self._is_content_paragraph(segment):
                    current_cau['options'][current_cau['last_marker']].append(segment)

        flush_cau()
        if in_group:
            all_blocks.append({'type': 'GROUP', 'data': current_group})
            self.error_log.append(("Hệ thống", "Lỗi cấu trúc nhóm", "Chạm đáy file nhưng chưa đóng nhóm. Tự động chốt nhóm."))
            
        return all_blocks


# ============================================================================
# 🔄 KHU VỰC 3: TRẠM KIỂM ĐỊNH LỖI ĐA TẦNG (KIỂM ĐẾM TỪ SỔ NAM TÀO)
# ============================================================================
    def _check_dinh_chu(self, opt_box):
        if not opt_box: return False
        p = opt_box[0]
        for node in p._element.iterchildren():
            tag = node.tag.split('}')[-1]
            if tag in ['drawing', 'oMath', 'oMathPara', 'shape', 'object', 'pict', 'AlternateContent']:
                return False 
            if tag == 'r':
                for child in node.iter():
                    ctag = child.tag.split('}')[-1]
                    if ctag in ['drawing', 'object', 'pict', 'AlternateContent']: return False
                    if ctag in ['tab', 'br', 'cr']: return False 
                    if ctag == 't' and child.text:
                        first_char = child.text[0]
                        if first_char in [' ', '\t', '\n', '\r', '\u00A0']: return False 
                        else: return True 
        return False

    def _khu_vuc_3_kiem_dinh_loi(self, all_blocks):
        all_qs_for_check = []
        for block in all_blocks:
            if block['type'] == 'SINGLE': all_qs_for_check.append(block['data'])
            else: all_qs_for_check.extend(block['data']['questions'])

        for q in all_qs_for_check:
            abs_id = q.get('absolute_id', '?')
            q_name = q.get('label', f"Câu {abs_id}")
            error_list = [] 

            # --- [BỘ LỌC CHUẨN]: THEO ĐÚNG THỨ TỰ ƯU TIÊN KEY ---
            key_val = ""
            detected_type = None
            
            # Lấy toàn bộ text để xét TL và TLN nhanh
            full_raw_text = "".join([Paragraph(el, self.doc).text for el in q['raw_xml'] if not el.tag.endswith('tbl')])
            
            has_tn_markers = bool(q['marker_counts']['A'] and q['marker_counts']['B'])
            has_ds_markers = bool(q['marker_counts']['a'] and q['marker_counts']['b'])

            # BƯỚC 1: ƯU TIÊN 1 - TỰ LUẬN
            if "<Tự luận>" in full_raw_text or "<TỰ LUẬN>" in full_raw_text:
                key_val = "TL"
                detected_type = 'PHAN_IV'
            else:
                # BƯỚC 2: ƯU TIÊN 2 - TRẢ LỜI NGẮN (Chỉ nhận <key=...>, tối đa 4 ký tự)
                m_tln = re.search(r"<key=([^>]+)>", full_raw_text, re.IGNORECASE)
                if m_tln:
                    inner_val = m_tln.group(1).strip()
                    if len(inner_val) <= 4:
                        key_val = inner_val
                        detected_type = 'PHAN_III'
                    else:
                        pass # Nếu dài hơn 4 ký tự -> Không nhận diện, rơi thẳng xuống Khu Cách Ly
                        
                # BƯỚC 3 & 4: XÉT TRẮC NGHIỆM VÀ ĐÚNG SAI DỰA TRÊN MÀU ĐỎ
                if not detected_type:
                    for xml_elem in q['raw_xml']:
                        if xml_elem.tag.endswith('tbl'): continue 
                        p_temp = Paragraph(xml_elem, self.doc)
                        text = p_temp.text
                        
                        # Chỉ tìm màu đỏ ở mốc A, B, C, D (Không quét chữ KEY: A)
                        if has_tn_markers and not has_ds_markers:
                            for m in REGEX_PA_ANCHOR.finditer(text):
                                if self._check_label_highlight(p_temp, m.group(1), m.start()):
                                    if m.group(1) not in key_val: key_val += m.group(1)
                                    
                        # Lưu dữ liệu màu đỏ cho a, b, c, d
                        if has_ds_markers:
                            for m in REGEX_PA_DS.finditer(text):
                                lbl = m.group(1).lower()
                                if lbl in ['a','b','c','d']:
                                    if self._check_label_highlight(p_temp, lbl, m.start()):
                                        q['tf_keys'][lbl] = True
                                        
                    # Phân loại sau khi có dữ liệu
                    if has_tn_markers and not has_ds_markers and len(key_val) == 1 and key_val in ['A', 'B', 'C', 'D']:
                        detected_type = 'PHAN_I'
                    elif has_ds_markers:
                        detected_type = 'PHAN_II'
            
            q['key_value'] = key_val
            is_tl = (detected_type == 'PHAN_IV')
            is_tln = (detected_type == 'PHAN_III')

            is_isolated = False

            if detected_type not in ['PHAN_II', 'PHAN_III', 'PHAN_IV']:
                if not key_val:
                    error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Thiếu KEY đáp án", "Không tìm thấy KEY hợp lệ. Vui lòng tô đỏ/gạch chân phương án đúng, hoặc bổ sung thẻ <key=...>, <Tự luận>."))
                    is_isolated = True
                elif len(key_val) > 1:
                    error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Nhiều KEY đáp án", f"Phát hiện nhiều phương án ({key_val}) cùng được bôi đỏ/gạch chân. Chỉ được phép chọn 1."))
                    is_isolated = True

            if detected_type not in ['PHAN_III', 'PHAN_IV']:
                if detected_type in ['PHAN_I', 'PHAN_II'] or has_tn_markers or has_ds_markers:
                    lbls_check = ['A', 'B', 'C', 'D'] if not has_ds_markers else ['a', 'b', 'c', 'd']
                    
                    # 1. KIỂM TRA THIẾU BẰNG SỔ ĐẾM TUYỆT ĐỐI
                    missing = [lbl for lbl in lbls_check if q['marker_counts'][lbl] == 0]
                    if missing:
                        missing_str = ", ".join(missing)
                        error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Thiếu phương án", f"Không tìm thấy phương án {missing_str}."))
                        is_isolated = True
                        
                    # 2. KIỂM TRA LẶP PHƯƠNG ÁN BẰNG SỔ ĐẾM TUYỆT ĐỐI
                    duplicates = [lbl for lbl in lbls_check if q['marker_counts'][lbl] > 1]
                    if duplicates:
                        for lbl in duplicates:
                            count = q['marker_counts'][lbl]
                            error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Lỗi lặp phương án", f"Phương án {lbl} xuất hiện {count} lần. Có thể do gõ trùng hoặc định dạng sai làm dồn 2 câu thành 1."))
                            is_isolated = True
                    
                    # 3. CHẶN DOMINO: Nếu đã bị cách ly vì thiếu/lặp mốc, KHÔNG soi lỗi Dính chữ/Rỗng nữa
                    if not is_isolated:
                        for lbl in lbls_check:
                            opt_box = q['options'][lbl]
                            if opt_box:
                                if self._check_dinh_chu(opt_box):
                                    error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Lỗi dính chữ phương án", f"Ký tự bị dính sát rạt vào nhãn {lbl} (Ví dụ: {lbl}.Hàm số). Vui lòng thêm 1 dấu cách."))
                                    is_isolated = True
                                
                                is_empty = True
                                for p_box in opt_box:
                                    if self._is_content_paragraph(p_box): is_empty = False; break
                                if is_empty:
                                    error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Lỗi phương án trống", f"Nội dung của phương án {lbl} trống rỗng."))
                                    is_isolated = True
            
            if detected_type is None and not is_isolated:
                error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Lỗi cấu trúc nặng", "Hệ thống không thể nhận diện được cấu trúc của câu hỏi này. Đã chuyển xuống Khu cách ly."))
                is_isolated = True

            if q.get('has_floating'): 
                error_list.append((f"Câu thứ {abs_id} (Gốc: {q_name})", "Lỗi chèn ảnh", "Câu này chứa hình ảnh ở chế độ trôi nổi. Vui lòng chọn Wrap Text -> 'In Line with Text'."))
                is_isolated = True
            
            for err in error_list: self.error_log.append(err)
            q['is_isolated'] = is_isolated
            q['detected_type'] = detected_type

            # [PHỤC HỒI QUYỀN TÁC GIẢ]: Xóa bỏ hộp bị cắt hỏng, bốc nguyên trạng từ Sổ gốc (raw_xml)
            if detected_type in ['PHAN_III', 'PHAN_IV'] or detected_type is None:
                q['stem'] = [] 
                for xml_elem in q['raw_xml']:
                    if xml_elem.tag.endswith('tbl'):
                        class DummyBlock:
                            def __init__(self, el): self._element = el
                        q['stem'].append(DummyBlock(deepcopy(xml_elem)))
                        continue
                    
                    p_temp = Paragraph(xml_elem, self.doc)
                    if REGEX_KEY_LINE.match(p_temp.text.strip()):
                        break # Dừng lấy dữ liệu khi đụng ranh giới Lời giải
                    
                    q['stem'].append(Paragraph(deepcopy(xml_elem), self.doc))
                    
                for lbl in ['A', 'B', 'C', 'D', 'a', 'b', 'c', 'd']:
                    q['options'][lbl] = []

        buckets = {'PHAN_I': [], 'PHAN_II': [], 'PHAN_III': [], 'PHAN_IV': [], 'PHAN_LOI': []}
        for block in all_blocks:
            if block['type'] == 'SINGLE':
                q = block['data']
                if q['is_isolated']: buckets['PHAN_LOI'].append(block)
                else: buckets[q['detected_type']].append(block)
            elif block['type'] == 'GROUP':
                valid_qs = []; faulty_qs = []
                for q in block['data']['questions']:
                    if q['is_isolated']: faulty_qs.append(q)
                    else: valid_qs.append(q)
                
                if valid_qs:
                    block['data']['questions'] = valid_qs
                    dt_group = valid_qs[0]['detected_type'] 
                    buckets[dt_group].append(block)
                for fq in faulty_qs:
                    buckets['PHAN_LOI'].append({'type': 'SINGLE', 'data': fq})
                    
        return buckets


# ============================================================================
# 🔒 KHU VỰC 4: VẼ BẢNG CHUẨN HÓA (NIÊM PHONG - CẤM XÂM PHẠM)
# ============================================================================
    def _khu_vuc_4_ve_bang_chuan_hoa(self, buckets):
        body = self.doc._body._element
        sectPr = None
        for child in list(body):
            if child.tag.endswith('sectPr'): sectPr = child
            else: body.remove(child)

        part_order = ['PHAN_I', 'PHAN_II', 'PHAN_III', 'PHAN_IV']
        for phan_key in part_order:
            blocks_in_phan = buckets[phan_key]
            if not blocks_in_phan: continue 
            local_count = 0
            for b in blocks_in_phan:
                if b['type'] == 'SINGLE': local_count += 1
                else: local_count += len(b['data']['questions'])
            self.stats[phan_key] = local_count
            s = 1; e = local_count
            template_info = STD_TEMPLATE.get(phan_key)
            if template_info:
                p = self.doc.add_paragraph()
                run_bold = p.add_run(template_info["bold"]); run_bold.bold = True; run_bold.font.color.rgb = COLOR_BLUE; run_bold.font.name = "Times New Roman"; run_bold.font.size = Pt(12)
                if local_count > 1:
                    run_normal = p.add_run(template_info["normal"].format(s=s, e=e)); run_normal.bold = False; run_normal.font.color.rgb = COLOR_BLUE; run_normal.font.name = "Times New Roman"; run_normal.font.size = Pt(12)
            
            current_table = self.doc.add_table(rows=0, cols=2)
            self._set_table_grid(current_table)
            self._set_col_width(current_table)
            
            local_q_idx = 1
            for block in blocks_in_phan:
                if block['type'] == 'GROUP':
                    self.doc.add_paragraph(block['data']['start_marker'].text)
                    if block['data'].get('header_content'):
                        header_paras = block['data']['header_content']
                        if any(p.text.strip() for p in header_paras):
                            tbl_p = self.doc.add_table(rows=0, cols=2)
                            self._set_table_grid(tbl_p); self._set_col_width(tbl_p)
                            row = tbl_p.add_row()
                            self._format_label_cell(row.cells[0], "p")
                            cell_content = row.cells[1]; self._wipe(cell_content._element)
                            for p_source in header_paras: cell_content._element.append(deepcopy(p_source._element))
                            self._trim_cell_safe(cell_content)
                            spacer = self.doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(1)
                    current_table = self.doc.add_table(rows=0, cols=2)
                    self._set_table_grid(current_table); self._set_col_width(current_table)

                questions = [block['data']] if block['type'] == 'SINGLE' else block['data']['questions']
                for q in questions:
                    prefix = q.get('marker', ''); new_label = f"{prefix}Câu {local_q_idx}"; local_q_idx += 1
                    if phan_key == 'PHAN_I':
                        r1 = current_table.add_row(); self._format_label_cell(r1.cells[0], new_label)
                        for p in q['stem']: 
                            new_p = deepcopy(p._element)
                            if self._sanitize_element(new_p) is not None: r1.cells[1]._element.append(new_p) 
                        self._trim_cell_safe(r1.cells[1], REGEX_CAU)
                        for lbl in ['A', 'B', 'C', 'D']:
                            r = current_table.add_row(); self._format_label_cell(r.cells[0], lbl)
                            for p in q['options'][lbl]: 
                                new_p = deepcopy(p._element)
                                if self._sanitize_element(new_p) is not None: r.cells[1]._element.append(new_p) 
                            self._trim_cell_safe(r.cells[1])
                        rk = current_table.add_row(); k_txt = f"KEY: {q['key_value']}"; self._format_label_cell(rk.cells[0], k_txt, COLOR_RED)
                        for p in q['solution']: 
                            new_p = deepcopy(p._element)
                            if self._sanitize_element(new_p) is not None: rk.cells[1]._element.append(new_p) 
                        self._clean_solution_cell(rk.cells[1])
                    elif phan_key == 'PHAN_II':
                        r1 = current_table.add_row(); self._format_label_cell(r1.cells[0], new_label)
                        for p in q['stem']: 
                            new_p = deepcopy(p._element)
                            if self._sanitize_element(new_p) is not None: r1.cells[1]._element.append(new_p) 
                        self._trim_cell_safe(r1.cells[1], REGEX_CAU)
                        map_lbl = {'a': 'a)', 'b': 'b)', 'c': 'c)', 'd': 'd)'}
                        for lbl in ['a', 'b', 'c', 'd']:
                            r = current_table.add_row(); self._format_label_cell(r.cells[0], map_lbl[lbl])
                            for p in q['options'][lbl]: 
                                new_p = deepcopy(p._element)
                                if self._sanitize_element(new_p) is not None: r.cells[1]._element.append(new_p) 
                            self._trim_cell_safe(r.cells[1])
                        rk = current_table.add_row()
                        k_val = self._generate_tf_key(q); k_txt = f"KEY: {k_val}"
                        self._format_label_cell(rk.cells[0], k_txt, COLOR_RED)
                        for p in q['solution']: 
                            new_p = deepcopy(p._element)
                            if self._sanitize_element(new_p) is not None: rk.cells[1]._element.append(new_p) 
                        self._clean_solution_cell(rk.cells[1])
                    else: 
                        r1 = current_table.add_row(); self._format_label_cell(r1.cells[0], new_label)
                        for p in q['stem']: 
                            new_p = deepcopy(p._element)
                            if self._sanitize_element(new_p) is not None: r1.cells[1]._element.append(new_p) 
                        self._trim_cell_safe(r1.cells[1], REGEX_CAU)
                        rk = current_table.add_row(); k_val = "TL" if phan_key == 'PHAN_IV' else q['key_value']; k_txt = f"KEY: {k_val}"
                        self._format_label_cell(rk.cells[0], k_txt, COLOR_RED)
                        for p in q['solution']: 
                            new_p = deepcopy(p._element)
                            if self._sanitize_element(new_p) is not None: rk.cells[1]._element.append(new_p) 
                        self._clean_solution_cell(rk.cells[1])
                    spacer = self.doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(1)
                    current_table = self.doc.add_table(rows=0, cols=2); self._set_table_grid(current_table); self._set_col_width(current_table)
                if block['type'] == 'GROUP' and block['data']['end_marker']:
                    self.doc.add_paragraph(block['data']['end_marker'].text)
                    current_table = self.doc.add_table(rows=0, cols=2); self._set_table_grid(current_table); self._set_col_width(current_table)

        if sectPr: self.doc._body._element.append(sectPr)

    def _khu_vuc_4_ve_bang_loi_va_cach_ly(self, buckets):
        unique_err_qs = set([err[0] for err in self.error_log])
        self.stats["LOI"] = len(unique_err_qs) 
        
        if self.error_log:
            try:
                self.doc.add_page_break()
                p_title = self.doc.add_paragraph(); p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_title.add_run("BẢNG TỔNG HỢP LỖI CẤU TRÚC ĐỀ GỐC"); run.bold = True; run.font.color.rgb = COLOR_RED; run.font.size = Pt(14)
                
                p_note = self.doc.add_paragraph()
                run_note = p_note.add_run("Chú ý: Có thể số thứ tự câu không khớp thông báo do đề gốc chưa được đánh đúng thứ tự. Vui lòng dựa vào nội dung câu để đối chiếu.")
                run_note.italic = True; run_note.font.color.rgb = COLOR_RED
                
                err_table = self.doc.add_table(rows=1, cols=3); self._set_table_grid(err_table)
                hdr_cells = err_table.rows[0].cells
                hdr_cells[0].paragraphs[0].add_run('Vị trí phát hiện').bold = True
                hdr_cells[1].paragraphs[0].add_run('Phân loại lỗi').bold = True
                hdr_cells[2].paragraphs[0].add_run('Hướng dẫn khắc phục').bold = True
                            
                for error in self.error_log:
                    row_cells = err_table.add_row().cells
                    row_cells[0].paragraphs[0].add_run(error[0]).font.color.rgb = COLOR_RED
                    row_cells[1].paragraphs[0].add_run(error[1]).font.color.rgb = COLOR_RED
                    row_cells[2].paragraphs[0].add_run(error[2]).font.color.rgb = COLOR_RED
            except: pass 

        if buckets['PHAN_LOI']:
            try:
                self.doc.add_paragraph() 
                p_loi_title = self.doc.add_paragraph(); p_loi_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_loi = p_loi_title.add_run("DANH SÁCH CÁC CÂU LỖI BỊ CÁCH LY"); run_loi.bold = True; run_loi.font.color.rgb = COLOR_RED; run_loi.font.size = Pt(14)
                
                p_note_loi = self.doc.add_paragraph()
                run_note_loi = p_note_loi.add_run("Các câu dưới đây bị thiếu dữ liệu trọng yếu hoặc sai định dạng. Nội dung gốc được giữ nguyên bên dưới để dễ dàng đối chiếu:")
                run_note_loi.italic = True; run_note_loi.font.color.rgb = COLOR_RED
                
                for block in buckets['PHAN_LOI']:
                    q = block['data']; p_label = self.doc.add_paragraph()
                    q_label_name = q.get('label', "Câu lỗi"); abs_id_val = q.get('absolute_id', '?')
                    r_label = p_label.add_run(f"--- Nội dung gốc của Câu thứ {abs_id_val} ({q_label_name}) ---"); r_label.bold = True; r_label.font.color.rgb = COLOR_RED
                    
                    # 1. Quét sổ đếm để tìm tất cả các phương án bị trùng (Cả A,B,C,D và a,b,c,d)
                    duplicates = [lbl for lbl in ['A', 'B', 'C', 'D', 'a', 'b', 'c', 'd'] if q.get('marker_counts', {}).get(lbl, 0) > 1]
                    
                    for xml_elem in q['raw_xml']: 
                        self.doc._body._element.append(deepcopy(xml_elem))
                        
                        # 2. Chiếu tia laser tìm tọa độ và tô xanh
                        if duplicates:
                            last_elem = self.doc._body._element[-1]
                            paras_to_check = []
                            if last_elem.tag.endswith('p'):
                                paras_to_check.append(Paragraph(last_elem, self.doc))
                            elif last_elem.tag.endswith('tbl'):
                                from docx.table import Table
                                paras_to_check.extend([p for row in Table(last_elem, self.doc).rows for cell in row.cells for p in cell.paragraphs])
                                
                            for new_p in paras_to_check:
                                text = new_p.text
                                
                                # Gộp cả 2 bộ Regex để quét song song (Trắc nghiệm và Đúng/Sai)
                                for reg_obj in [REGEX_PA_ANCHOR, REGEX_PA_DS]:
                                    for m in reg_obj.finditer(text):
                                        lbl = m.group(1) 
                                        if lbl in duplicates:
                                            start_idx = m.start()
                                            end_idx = m.end()
                                            curr_pos = 0
                                            for run in new_p.runs:
                                                run_len = len(run.text)
                                                # Nếu đoạn văn bản nằm khớp tọa độ của chữ bị lặp
                                                if max(curr_pos, start_idx) < min(curr_pos + run_len, end_idx):
                                                    run.font.color.rgb = COLOR_BLUE
                                                    run.font.bold = True
                                                curr_pos += run_len
                                                
                    spacer = self.doc.add_paragraph(); spacer.paragraph_format.space_after = Pt(1)
            except: pass

# ============================================================================
# 🔄 KHU VỰC 5: TRẠM ĐIỀU KHIỂN & TÁCH LUỒNG (XUẤT FILE)
# ============================================================================
    def xu_ly(self, input_source):
        self.doc = Document(input_source)
        self._setup_page_layout()
        self._clear_footers()
        
        # [MỚI THÊM] Tiền xử lý Phẫu thuật Shift+Enter an toàn (Ép cứng hoa/thường)
        self._tien_xu_ly_shift_enter()

        all_blocks = self._khu_vuc_2_nhan_dang_du_lieu()
        buckets = self._khu_vuc_3_kiem_dinh_loi(all_blocks)
        self._khu_vuc_4_ve_bang_chuan_hoa(buckets)

        stream_sach = io.BytesIO()
        self.doc.save(stream_sach)
        stream_sach.seek(0)

        self._khu_vuc_4_ve_bang_loi_va_cach_ly(buckets)
        stream_vat_ly = io.BytesIO()
        self.doc.save(stream_vat_ly)
        stream_vat_ly.seek(0)

        return stream_sach, stream_vat_ly

def xu_ly_va_chuan_hoa(file_goc):
    proc = XuLyDeChuanHoa()
    stream_sach, stream_vat_ly = proc.xu_ly(file_goc)
    return stream_sach, stream_vat_ly, proc.stats, proc.stats["LOI"]